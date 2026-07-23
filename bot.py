from __future__ import annotations
import asyncio
import os
import json
import urllib.request
import urllib.error
import gspread
from google.oauth2.service_account import Credentials
from aiogram import Bot, Dispatcher, types, F
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.filters import Command
from aiogram.utils.keyboard import InlineKeyboardBuilder
from docx import Document
from datetime import datetime, timedelta
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import calendar
import subprocess
import time
from aiogram.exceptions import TelegramBadRequest
from docxcompose.composer import Composer
import re
import tempfile
from pathlib import Path
import requests
import socket
from dotenv import load_dotenv
from aiogram.client.session.aiohttp import AiohttpSession
from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode
import aiohttp

# Load environment variables from .env file
load_dotenv()

try:
    from aiohttp_socks import ProxyConnector
except ImportError:
    ProxyConnector = None

# Try to import Google Generative AI SDK (google-genai)
try:
    from google import genai
except Exception:
    genai = None

# ---------------- НАСТРОЙКИ ---------------- #
BOT_TOKEN = os.getenv("BOT_TOKEN", "8374018707:AAEqM0IO-SKvWYCufdqyMGih5WklxpuYoC4")
GOOGLE_SHEET_URL1 = os.getenv("GOOGLE_SHEET_URL1", "https://docs.google.com/spreadsheets/d/16S4ImCcjoBjIMyUMmv4bdeAYGMPzVcGv5g_yMJaxDIY/edit?usp=sharing")
GOOGLE_SHEET_URL2 = os.getenv("GOOGLE_SHEET_URL2", "https://docs.google.com/spreadsheets/d/1lv0LSgPFRuBT6ELjTza2M-D_Tc21H9gTRNTEs4SbFlI/edit?usp=sharing")
GOOGLE_PHONE_GVIZ_URL = os.getenv("GOOGLE_PHONE_GVIZ_URL", "https://docs.google.com/spreadsheets/d/16S4ImCcjoBjIMyUMmv4bdeAYGMPzVcGv5g_yMJaxDIY/gviz/tq?tqx=out:json&tq=select%20A,B")
# Налаштування проксі (підтримка PythonAnywhere та .env)
# Пріоритет: .env PROXY_URL -> системний http_proxy -> стандартний для PythonAnywhere
PROXY_URL = os.getenv("PROXY_URL") or os.environ.get("http_proxy") or os.environ.get("https_proxy")
if not PROXY_URL and os.environ.get("PYTHONANYWHERE_DOMAIN"):
    PROXY_URL = "http://proxy.server:3128"
SCOPES = ["https://www.googleapis.com/auth/spreadsheets.readonly"]

# Resolve credentials.json relative to this script directory, allow override by env
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CREDENTIALS_FILE = os.getenv("GOOGLE_CREDENTIALS_FILE", os.path.join(BASE_DIR, "credentials.json"))

EXECUTORS = {
    "dilin": "Дилін.docx",
    "tsarenko": "Царенко.docx",
    "fedun": "Федун.docx",
    "pavlenko": "Павленко.docx"
}

PIDPYS = {
    "korzhov": "Коржов.docx",
    "usyk": "Усик.docx",
    "sereda": "Середа.docx"
}

if os.name == 'nt':  # Windows
    LIBREOFFICE_PATH = r"C:\Program Files\LibreOffice\program\soffice.exe"
else:  # Linux (PythonAnywhere и др.)
    # Стандартные пути для Linux систем
    _possible_paths = ["/usr/bin/soffice", "/usr/bin/libreoffice", "/usr/local/bin/soffice"]
    LIBREOFFICE_PATH = next((p for p in _possible_paths if os.path.exists(p)), "/usr/bin/soffice")

# --- Настройки временных файлов --- #
TEMP_DIR = Path(tempfile.gettempdir()) / "bot_address_cache"
TEMP_DIR.mkdir(parents=True, exist_ok=True)
ADDRESS_REQUEST_TTL = 300  # 5 минут для завершения ввода адреса

# ------------------------------------------- #
CACHE_TTL_SECONDS = 300
SHEET_CACHE = {}

# --- Логи и статус кеша/онлайн --- #
DEBUG_LOG = False  # при необходимости можно включить подробные логи в консоль

def _log_debug(message: str):
    if DEBUG_LOG:
        print(message)

# Флаги состояния кеша/мережі для отображения статуса пользователю
CACHE_READY = False
CACHE_LAST_ERROR: str | None = None
CACHE_PROGRESS: int = 0  # 0-100, приблизний прогрес прогріву кеша

# Multiple API keys support (from .env: GEMINI_API_KEY, GEMINI_API_KEY_2, GEMINI_API_KEY_3, ...)
def _get_gemini_api_keys():
    """Отримати список API ключів Gemini з .env."""
    keys = []
    # Primary key
    primary = os.getenv("GEMINI_API_KEY", "")
    if primary:
        keys.append(primary)
    # Additional keys: GEMINI_API_KEY_2, GEMINI_API_KEY_3, ...
    for i in range(2, 11):
        key = os.getenv(f"GEMINI_API_KEY_{i}", "")
        if key:
            keys.append(key)
    return keys if keys else [GEMINI_API_KEY]

# Rate limiting
_GEMINI_KEY_INDEX = 0
_GEMINI_LAST_CALL = 0.0
_GEMINI_RATE_LIMIT = 1.5  # seconds between calls per key

def _call_gemini_decline_one(name: str, api_key: str) -> dict:
    """Відмінює ОДНЕ ПІБ через Gemini REST API. Повертає {H, P, U, V} або {}."""
    # Компактний промпт для одного імені
    prompt = (
        f"Відміняй українське ПІБ «{name}» (називний відмінок).\n"
        "Виведи РІВНО один рядок у форматі:\n"
        "оригінал|H|P|U|V\n"
        "де:\n"
        "H = коротка форма родовий (кого?): Ім'я ПРІЗВИЩЕ — наприклад «Ігоря АНДРОСЕНКА»\n"
        "P = коротка форма давальний (кому?): Ім'я ПРІЗВИЩЕ — наприклад «Ігорю АНДРОСЕНКУ»\n"
        "U = повна форма давальний: ПРІЗВИЩЕ Ім'я По-батькові — наприклад «АНДРОСЕНКУ Ігорю Сергійовичу»\n"
        "V = повна форма родовий: ПРІЗВИЩЕ Ім'я По-батькові — наприклад «АНДРОСЕНКА Ігоря Сергійовича»\n"
        "Прізвище ВЕЛИКИМИ літерами. Ім'я та по батькові з великої літери.\n"
        "Нічого крім рядка з | не виводь."
    )
    payload = json.dumps({
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.1, "maxOutputTokens": 200}
    }).encode("utf-8")

    last_error = ""
    for attempt in range(1):  # one attempt only, fallback handles the rest
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash-lite:generateContent?key={api_key}"
            req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json"})
            with urllib.request.urlopen(req, timeout=20) as resp:
                obj = json.loads(resp.read().decode("utf-8"))
            # Check for error in response
            if "error" in obj:
                last_error = f"API error: {obj['error']}"
                print(f"[GEMINI] {last_error}")
                if "429" in str(obj.get("error", {}).get("code", "")):
                    time.sleep((attempt + 1) * 10)
                    continue
                return {}
            for cand in (obj.get("candidates") or []):
                for p in ((cand.get("content") or {}).get("parts") or []):
                    text = (p.get("text") or "").strip()
                    if '|' in text:
                        parts = text.split('|')
                        if len(parts) >= 5:
                            return {"H": parts[1].strip(), "P": parts[2].strip(),
                                    "U": parts[3].strip(), "V": parts[4].strip()}
            last_error = f"empty response: {str(obj)[:200]}"
            print(f"[GEMINI] No valid response: {last_error}")
            return {}
        except urllib.error.HTTPError as e:
            last_error = f"HTTP {e.code}: {e.reason}"
            print(f"[GEMINI] {last_error} (attempt {attempt+1}/3)")
            if e.code == 429 and attempt < 2:
                time.sleep((attempt + 1) * 10)
            elif e.code in (400, 403):
                return {}
            elif attempt < 2:
                time.sleep(3)
        except Exception as e:
            last_error = str(e)[:200]
            print(f"[GEMINI] Exception: {last_error}")
            if attempt < 2:
                time.sleep(3)
    return {}


# --- Gemini AI (для генерації варіантів 'що саме') --- #
GEMINI_PROMPT_SHO = (
    "Ти — AI-помічник для створення преамбул до наказів про залучення підрозділів ДСНС для гасіння пожеж та виконання робіт. Користувач надасть дані про інцидент у форматі, наприклад: 'м. Чернігів вул. Першотравнева, 13 внаслідок влучання БПЛА горить будинок 9 ДПРЧ'. Твоє завдання — переписати лише опис інциденту формальною діловою мовою українською, додаючи перед основним текстом на вибір по ситуації або свій варіант по контексту присланого повідомлення одну з фраз: «із застосуванням ракет невстановленого типу», «із застосуванням ударних дронів невстановленого типу», «із застосуванням засобів вогневого ураження невстановленого типу», починаючи з 'здійснено ураження (чого саме) за адресою: ' та вказуючи повну адресу у форматі: ' область, район, громада, населений пункт, вулиця [вулиця], будинок [номер]. 'Використовуєм Чернігівську область Чернігівській та Корюківський райони' 'Виключай будь-які згадки підрозділів ДСНС, такі як ДПРЧ та ДПРП' або інші. Не додавай дату чи інформацію про збройні сили — тільки генеровану частину в лапках. Виводь лише результат, без зайвих коментарів. не використовуй лапки у відповідях: 'Ти — AI-помічник для створення преамбул до наказів про залучення підрозділів ДСНС для гасіння пожеж та виконання робіт. Користувач надасть дані про інцидент у форматі, наприклад: 'м. Чернігів вул. Першотравнева, 13 внаслідок влучання БПЛА горить будинок 9 ДПРЧ'. Твоє завдання — переписати лише опис інциденту формальною діловою мовою українською, починаючи з 'здійснено ураження (чого саме) за адресою: ' та вказуючи повну адресу у форматі: ' область, район, територіальна (міська) громада, населений пункт, вулиця [вулиця], будинок [номер] (якщо у повідомленні немає номеру будинку нічого не ставимо не вигадуємо). 'Використовуєм Чернігівську область Чернігівській та Корюківський райони' 'Виключай будь-які згадки підрозділів ДСНС, такі як ДПРЧ та ДПРП' або інші. Не додавай дату чи інформацію про збройні сили — тільки генеровану частину в лапках. Виводь лише результат, без зайвих коментарів. не використовуй лапки у відповідях' до складу Чернігівського району входять: Чернігівська, Городнянська, Остерська міські, Березнянська, Гончарівська, Деснянська, Добрянська, Козелецька, Куликівська, Любецька, Михайло-Коцюбинська, Олишівська, Ріпкинська, Седнівська селищні та Іванівська, Кіптівська, Киїнська, Киселівська, Новобілоуська, Тупичівська сільські територіальні громади. до складу Корюківського району входять: Корюківська, Менська, Сновська міські та Сосницька і Холминська селищні територіальні громади"
)

# API ключ можна задати через змінну середовища GEMINI_API_KEY
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyCsatzqBbKHV7wjiJwQ-zGIzl0PQNGZwY4")

# Cache for selected model
_GEM_MODEL = None
_GEM_MODEL_NAME = None

# --- Cache for gspread Client --- #
_GSPREAD_CLIENT = None

def get_gspread_client():
    global _GSPREAD_CLIENT
    if _GSPREAD_CLIENT is None:
        creds = Credentials.from_service_account_file(CREDENTIALS_FILE, scopes=SCOPES)
        _GSPREAD_CLIENT = gspread.authorize(creds)
    return _GSPREAD_CLIENT

# --- Функции для работы с временными файлами адресов --- #
def get_address_request_file(uid: int) -> Path:
    """Получить путь к временному файлу запроса адреса для пользователя"""
    return TEMP_DIR / f"address_request_{uid}.json"

def save_address_request(uid: int, sho_text: str, context: dict = None):
    """Сохранить запрос адреса во временный файл"""
    request_data = {
        'user_id': uid,
        'sho_text': sho_text,
        'timestamp': time.time(),
        'context': context or {}
    }
    
    file_path = get_address_request_file(uid)
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(request_data, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"Ошибка сохранения запроса адреса: {e}")
        return False

def load_address_request(uid: int) -> dict:
    """Загрузить запрос адреса из временного файла"""
    file_path = get_address_request_file(uid)
    try:
        if file_path.exists():
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            # Проверить TTL
            if time.time() - data.get('timestamp', 0) < ADDRESS_REQUEST_TTL:
                return data
            else:
                # Удалить просроченный файл
                cleanup_address_request(uid)
    except Exception as e:
        print(f"Ошибка загрузки запроса адреса: {e}")
    return None

def fetch_phone_mapping():
    try:
        resp = requests.get(GOOGLE_PHONE_GVIZ_URL, timeout=15)
        resp.raise_for_status()
    except Exception:
        return {}
    text = resp.text
    start = text.find('{')
    end = text.rfind('}')
    if start == -1 or end == -1:
        return {}
    try:
        data = json.loads(text[start:end+1])
    except json.JSONDecodeError:
        return {}
    mapping = {}
    for row in data.get('table', {}).get('rows', []):
        cells = row.get('c') or []
        name_cell = cells[1] if len(cells) > 1 else None
        phone_cell = cells[0] if cells else None
        name = (name_cell or {}).get('v')
        if not name:
            continue
        key = _normalize_lookup_key(name)
        if not key:
            continue
        phone_val = ""
        if phone_cell:
            phone_val = str(phone_cell.get('f') or phone_cell.get('v') or '').strip()
        mapping[key] = phone_val
    return mapping

def get_phone_mapping(force_refresh=False):
    key = "phones"
    now_ts = time.time()
    entry = SHEET_CACHE.get(key)
    if (not force_refresh) and entry:
        data, ts = entry
        if now_ts - ts < CACHE_TTL_SECONDS:
            return data
    mapping = fetch_phone_mapping()
    SHEET_CACHE[key] = (mapping, now_ts)
    return mapping

def cleanup_address_request(uid: int):
    """Очистить временный файл запроса адреса"""
    file_path = get_address_request_file(uid)
    try:
        if file_path.exists():
            file_path.unlink()
            return True
    except Exception as e:
        print(f"Ошибка удаления запроса адреса: {e}")
    return False

def _select_gemini_model() -> None:
    """Initialize global genai.Client (new SDK)."""
    global _GEM_MODEL, _GEM_MODEL_NAME
    if _GEM_MODEL is not None:
        return
    if not GEMINI_API_KEY or genai is None:
        return
    try:
        # Створюємо клієнт нового SDK
        client = genai.Client(api_key=GEMINI_API_KEY)
        _GEM_MODEL = client
        # Використовуємо 2.5 Flash Lite як стандарт
        _GEM_MODEL_NAME = "gemini-2.5-flash-lite"
    except Exception as e:
        print(f"⚠️ Помилка ініціалізації Gemini SDK: {e}")
        _GEM_MODEL = None
        _GEM_MODEL_NAME = None

def _get_user_model(uid: int) -> str:
    """Отримати модель, вибрану користувачем, або стандартну."""
    return user_state.get(uid, {}).get("ai_model", "gemini-2.5-flash-lite")

def _call_gemini_generate_sho(seed_text: str, izchogo_text: str, uid: int = 0) -> list:
    """Generate 2-3 SHO suggestions using SDK if available, else REST fallback."""
    model_to_use = _get_user_model(uid) if uid else "gemini-2.5-flash-lite"
    # Compose content once
    context_parts = [GEMINI_PROMPT_SHO]
    if izchogo_text:
        context_parts.append(f"Контекст (із чого): {izchogo_text}")
    if seed_text:
        context_parts.append(f"Деталі від користувача: {seed_text}")
    composed = "\n\n".join(context_parts)

    # Try SDK first
    if GEMINI_API_KEY and genai is not None:
        try:
            _select_gemini_model()
            if _GEM_MODEL is not None:
                # Новий синтаксис: client.models.generate_content
                resp = _GEM_MODEL.models.generate_content(
                    model=model_to_use,
                    contents=composed
                )
                full_text = ""
                try:
                    full_text = (getattr(resp, "text", None) or "").strip()
                except Exception:
                    full_text = ""
                texts = [full_text] if full_text else []
                # Parse lines
                lines = []
                for t in texts:
                    for ln in (t.splitlines() if t else []):
                        ln = ln.strip().lstrip("-•·0123456789. ")
                        if ln:
                            lines.append(ln)
                # Deduplicate top 3
                out, seen = [], set()
                for ln in lines:
                    k = ln.lower()
                    if k not in seen:
                        seen.add(k)
                        out.append(ln)
                    if len(out) >= 3:
                        break
                if out:
                    return out
        except Exception:
            pass

    # REST fallback
    try:
        model = model_to_use
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={GEMINI_API_KEY}"
            payload = {
                "contents": [
                    {"parts": [{"text": composed}]}
                ]
            }
            data = json.dumps(payload).encode("utf-8")
            req = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"})
            with urllib.request.urlopen(req, timeout=30) as resp:
                body = resp.read().decode("utf-8")
                obj = json.loads(body)
            texts = []
            for cand in (obj.get("candidates") or []):
                parts = (((cand.get("content") or {}).get("parts")) or [])
                for p in parts:
                    t = (p.get("text") or "").strip()
                    if t:
                        texts.append(t)
                lines = []
                for t in texts:
                    for ln in t.splitlines():
                        ln = ln.strip().lstrip("-•·0123456789. ")
                        if ln:
                            lines.append(ln)
                out, seen = [], set()
                for ln in lines:
                    k = ln.lower()
                    if k not in seen:
                        seen.add(k)
                        out.append(ln)
                    if len(out) >= 3:
                        break
                if out:
                    return out
        except Exception as e:
            print(f"⚠️ REST AI error: {e}")
            return []
    except Exception:
        return []

def _call_gemini_extract_address_improved(sho_text: str, uid: int = 0) -> str:
    """Улучшенная функция извлечения адреса с повторными попытками и обработкой ошибок"""
    if not sho_text:
        return ""
    
    model_to_use = _get_user_model(uid) if uid else "gemini-2.5-flash-lite"
    # Более точный промпт для извлечения адреса
    prompt = (
        "Вихідний текст: \n" + sho_text.strip() + "\n\n"
        "Завдання: Виділи лише адресу з тексту та поверни рівно одну фразу у форматі: \"за адресою: <повна адреса>\".\n"
        "Формат адреси: область, район, громада, населений пункт, вулиця [вулиця], будинок [номер].\n"
        "Використовуй Чернігівську область, Чернігівський та Корюківський райони.\n"
        "Приклад: 'за адресою: Чернігівська область, Чернігівський район, Чернігівська міська громада, м. Чернігів, вул. Першотравнева, 13'\n"
        "Якщо номер будинку не вказано - не вигадуй.\n"
        "Не додавай жодних інших слів, пояснень чи переносів рядка."
        "Не вживай підрозділи, які залучаються (ДПРП, ДПРЧ, ПСО, МПК)."
    )
    
    # Try SDK first
    try:
        if GEMINI_API_KEY and genai is not None:
            _select_gemini_model()
            if _GEM_MODEL is not None:
                # Новий синтаксис: client.models.generate_content
                resp = _GEM_MODEL.models.generate_content(
                    model=model_to_use,
                    contents=prompt
                )
                try:
                    text = (getattr(resp, "text", None) or "").strip()
                    # Проверим, что ответ содержит ключевые слова адреса
                    if "за адресою:" in text.lower():
                        return text
                except Exception:
                    pass
    except Exception:
        pass
    
    # REST fallback
    try:
        model = model_to_use
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={GEMINI_API_KEY}"
        payload = {"contents": [{"parts": [{"text": prompt}]}]}
        data = json.dumps(payload).encode("utf-8")
        req = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=30) as resp:
            body = resp.read().decode("utf-8")
            obj = json.loads(body)
        texts = []
        for cand in (obj.get("candidates") or []):
            parts = (((cand.get("content") or {}).get("parts")) or [])
            for p in parts:
                t = (p.get("text") or "").strip()
                if t and "за адресою:" in t.lower():
                    texts.append(t)
        return texts[0] if texts else ""
    except Exception:
        return ""

def _call_gemini_chat_simple(user_text: str, uid: int = 0) -> str:
    """Simple chat call to Gemini. Returns plain text or empty string on failure."""
    if not user_text:
        return ""
    model_to_use = _get_user_model(uid) if uid else "gemini-2.5-flash-lite"
    # Try SDK
    try:
        if GEMINI_API_KEY and genai is not None:
            _select_gemini_model()
            if _GEM_MODEL is not None:
                # Новий синтаксис: client.models.generate_content
                resp = _GEM_MODEL.models.generate_content(
                    model=model_to_use,
                    contents=user_text
                )
                try:
                    text = (getattr(resp, "text", None) or "").strip()
                except Exception:
                    text = ""
                return text
    except Exception:
        pass
    # REST fallback
    try:
        model = model_to_use
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={GEMINI_API_KEY}"
        payload = {"contents": [{"parts": [{"text": user_text}]}]}
        data = json.dumps(payload).encode("utf-8")
        req = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=30) as resp:
            body = resp.read().decode("utf-8")
            obj = json.loads(body)
        texts = []
        for cand in (obj.get("candidates") or []):
            parts = (((cand.get("content") or {}).get("parts")) or [])
            for p in parts:
                t = (p.get("text") or "").strip()
                if t:
                    texts.append(t)
        return texts[0] if texts else ""
    except Exception:
        return ""

# --- Фабрика бота (для безпечного перезапуску) --- #
def create_bot():
    print(f"Ініціалізація нового з'єднання... (Проксі: {PROXY_URL or 'відсутній'})")
    
    current_session = None
    if PROXY_URL:
        t_val = 60.0  # Суворо число float для уникнення TypeError в aiogram
        if PROXY_URL.startswith("socks") and ProxyConnector:
            try:
                connector = ProxyConnector.from_url(PROXY_URL)
                current_session = AiohttpSession(connector=connector, timeout=t_val)
            except Exception as e:
                print(f"⚠️ Помилка SOCKS: {e}. Перехід на пряме з'єднання.")
                current_session = AiohttpSession(timeout=t_val)
        else:
            current_session = AiohttpSession(proxy=PROXY_URL, timeout=t_val)
    else:
        current_session = AiohttpSession(timeout=60.0)

    return Bot(
        token=BOT_TOKEN,
        session=current_session,
        default=DefaultBotProperties(parse_mode=ParseMode.HTML)
    )

# Глобальний диспетчер залишається один
dp = Dispatcher()
user_state = {}

def _get_required_addresses(uid: int) -> int:
    return 1

def _get_addresses_list(uid: int) -> list:
    return user_state.setdefault(uid, {}).setdefault("addresses_list", [])

def _format_addresses_block(addresses: list) -> str:
    if not addresses:
        return ""
    return addresses[0] if addresses else ""

async def _reply_message(target, text: str, reply_markup=None):
    if isinstance(target, types.CallbackQuery):
        await target.message.answer(text, reply_markup=reply_markup)
    else:
        await target.answer(text, reply_markup=reply_markup)

async def _prompt_manual_address_input(target, uid: int):
    kb = InlineKeyboardBuilder()
    kb.button(text="⏭️ Пропустити", callback_data="skip_adresa")
    kb.button(text="🔙 Назад", callback_data="back_to_pid_units")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    user_state.setdefault(uid, {})["state"] = "enter_adresa"
    await _reply_message(target, "Введіть адресу:", kb.as_markup())

async def _finalize_address_collection(target, uid: int):
    user_state.setdefault(uid, {})
    user_state[uid]["kgp_mode"] = True
    user_state[uid]["dodatok"] = "1"
    user_state[uid]["unit_order"] = []
    user_state[uid]["selected_indices_by_unit"] = {}
    user_state[uid]["quantities_by_unit"] = {}
    addresses_text = _format_addresses_block(_get_addresses_list(uid))
    if addresses_text:
        await _reply_message(target, f"✅ Усі адреси збережено:\n{addresses_text}")
    else:
        await _reply_message(target, "✅ Адресу збережено")
    await show_unit_selection(target, uid)

async def _append_address_and_continue(target, uid: int, address_text: str):
    text = (address_text or "").strip()
    if not text:
        return False
    addresses = _get_addresses_list(uid)
    addresses.append(text)
    collected = len(addresses)
    required = _get_required_addresses(uid)
    user_state.setdefault(uid, {})["state"] = None
    cleanup_address_request(uid)
    if collected >= required:
        user_state[uid]["adresa"] = _format_addresses_block(addresses)
        await _finalize_address_collection(target, uid)
    else:
        remaining = required - collected
        await _reply_message(
            target,
            f"✅ Адресу {collected} збережено. Залишилось додати {remaining}.")
        await _prompt_manual_address_input(target, uid)
    return True

# --- Допоміжні функції --- #
def is_authenticated(uid):
    # Password protection disabled: everyone is authenticated
    return True

def touch_activity(uid: int):
    if uid not in user_state:
        user_state[uid] = {'authenticated': False, 'state': None}
    user_state[uid]['last_activity'] = time.time()

def clear_workflow_state(uid: int):
    if uid not in user_state:
        return
    keys_to_clear = [
        "mode","unit_order","selected_indices_by_unit","quantities_by_unit",
        "units_list","persons_list","current_unit","selected_person_indices",
        "executor","selected_date","quantities","current_qty_index","next_action",
        "dodatok","rozrakh_selected_step1","rozrakh_selected_step2","rozrakh_selected_step3",
        "rozrakh_date_range","rozrakh_nomber","rozrakh_selected_dates","rozrakh_date_text",
        "1date", "2date", "izchogo", "sho", "kontрол", "pidrozdily", "adresa", "kgp_mode",
        "izchogo_list", "sho_list", "pidrozdily_units", "pidpys", "is_vchera"
    ]
    keys_to_clear.extend(["tel_mode", "addresses_required", "addresses_list"])
    # Vidpustky keys
    keys_to_clear.extend([
        "vidp_data_nakazu", "vidpustky_items", "vidpustky_sick_items",
        "vidpustky_return_items", "vidp_current_unit", "vidp_current_person",
        "vidp_leave_type", "vidp_year", "vidp_days", "vidp_date_from",
        "vidp_date_to", "vidp_destination", "vidp_ln_number", "vidp_ln_date",
        "vidp_rapport_pib", "vidp_date_return", "vidp_extra_notes",
        "vidp_editing_item", "vidp_pidpys", "vidp_mode"
    ])
    for k in keys_to_clear:
        user_state[uid].pop(k, None)
    # Также очищаем временный файл адреса
    cleanup_address_request(uid)

def safe_remove(filename):
    for attempt in range(10):
        try:
            os.remove(filename)
            return True
        except PermissionError:
            time.sleep(1)
    print(f"Не вдалося видалити файл {filename} після 10 спроб.")
    return False

# --- Візуалізація прогресу --- #
async def _spinner_message(message_obj: types.Message, base_text: str, stop_event: asyncio.Event):
    frames = ["☕", "😴", "🤔", "🧠", "💭", "💡", "🤯", "🥱", "📚", "🤓", "🚀", "✅"]
    i = 0
    try:
        while not stop_event.is_set():
            frame = frames[i % len(frames)]
            try:
                await message_obj.edit_text(f"{base_text} {frame}")
            except TelegramBadRequest:
                pass
            i += 1
            await asyncio.sleep(0.3)
    except asyncio.CancelledError:
        return

# --- Робота з Google Sheets --- #
def _expected_headers(dodatok_type: str):
    if dodatok_type == "1":
        return [
            "Підрозділ",
            "Спеціальне звання",
            "Посада",
            "ПІБ",
            "kgp",
            "izchogo",
            "sho",
            "kontrol"
        ]
    elif dodatok_type == "3":
        return [
            "izchogo",
            "sho",
            "kontrol",
        ]
    else:
        return [
            "Підрозділ, що комплектує",
            "Марка техніки",
            "Номерний знак",
        ]

def _normalize_lookup_key(value: str) -> str:
    if not isinstance(value, str):
        return ""
    cleaned = value.replace('\u00A0', ' ').replace('\u2060', ' ')
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip().lower()

_CUSTOM_UNIT_ORDER = [
    ("ДПРЧ", ["1", "2", "3", "4", "5", "6", "8", "9", "10", "20", "25", "33"]),
    ("ДПРП", ["2", "3", "4", "6", "8", "10"]),
]

def sort_units(units: list[str]) -> list[str]:
    type_priority = {unit_type: idx for idx, (unit_type, _) in enumerate(_CUSTOM_UNIT_ORDER)}
    number_priority = {
        (unit_type, number): pos
        for unit_type, numbers in _CUSTOM_UNIT_ORDER
        for pos, number in enumerate(numbers)
    }

    def _key(unit: str):
        unit = unit or ""
        stripped = unit.strip()
        match = re.match(r"^(\d+)\s*(ДПРЧ|ДПРП)\b", stripped, flags=re.IGNORECASE)
        if match:
            number = match.group(1)
            unit_type = match.group(2).upper()
            if (unit_type, number) in number_priority:
                type_idx = type_priority.get(unit_type, len(_CUSTOM_UNIT_ORDER))
                num_idx = number_priority[(unit_type, number)]
                return (1, type_idx, num_idx, stripped)
        return (0, stripped)

    return sorted(units, key=_key)

def _fetch_sheet_data(dodatok_type):
    # Set a global default timeout for blocking operations (like requests)
    # This prevents indefinite hanging on network calls
    socket.setdefaulttimeout(30)
    
    url = GOOGLE_SHEET_URL1 if dodatok_type == "1" else GOOGLE_SHEET_URL2
    try:
        client = get_gspread_client()
        sheet = client.open_by_url(url).sheet1
    except Exception as e:
        print(f"Error connecting to Google Sheets: {e}")
        # Retry once with fresh client
        global _GSPREAD_CLIENT
        _GSPREAD_CLIENT = None
        client = get_gspread_client()
        sheet = client.open_by_url(url).sheet1

    headers = _expected_headers(dodatok_type)
    # For sheet type "1" we must manually map fixed columns and skip header row
    if dodatok_type == "1":
        values = sheet.get_all_values()
        # Sheet has headers on the first row; actual data starts from row 2
        data_rows = values[1:] if values else []
        header_index = {h: None for h in headers}
        if values:
            first_row = values[0]
            for h in headers:
                try:
                    header_index[h] = first_row.index(h)
                except ValueError:
                    header_index[h] = None
        # Force fixed columns for specific fields only, regardless of header names
        # H (7) -> kgp, L (11) -> izchogo, M (12) -> sho, N (13) -> kontrol
        header_index["kgp"] = 7
        header_index["izchogo"] = 11
        header_index["sho"] = 12
        header_index["kontrol"] = 13
        # Ensure ПІБ comes from column B (index 1)
        header_index["ПІБ"] = 1
        records = []
        for row in data_rows:
            item = {}
            for h in headers:
                idx = header_index.get(h)
                item[h] = row[idx].strip() if (idx is not None and idx < len(row)) else ""
            records.append(item)
    else:
        try:
            records = sheet.get_all_records(expected_headers=headers, default_blank="")
        except Exception:
            values = sheet.get_all_values()
            data_rows = values[1:] if len(values) > 1 else []
            header_index = {h: None for h in headers}
            if values:
                first_row = values[0]
                for h in headers:
                    try:
                        header_index[h] = first_row.index(h)
                    except ValueError:
                        header_index[h] = None
            records = []
            for row in data_rows:
                item = {}
                for h in headers:
                    idx = header_index.get(h)
                    item[h] = row[idx].strip() if (idx is not None and idx < len(row)) else ""
                records.append(item)
    for row in records:
        for key in list(row.keys()):
            val = row[key]
            if isinstance(val, str):
                # Strip and normalize common whitespace anomalies from Google Sheets
                row[key] = normalize_text(val)
    # If dodatok type is "1", ensure empty 'Спеціальне звання' defaults to 'Працівник'
    if dodatok_type == "1":
        for row in records:
            if not row.get("Спеціальне звання", "").strip():
                row["Спеціальне звання"] = "Працівник"
    return records

async def fetch_sheet_with_progress(callback: types.CallbackQuery, dodatok_type: str, force_refresh: bool = False):
    now_ts = time.time()
    entry = SHEET_CACHE.get(dodatok_type)
    if (not force_refresh) and entry and (now_ts - entry[1] < CACHE_TTL_SECONDS):
        return entry[0]
    
    # Негайно відповідаємо
    await safe_answer_callback(callback, "Завантаження даних...")
    
    loading_msg = await callback.message.answer("Завантаження даних…")
    try:
        uid = callback.from_user.id
        user_state.setdefault(uid, {}).setdefault('sent_message_ids', []).append((callback.message.chat.id, loading_msg.message_id))
    except Exception:
        pass
    stop_event = asyncio.Event()
    spinner_task = asyncio.create_task(_spinner_message(loading_msg, "Завантаження даних…", stop_event))
    try:
        data = await asyncio.to_thread(_fetch_sheet_data, dodatok_type)
        SHEET_CACHE[dodatok_type] = (data, time.time())
    finally:
        stop_event.set()
        try:
            await spinner_task
        except Exception:
            pass
        try:
            await loading_msg.edit_text("Дані завантажено ✅")
            await asyncio.sleep(0.4)
            await loading_msg.delete()
        except Exception:
            pass
    return data

# --- Aggressive prefetch/warm cache for instant menus --- #
async def _warm_cache_once():
    try:
        global CACHE_READY, CACHE_LAST_ERROR, CACHE_PROGRESS
        _log_debug("[CACHE] Starting cache warm-up...")
        start_time = time.time()
        CACHE_PROGRESS = 5
        
        # Fetch both sheets in thread to avoid blocking loop
        data1 = await asyncio.to_thread(_fetch_sheet_data, "1")
        data2 = await asyncio.to_thread(_fetch_sheet_data, "2")
        now_ts = time.time()
        SHEET_CACHE["1"] = (data1, now_ts)
        SHEET_CACHE["2"] = (data2, now_ts)
        _log_debug(f"[CACHE] Loaded sheet data: {len(data1)} rows from sheet1, {len(data2)} rows from sheet2")

        CACHE_PROGRESS = 25

        # Also cache raw values of sheets in a worker thread to avoid blocking the event loop
        try:
            def _load_raw_values():
                try:
                    creds_local = Credentials.from_service_account_file(CREDENTIALS_FILE, scopes=SCOPES)
                    client_local = gspread.authorize(creds_local)
                    sheet1_local = client_local.open_by_url(GOOGLE_SHEET_URL1).sheet1
                    values1_local = sheet1_local.get_all_values() or []
                    sheet2_local = client_local.open_by_url(GOOGLE_SHEET_URL2).sheet1
                    values2_local = sheet2_local.get_all_values() or []
                    return values1_local, values2_local
                except Exception:
                    return [], []

            values1, values2 = await asyncio.to_thread(_load_raw_values)
        except Exception:
            values1 = []
            values2 = []
        SHEET_CACHE["1_values"] = (values1, now_ts)
        SHEET_CACHE["2_values"] = (values2, now_ts)

        CACHE_PROGRESS = 45

        # Precompute common lists from sheet 1
        def _uniq(values):
            seen = {}
            for v in values:
                if not v:
                    continue
                k = normalize_text(v)
                if k and k not in seen:
                    seen[k] = v
            return list(seen.values())

        iz_list = _uniq([row.get("izchogo", "") for row in data1])
        sho_list = _uniq([row.get("sho", "") for row in data1])
        kontrol_list = _uniq([row.get("kontrol", "") for row in data1])
        units_list = sort_units(list({row.get("Підрозділ", "") for row in data1 if row.get("Підрозділ")}))

        SHEET_CACHE["pre_izchogo"] = (iz_list, now_ts)
        SHEET_CACHE["pre_sho"] = (sho_list, now_ts)
        SHEET_CACHE["pre_kontrol"] = (kontrol_list, now_ts)
        SHEET_CACHE["pre_units1"] = (units_list, now_ts)
        
        # Also cache units for sheet 2 (Додаток 2)
        units_list2 = sort_units(list({row.get("Підрозділ, що комплектує", "") for row in data2 if row.get("Підрозділ, що комплектує")}))
        SHEET_CACHE["pre_units2"] = (units_list2, now_ts)

        CACHE_PROGRESS = 65

        # Cache phone mapping for tel mode
        phone_mapping = await asyncio.to_thread(fetch_phone_mapping)
        SHEET_CACHE["phones"] = (phone_mapping, now_ts)

        CACHE_PROGRESS = 80

        # Addresses from column K via helper (only if not already cached)
        if not SHEET_CACHE.get("addresses"):
            addrs = await asyncio.to_thread(get_addresses_from_sheet)
        else:
            addrs = SHEET_CACHE["addresses"][0]  # Use cached version
        SHEET_CACHE["addresses"] = (addrs, now_ts)

        # Обновляем флаги статуса кеша/мережі для користувача
        CACHE_READY = True
        CACHE_LAST_ERROR = None
        CACHE_PROGRESS = 100

        end_time = time.time()
        cache_stats = {
            "sheets": f"sheet1:{len(data1)}, sheet2:{len(data2)}",
            "units": f"sheet1:{len(units_list)}, sheet2:{len(units_list2)}",
            "lists": f"iz:{len(iz_list)}, sho:{len(sho_list)}, kontrol:{len(kontrol_list)}",
            "phones": len(phone_mapping),
            "addresses": len(addrs),
            "time": f"{end_time - start_time:.2f}s"
        }
        _log_debug(f"[CACHE] Warm-up completed in {cache_stats['time']}. Stats: {cache_stats}")
    except Exception as e:
        CACHE_READY = False
        CACHE_LAST_ERROR = str(e)
        CACHE_PROGRESS = 0
        _log_debug(f"[CACHE] Error during warm-up: {e}")

def get_cache_stats():
    """Get current cache statistics for monitoring"""
    now_ts = time.time()
    stats = {}
    
    cache_keys = ["1", "2", "1_values", "2_values", "pre_izchogo", "pre_sho", "pre_kontrol", "pre_units1", "pre_units2", "phones", "addresses"]
    
    for key in cache_keys:
        entry = SHEET_CACHE.get(key)
        if entry:
            data, ts = entry
            age = now_ts - ts
            stats[key] = {
                "size": len(data) if hasattr(data, '__len__') else "N/A",
                "age_seconds": f"{age:.0f}",
                "fresh": age < CACHE_TTL_SECONDS
            }
        else:
            stats[key] = {"size": 0, "age_seconds": "N/A", "fresh": False}
    
    return stats

async def warm_cache_loop():
    """Один ручний прогрів кешу (використовується для кнопки 'Оновити кеш')."""
    try:
        await _warm_cache_once()
    except Exception as e:
        _log_debug(f"[CACHE] Error in manual warm-up: {e}")

def get_sheet_data(dodatok_type, force_refresh=False):
    now_ts = time.time()
    entry = SHEET_CACHE.get(dodatok_type)
    if (not force_refresh) and entry:
        data, ts = entry
        if now_ts - ts < CACHE_TTL_SECONDS:
            return data
    data = _fetch_sheet_data(dodatok_type)
    SHEET_CACHE[dodatok_type] = (data, now_ts)
    return data

async def get_sheet_data_async(dodatok_type, force_refresh: bool = False):
    now_ts = time.time()
    entry = SHEET_CACHE.get(dodatok_type)
    if (not force_refresh) and entry:
        data, ts = entry
        if now_ts - ts < CACHE_TTL_SECONDS:
            return data
    data = await asyncio.to_thread(_fetch_sheet_data, dodatok_type)
    SHEET_CACHE[dodatok_type] = (data, time.time())
    return data

# --- Наказ: вибір izchogo/sho/kontrol (old flow disabled) --- #
@dp.callback_query(F.data == "nakaz_old_disabled")
async def start_nakaz(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    clear_workflow_state(uid)
    user_state[uid]["mode"] = "nakaz"
    user_state[uid]["dodatok"] = "1"
    await start(callback.message)

async def show_izchogo_selection(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    await start(callback.message)

async def show_sho_selection(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    await start(callback.message)

@dp.callback_query(F.data.startswith("iz_page_"))
async def iz_page_nav(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await start(callback.message)
    await safe_answer_callback(callback)

@dp.callback_query(F.data.startswith("sho_page_"))
async def sho_page_nav(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await start(callback.message)
    await safe_answer_callback(callback)

@dp.callback_query(F.data.regexp(r"^sho_\d+$"))
async def select_sho(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await start(callback.message)
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "sho_ai_prompt")
async def sho_ai_prompt(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state[uid]["state"] = "sho_ai_seed"
    await callback.message.answer("Введіть деталі (одним повідомленням), які допоможуть ШІ сформувати варіанти 'що саме'. Після цього натисніть кнопку 'Варіант ШІ'.")
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "sho_ai_generate")
async def sho_ai_generate(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    seed = user_state.get(uid, {}).get("sho_ai_seed", "")
    iz = user_state.get(uid, {}).get("izchogo", "")
    suggestions = await asyncio.to_thread(_call_gemini_generate_sho, seed, iz, uid)
    if not suggestions:
        await safe_answer_callback(callback, "Не вдалося отримати варіанти ШІ", show_alert=True)
        return
    user_state[uid]["sho_ai_suggestions"] = suggestions
    # Render suggestions with numeric buttons
    lines = [f"{i+1}. {s}" for i, s in enumerate(suggestions)]
    text = "Варіанти від ШІ:\n\n" + "\n\u2009\n".join(lines)
    kb = InlineKeyboardBuilder()
    for i, _ in enumerate(suggestions):
        kb.button(text=str(i+1), callback_data=f"sho_ai_pick_{i}")
    kb.button(text="🔁 Згенерувати ще раз", callback_data="sho_ai_generate")
    kb.button(text="🔙 Назад", callback_data="back_to_start")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(2)
    try:
        await callback.message.edit_text(text, reply_markup=kb.as_markup())
    except Exception:
        await callback.message.answer(text, reply_markup=kb.as_markup())
    await safe_answer_callback(callback)

@dp.callback_query(F.data.startswith("sho_ai_pick_"))
async def sho_ai_pick(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    try:
        idx = int(callback.data.replace("sho_ai_pick_", ""))
    except Exception:
        await safe_answer_callback(callback)
        return
    lst = user_state.get(uid, {}).get("sho_ai_suggestions", [])
    if not (0 <= idx < len(lst)):
        await safe_answer_callback(callback)
        return
    user_state[uid]["sho"] = lst[idx]
    await show_calendar_for_nakaz(callback, uid, date_key="2date")
    await safe_answer_callback(callback)

async def show_kontrol_selection(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    # Guard: in new flow, kontrol is chosen at the end; skip this old menu
    if user_state.get(uid, {}).get("nakaz_flow") == "new":
        return
    data = await get_sheet_data_async("1", True)
    # Build a global distinct list of 'kontrol' independent from 'izchogo' and 'sho'
    kontrol_seen = {}
    for r in data:
        ko = r.get("kontrol", "").strip()
        key = normalize_text(ko)
        if key and key not in kontrol_seen:
            kontrol_seen[key] = ko
    kontrol_list = list(kontrol_seen.values())
    user_state[uid]["kontrol_list"] = kontrol_list
    if not kontrol_list:
        user_state[uid]["kontrol"] = ""
        await show_calendar_for_nakaz(callback, uid, date_key="2date")
        return
    kb = InlineKeyboardBuilder()
    for i, ko in enumerate(kontrol_list):
        kb.button(text=ko, callback_data=f"kontrol_{i}")
    kb.button(text="🔙 Назад", callback_data="back_to_izchogo")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    text = "Оберіть хто здыйнює контроль виконання наказу"
    try:
        await callback.message.edit_text(text, reply_markup=kb.as_markup())
    except Exception:
        await callback.message.answer(text, reply_markup=kb.as_markup())

@dp.callback_query(F.data.startswith("kontrol_"))
async def select_kontrol(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    
    try:
        index = int(callback.data.replace("kontrol_", ""))
        k_list = user_state[uid].get("kontrol_list", [])
        if 0 <= index < len(k_list):
            user_state[uid]["kontrol"] = k_list[index]
        else:
            user_state[uid]["kontrol"] = ""
        
        # Швидко відповідаємо
        await safe_answer_callback(callback, "Контроль обрано")
        
        # After selecting kontrol, show pidpys selection
        await show_pidpys_selection(callback, uid)
    except Exception as e:
        print(f"Помилка в select_kontrol: {e}")
        await safe_answer_callback(callback, "Сталася помилка", show_alert=True)

# --- Робота з Word документами --- #
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def disable_auto_hyphenation_paragraph(paragraph):
    try:
        pPr = paragraph._p.get_or_add_pPr()
        el = OxmlElement('w:suppressAutoHyphenation')
        el.set(qn('w:val'), '1')
        pPr.append(el)
    except Exception:
        pass

def disable_auto_hyphenation_document(doc: Document):
    try:
        for p in doc.paragraphs:
            disable_auto_hyphenation_paragraph(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        disable_auto_hyphenation_paragraph(p)
    except Exception:
        pass

def apply_protect_no_wrap_document(doc: Document):
    def _has_drawing(run) -> bool:
        try:
            r = run._r
            return bool(r.xpath('.//w:drawing') or r.xpath('.//w:pict'))
        except Exception:
            return False
    try:
        for p in doc.paragraphs:
            for r in p.runs:
                try:
                    # Skip runs with images/drawings and empty-text runs
                    if _has_drawing(r) or not r.text:
                        continue
                    r.text = protect_no_wrap(r.text)
                except Exception:
                    pass
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            try:
                                if _has_drawing(r) or not r.text:
                                    continue
                                r.text = protect_no_wrap(r.text)
                            except Exception:
                                pass
    except Exception:
        pass

def decline_rank(rank):
    rank = rank.lower().strip()
    ranks_dict = {
        "рядовий": "рядового",
        "молодший сержант": "молодшого сержанта",
        "сержант": "сержанта",
        "старший сержант": "старшого сержанта",
        "старшина": "старшину",
        "прапорщик": "прапорщика",
        "старший прапорщик": "старшого прапорщика",
        # Додайте більше за потребою
    }
    return ranks_dict.get(rank, rank) # TODO: покращити відмінювання

def decline_position(position):
    position = position.lower().strip()
    pos_dict = {
        "начальник": "начальника",
        # Додайте більше за потребою
    }
    return pos_dict.get(position, position) # TODO: покращити відмінювання

def decline_pib(pib):
    # Повертаємо повне ПІБ без скорочень
    return pib

# --- Відмінювання для наказу відпустки --- #

def _is_female(pib: str) -> bool:
    """Визначає чи ПІБ належить жінці (за іменем або по батькові)."""
    parts = pib.strip().split()
    if len(parts) < 2:
        return False
    # Жіночі закінчення імен
    female_name_endings = ('а', 'я', 'ія', 'ія')
    # Жіночі закінчення по батькові
    female_patronymic_endings = ('івна', 'ївна', 'овна')
    # Перевіряємо ім'я (друга частина)
    if len(parts) >= 2:
        given = parts[1].lower()
        if given.endswith(female_name_endings) and given not in ('микола', 'микита', 'ілля', 'лука', 'кузьма', 'сава', 'йона'):
            return True
    # Перевіряємо по батькові
    if len(parts) >= 3:
        patr = parts[2].lower()
        if patr.endswith(female_patronymic_endings):
            return True
    return False

# Full given name genitive mapping
_GIVEN_NAME_GEN = {
    # -й ending → -я
    "анатолій": "Анатолія", "андрій": "Андрія", "валерій": "Валерія",
    "василь": "Василя", "віктор": "Віктора", "віталій": "Віталія",
    "володимир": "Володимира", "геннадій": "Геннадія", "григорій": "Григорія",
    "денис": "Дениса", "дмитро": "Дмитра", "євген": "Євгена",
    "євгеній": "Євгенія", "іван": "Івана", "ігор": "Ігоря",
    "кирило": "Кирила", "костянтин": "Костянтина", "максим": "Максима",
    "микола": "Миколи", "михайло": "Михайла", "олег": "Олега",
    "олександр": "Олександра", "олексій": "Олексія", "павло": "Павла",
    "петро": "Петра", "роман": "Романа", "руслан": "Руслана",
    "сергій": "Сергія", "степан": "Степана", "тарас": "Тараса",
    "юрий": "Юрія", "юрій": "Юрія", "ярослав": "Ярослава",
    "артем": "Артема", "богдан": "Богдана", "вадим": "Вадима",
    "валентин": "Валентина", "в'ячеслав": "В'ячеслава", "влас": "Власа",
    "данило": "Данила", "захар": "Захара", "ллія": "Іллі",
    "лев": "Лева", "леонід": "Леоніда", "марк": "Марка",
    "микита": "Микити", "назар": "Назара", "олесь": "Олеся",
    "ростислав": "Ростислава", "святослав": "Святослава", "семен": "Семена",
    "станіслав": "Станіслава", "тимофій": "Тимофія", "федір": "Федора",
    "юхим": "Юхима", "яків": "Якова", "ян": "Яна",
    "болеслав": "Болеслава", "владислав": "Владислава",
    "всеволод": "Всеволода", "арсен": "Арсена", "арсеній": "Арсенія",
    "єгор": "Єгора", "емануїл": "Емануїла",
    # Female names (for full name genitive)
    "алла": "Алли", "анна": "Анни", "валентина": "Валентини",
    "вікторія": "Вікторії", "галина": "Галини", "дарія": "Дарії",
    "еліна": "Еліни", "елізавета": "Елізавети", "єва": "Єви",
    "інна": "Інни", "ірина": "Ірини", "катерина": "Катерини",
    "лариса": "Лариси", "лідія": "Лідії", "любов": "Любові",
    "людмила": "Людмили", "маргарита": "Маргарити", "марина": "Марини",
    "марія": "Марії", "надія": "Надії", "наталія": "Наталії",
    "оксана": "Оксани", "олена": "Олени", "ольга": "Ольги",
    "раїса": "Раїси", "світлана": "Світлани", "софія": "Софії",
    "тамара": "Тамари", "тетяна": "Тетяни", "юлія": "Юлії",
    "яна": "Яни", "евгенія": "Євгенії", "соломія": "Соломії",
}

# Full given name dative mapping
_GIVEN_NAME_DAT = {
    "анатолій": "Анатолію", "андрій": "Андрію", "валерій": "Валерію",
    "василь": "Василю", "віктор": "Віктору", "віталій": "Віталію",
    "володимир": "Володимиру", "геннадій": "Геннадію", "григорій": "Григорію",
    "денис": "Денису", "дмитро": "Дмитру", "євген": "Євгену",
    "євгеній": "Євгенію", "іван": "Івану", "ігор": "Ігорю",
    "кирило": "Кирилу", "костянтин": "Костянтину", "максим": "Максиму",
    "микола": "Миколі", "михайло": "Михайлу", "олег": "Олегу",
    "олександр": "Олександру", "олексій": "Олексію", "павло": "Павлу",
    "петро": "Петру", "роман": "Роману", "руслан": "Руслану",
    "сергій": "Сергію", "степан": "Степану", "тарас": "Тарасу",
    "юрий": "Юрію", "юрій": "Юрію", "ярослав": "Ярославу",
    "артем": "Артему", "богдан": "Богдану", "вадим": "Вадиму",
    "валентин": "Валентину", "в'ячеслав": "В'ячеславу",
    "данило": "Данилу", "ілля": "Іллі", "леонід": "Леоніду",
    "назар": "Назару", "святослав": "Святославу",
    "тимофій": "Тимофію", "арсен": "Арсену", "арсеній": "Арсенію",
    "микита": "Микиті", "данило": "Данилу", "данiiл": "Даніілу",
    "дмитрій": "Дмитрію", "владислав": "Владиславу", "в'ячеслав": "В'ячеславу",
    "вячеслав": "Вячеславу", "леонід": "Леоніду", "семен": "Семену",
    "ростислав": "Ростиславу", "станіслав": "Станіславу",
    "артур": "Артуру", "ернівович": "Ернівовичу",
    # Female names
    "алла": "Аллі", "анна": "Анні", "валентина": "Валентині",
    "галина": "Галині", "ірина": "Ірині", "катерина": "Катерині",
    "людмила": "Людмилі", "марія": "Марії", "надія": "Надії",
    "наталія": "Наталії", "оксана": "Оксані", "олена": "Олені",
    "ольга": "Ользі", "світлана": "Світлані", "тетяна": "Тетяні",
    "юлія": "Юлії",
}

def decline_rank_dative(rank: str) -> str:
    """Звання в давальному відмінку (кому?).
    Додає 'служби цивільного захисту' якщо звання не 'Працівник'."""
    rank_clean = rank.strip()
    # Якщо звання вже містить "служби цивільного захисту" — відмінюємо тільки звання
    suffix = ""
    if "служби цивільного захисту" in rank_clean.lower():
        # Витягуємо чисте звання
        rank_clean = rank_clean.replace("служби цивільного захисту", "").strip()
        suffix = " служби цивільного захисту"

    rank_lower = rank_clean.lower()
    dative_map = {
        "рядовий": "рядовому",
        "молодший сержант": "молодшому сержанту",
        "сержант": "сержанту",
        "старший сержант": "старшому сержанту",
        "старшина": "старшині",
        "прапорщик": "прапорщику",
        "старший прапорщик": "старшому прапорщику",
        "лейтенант": "лейтенанту",
        "старший лейтенант": "старшому лейтенанту",
        "капітан": "капітану",
        "майор": "майору",
        "підполковник": "підполковнику",
        "полковник": "полковнику",
        "майстер-сержант": "майстер-сержанту",
        "головний майстер-сержант": "головному майстер-сержанту",
        "працівник": "працівнику",
    }
    declined = dative_map.get(rank_lower, rank_clean)

    # Додаємо «служби цивільного захисту» якщо це військове звання (не Працівник)
    if suffix:
        return declined + suffix
    if rank_lower != "працівник":
        return declined + " служби цивільного захисту"
    return declined


def _capitalize_ending(base: str, ending: str) -> str:
    """Додає закінчення з урахуванням регістру бази."""
    letters = [c for c in base if c.isalpha()]
    is_all_caps = len(letters) > 0 and all(c.isupper() for c in letters)
    if is_all_caps:
        return base + ending.upper()
    elif base and base[0].isupper():
        return base + ending
    else:
        return base + ending


def decline_name_dative(pib: str) -> str:
    """ПІБ у давальному відмінку (кому?). Жіночі прізвища на приголосний/-о НЕ відмінюються."""
    parts = pib.strip().split()
    if not parts:
        return pib
    is_female = _is_female(pib)

    def decline_surname_dative(s: str) -> str:
        """Базове відмінювання прізвищ"""
        s_lower = s.lower()
        # Спеціальні випадки
        special = {
            "орел": "Орлу", "ревенок": "Ревенку", "швець": "Швецю",
            "беленок": "Беленку", "коток": "Котку", "козачек": "Козачку",
            "богдан": "Богдану", "гапон": "Гапону", "єлагін": "Єлагіну",
            "кравченко": "Кравченку", "лагута": "Лагуті", "лимар": "Лимару",
            "лось": "Лосю", "павленко": "Павленку", "полкозак": "Полкозаку",
            "хобел": "Хобелу", "беспалий": "Беспалому", "пеньковець": "Пеньковцю",
            "гець": "Гецю", "чумак": "Чумаку", "силенко": "Силенку",
            "омеляненко": "Омеляненку", "паньковський": "Паньковському",
            "коржов": "Коржову", "усик": "Усику", "середа": "Середі",
            "дилін": "Диліну", "царенко": "Царенку", "федун": "Федуну",
            "закотій": "Закотію", "антипенко": "Антипенку", "онищенко": "Онищенку",
            "чуркіна": "Чуркіній", "глянь": "Гляню",
        }
        if s_lower in special:
            result = special[s_lower]
            letters = [c for c in s if c.isalpha()]
            if letters and all(c.isupper() for c in letters):
                return result.upper()
            return result
        # -ів/-їв → -еву/-єву (Ковалів→Ковалеву, Андрусів→Андрусеву)
        if s_lower.endswith("ів") and len(s) > 3:
            return _capitalize_ending(s[:-2], "еву")
        if s_lower.endswith("їв") and len(s) > 3:
            return _capitalize_ending(s[:-2], "єву")
        # Загальні правила (з урахуванням регістру)
        if s_lower.endswith("ко"):
            return _capitalize_ending(s[:-1], "у")  # АНДРОСЕНКО → АНДРОСЕНК+У
        if s_lower.endswith("ук") or s_lower.endswith("юк") or s_lower.endswith("чук"):
            return _capitalize_ending(s, "у")  # НЕЧЕПОРУК→НЕЧЕПОРУКУ (не обрізати!)
        if s_lower.endswith(("ов", "ев", "єв", "ин", "ін", "їн")):
            return _capitalize_ending(s, "у")
        # Female adjective-type: -ова/-ева/-іна/-ська → -овій/-евій/-іній/-ській
        if s_lower.endswith(("ова", "ева", "єва", "іна", "ська", "цька", "зька")):
            return _capitalize_ending(s[:-1], "ій")
        if s_lower.endswith("ній"):
            return _capitalize_ending(s[:-2], "ньому")  # КУТНІЙ→КУТНЬОМУ
        if s_lower.endswith("ий") or s_lower.endswith("ій"):
            return _capitalize_ending(s[:-2], "ому")
        if s_lower.endswith(("ка", "га", "ха")):
            # consonant mutation: к→ц, г→з, х→с before -і
            last = s[-2].lower()
            if last == 'г': repl = 'З' if s[-2].isupper() else 'з'
            elif last == 'к': repl = 'Ц' if s[-2].isupper() else 'ц'
            elif last == 'х': repl = 'С' if s[-2].isupper() else 'с'
            else: repl = s[-2]
            return _capitalize_ending(s[:-2] + repl, "і")
        if s_lower.endswith("а"):
            return s[:-1] + "і"  # -а → -і
        if s_lower.endswith("я"):
            return _capitalize_ending(s[:-1], "і")  # -я → -і
        if s_lower.endswith("й"):
            return _capitalize_ending(s[:-1], "ю")  # -й → -ю (ПОЛЬГУЙ→ПОЛЬГУЮ)
        if s_lower.endswith("ець"):
            return _capitalize_ending(s[:-3], "цю")
        if s_lower.endswith("ь"):
            return _capitalize_ending(s[:-1], "ю")
        # Default: add -у for consonants
        if s_lower[-1] not in 'аеєиіїоуюя':
            return _capitalize_ending(s, "у")
        return s

    def decline_patronymic_dative(s: str) -> str:
        """По батькові у давальному відмінку"""
        s_lower = s.lower()
        if s_lower.endswith("ович"):
            return s[:-4] + "овичу"
        if s_lower.endswith("івна"):
            return s + ""  # без змін
        if s_lower.endswith("ич"):
            return s + "у"
        return s

    result_parts = []
    for i, part in enumerate(parts):
        p_lower = part.lower()
        if i == 0:  # Прізвище
            # Жіночі прізвища на приголосний/-о НЕ відмінюються
            # Жіночі прізвища на -а/-я ВІДМІНЮЮТЬСЯ
            if is_female and not part.lower().endswith(('а', 'я')):
                result_parts.append(part)
            else:
                result_parts.append(decline_surname_dative(part))
        elif i == len(parts) - 1:  # По батькові
            result_parts.append(decline_patronymic_dative(part))
        else:  # Ім'я
            result_parts.append(_GIVEN_NAME_DAT.get(p_lower, part))

    return " ".join(result_parts)


def decline_rank_genitive(rank: str) -> str:
    """Звання в родовому відмінку (кого?). Для §2, §3."""
    rank_clean = rank.strip()
    suffix = ""
    if "служби цивільного захисту" in rank_clean.lower():
        rank_clean = rank_clean.replace("служби цивільного захисту", "").strip()
        suffix = " служби цивільного захисту"

    rank_lower = rank_clean.lower()
    genitive_map = {
        "рядовий": "рядового",
        "молодший сержант": "молодшого сержанта",
        "сержант": "сержанта",
        "старший сержант": "старшого сержанта",
        "старшина": "старшини",
        "прапорщик": "прапорщика",
        "старший прапорщик": "старшого прапорщика",
        "лейтенант": "лейтенанта",
        "старший лейтенант": "старшого лейтенанта",
        "капітан": "капітана",
        "майор": "майора",
        "підполковник": "підполковника",
        "полковник": "полковника",
        "майстер-сержант": "майстер-сержанта",
        "головний майстер-сержант": "головного майстер-сержанта",
        "працівник": "працівника",
    }
    declined = genitive_map.get(rank_lower, rank_clean)

    if suffix:
        return declined + suffix
    if rank_lower != "працівник":
        return declined + " служби цивільного захисту"
    return declined


def decline_name_genitive(pib: str) -> str:
    """ПІБ у родовому відмінку (кого?). Жіночі прізвища на приголосний/-о НЕ відмінюються."""
    parts = pib.strip().split()
    if not parts:
        return pib
    is_female = _is_female(pib)

    def decline_surname_genitive(s: str) -> str:
        s_lower = s.lower()
        special = {
            "орел": "Орла", "ревенок": "Ревенка", "швець": "Швеця",
            "богдан": "Богдана", "гапон": "Гапона", "єлагін": "Єлагіна",
            "кравченко": "Кравченка", "лагута": "Лагути", "лимар": "Лимара",
            "лось": "Лося", "павленко": "Павленка", "полкозак": "Полкозака",
            "хобел": "Хобела", "беспалий": "Беспалого", "пеньковець": "Пеньковця",
            "гець": "Геця", "чумак": "Чумака", "силенко": "Силенка",
            "омеляненко": "Омеляненка", "паньковський": "Паньковського",
            "коржов": "Коржова", "усик": "Усика", "середа": "Середи",
            "чуркіна": "Чуркіної",
        }
        if s_lower in special:
            result = special[s_lower]
            # Preserve all-caps input
            letters = [c for c in s if c.isalpha()]
            if letters and all(c.isupper() for c in letters):
                return result.upper()
            return result
        # -ів/-їв → -ева/-єва (Ковалів→Ковалева)
        if s_lower.endswith("ів") and len(s) > 3:
            return _capitalize_ending(s[:-2], "ева")
        if s_lower.endswith("їв") and len(s) > 3:
            return _capitalize_ending(s[:-2], "єва")
        if s_lower.endswith("ко"):
            return _capitalize_ending(s[:-1], "а")  # АНДРОСЕНКО → АНДРОСЕНК+А
        if s_lower.endswith(("ук", "юк", "чук")):
            return _capitalize_ending(s, "а")  # НЕЧЕПОРУК→НЕЧЕПОРУКА (keep stem)
        if s_lower.endswith(("ов", "ев", "єв", "ин", "ін", "їн")):
            return _capitalize_ending(s, "а")
        # Female adjective-type: -ова/-ева/-іна/-ська → -ової/-евої/-іної/-ської
        if s_lower.endswith(("ова", "ева", "єва", "іна", "ська", "цька", "зька")):
            return _capitalize_ending(s[:-1], "ої")
        if s_lower.endswith("ий") or s_lower.endswith("ій"):
            return _capitalize_ending(s[:-2], "ого")
        if s_lower.endswith("а"):
            return s[:-1] + "и"  # -а → -и (no consonant mutation for genitive)
        if s_lower.endswith("я"):
            return _capitalize_ending(s[:-1], "і")  # -я → -і
        if s_lower.endswith("й"):
            return _capitalize_ending(s[:-1], "я")  # -й → -я (ЧУГАЙ→ЧУГАЯ)
        if s_lower.endswith("ець"):
            return _capitalize_ending(s[:-3], "ця")
        if s_lower.endswith("ь"):
            return _capitalize_ending(s[:-1], "я")
        if s_lower[-1] not in 'аеєиіїоуюя':
            return _capitalize_ending(s, "а")
        return s

    def decline_patronymic_genitive(s: str) -> str:
        s_lower = s.lower()
        if s_lower.endswith("ович"):
            return s[:-4] + "овича"
        if s_lower.endswith("івна"):
            return s
        if s_lower.endswith("ич"):
            return s + "а"
        return s

    result_parts = []
    for i, part in enumerate(parts):
        p_lower = part.lower()
        if i == 0:
            if is_female and not part.lower().endswith(('а', 'я')):
                result_parts.append(part)  # жіноче прізвище на приголосний не відмінюється
            else:
                result_parts.append(decline_surname_genitive(part))
        elif i == len(parts) - 1:
            result_parts.append(decline_patronymic_genitive(part))
        else:
            result_parts.append(_GIVEN_NAME_GEN.get(p_lower, part))

    return " ".join(result_parts)


def format_name_short(pib: str) -> str:
    """Коротка форма ПІБ для рапорту: 'БОГДАН Владислав Ігорович' -> 'Богдана В.І.'"""
    parts = pib.strip().split()
    if not parts:
        return pib
    surname = parts[0]
    # Відмінюємо прізвище в родовий відмінок
    surname_gen = decline_surname_genitive_for_short(surname)
    # Приводимо до формату: перша літера велика, решта малі
    surname_formatted = surname_gen[0].upper() + surname_gen[1:].lower() if len(surname_gen) > 1 else surname_gen.upper()
    initials = ""
    for part in parts[1:]:
        if part:
            initials += part[0].upper() + "."
    return f"{surname_formatted} {initials}"


def decline_surname_genitive_for_short(s: str) -> str:
    """Прізвище в родовому відмінку для короткої форми."""
    s_lower = s.lower()
    special = {
        "богдан": "Богдана", "гапон": "Гапона", "єлагін": "Єлагіна",
        "кравченко": "Кравченка", "лагута": "Лагути", "лимар": "Лимара",
        "лось": "Лося", "павленко": "Павленка", "полкозак": "Полкозака",
        "хобел": "Хобела", "беспалий": "Беспалого", "пеньковець": "Пеньковця",
        "гець": "Геця", "чумак": "Чумака", "силенко": "Силенка",
        "омеляненко": "Омеляненка", "паньковський": "Паньковського",
        "коржов": "Коржова", "усик": "Усика", "середа": "Середи",
        "дилін": "Диліна", "царенко": "Царенка", "федун": "Федуна",
    }
    if s_lower in special:
        return special[s_lower]
    if s_lower.endswith("ко"):
        return s[:-1] + "а"
    if s_lower.endswith("ов") or s_lower.endswith("ев") or s_lower.endswith("єв"):
        return s + "а"
    if s_lower.endswith("ий") or s_lower.endswith("ій"):
        return s[:-2] + "ого"
    if s_lower.endswith("ець"):
        return s[:-2] + "ця"
    if s_lower.endswith("а"):
        return s[:-1] + "и"
    if s_lower[-1] not in 'аеєиіїоуюя':
        return s + "а"
    return s


# --- Форматування пунктів наказу відпустки --- #

def _format_vidpustky_item(item: dict, number: int) -> tuple:
    """Форматує один пункт §1 (відпустка).
    Використовує готові відмінки з Google Sheets (колонки O, U, Q).
    Повертає (текст_пункту, текст_підстави)."""
    # Використовуємо готові давальні відмінки з таблиці (колонки O, U, Q)
    rank_dative = item.get("rank_dative", "") or decline_rank_dative(item.get("rank", ""))
    name_dative = item.get("pib_dative", "") or decline_name_dative(item.get("pib", ""))
    position_full = item.get("position_full", "") or item.get("position", "")
    leave_type = item.get("leave_type", "shchorichna")
    year = item.get("year", "")
    days = item.get("days", "")
    date_from = item.get("date_from", "")
    date_to = item.get("date_to", "")
    destination = item.get("destination", "")

    # Велика перша літера у званні
    if rank_dative and rank_dative[0].islower():
        rank_dative = rank_dative[0].upper() + rank_dative[1:]

    leave_type_text = _get_leave_type_text(leave_type)
    travel_days = int(item.get("travel_days", 0) or 0)
    total_days = int(days) + travel_days
    days_word = _plural_days(total_days)

    # Розшифровуємо абревіатури
    position_expanded = _expand_abbreviations(position_full)

    # Визначаємо чи є звання (колонка E — якщо порожня, це працівник без звання)
    has_rank = item.get("has_rank", bool(rank_dative and rank_dative.strip()))

    # Формат без звання: "N. [ПІБ_дав], [посада], [тип] за період з [date_from] по [date_to]"
    if not has_rank:
        period_text = f"за період з {date_from} по {date_to}"
        text = f"{number}. {name_dative}, {position_expanded}, {leave_type_text} {period_text}"
    else:
        # Формат зі званням: "N. [звання_дав] [ПІБ_дав], [посада], [тип] за [рік] рік терміном на N днів з ... по ..."
        text = f"{number}. {rank_dative} {name_dative}, {position_expanded}, {leave_type_text} за {year} рік терміном на {total_days} календарних {days_word} з {date_from} по {date_to}"

    if destination:
        text += f", з виїздом до {destination}"
        if travel_days > 0 and has_rank:
            travel_word = _plural_days(travel_days)
            text += f" (з них {travel_days} календарних {travel_word} на проїзд)"
    text += "."

    # Підстава: рапорт (зі званням) / заява (без звання)
    name_short = format_name_short(item.get("pib", ""))
    pidstava_type = "рапорт" if has_rank else "заява"
    pidstava = f"Підстава: {pidstava_type} {name_short.rstrip('.')}."

    return text, pidstava


def _format_sick_item(item: dict, number: int) -> tuple:
    """Форматує один пункт §2 (лікарняний). Родовий відмінок.
    Використовує колонки W (звання), V (ПІБ), J (посада)."""
    rank_g = item.get("rank_genitive", "") or decline_rank_genitive(item.get("rank", ""))
    if rank_g and rank_g[0].islower():
        rank_g = rank_g[0].upper() + rank_g[1:]
    name_g = item.get("pib_genitive", "") or decline_name_genitive(item.get("pib", ""))
    position = _expand_abbreviations(item.get("position_genitive", "") or item.get("position", ""))
    date_from = item.get("date_from", "")
    ln_number = item.get("ln_number", "")
    ln_date = item.get("ln_date", "")
    extra_notes = item.get("extra_notes", "")

    text = f"{number}. {rank_g} {name_g}, {position}, з {date_from}."

    # Формуємо підставу — колонка B (називний), формат "ПРІЗВИЩЕ І.П."
    rapport_pib = item.get("rapport_pib", "") or item.get("pib", "")
    parts = rapport_pib.strip().split()
    if len(parts) >= 2:
        # Format: колонка V (родовий) + ініціали "АНДРОСЕНКА І.С"
        surname = parts[0]  # вже в родовому відмінку
        init_list = [p[0].upper() + "." for p in parts[1:-1]]  # всі крім останнього
        if len(parts) > 1:
            init_list.append(parts[-1][0].upper())  # останній ініціал без крапки
        initials = "".join(init_list)
        pidstava = f"Підстава: рапорт {surname} {initials}."
    else:
        pidstava = f"Підстава: рапорт {rapport_pib}."
    if ln_number:
        pidstava += f", запис в Електронному реєстрі листків непрацездатності"
        if ln_date:
            pidstava += f" від {ln_date}"
        pidstava += f" ЛН {ln_number}"
    if extra_notes:
        pidstava += f", {extra_notes}"
    pidstava += "."

    return text, pidstava


def _format_return_item(item: dict, number: int) -> tuple:
    """Форматує один пункт §3 (повернення). Родовий відмінок.
    Використовує колонки W (звання), V (ПІБ), J (посада)."""
    rank_g = item.get("rank_genitive", "") or decline_rank_genitive(item.get("rank", ""))
    if rank_g and rank_g[0].islower():
        rank_g = rank_g[0].upper() + rank_g[1:]
    name_g = item.get("pib_genitive", "") or decline_name_genitive(item.get("pib", ""))
    position = _expand_abbreviations(item.get("position_genitive", "") or item.get("position", ""))
    date_return = item.get("date_return", "")

    text = f"{number}. {rank_g} {name_g}, {position}, з {date_return}."

    rapport_name = format_name_short(item.get("rapport_pib") or item.get("pib", ""))
    pidstava = f"Підстава: рапорт {rapport_name.rstrip('.')}."

    return text, pidstava


def _expand_abbreviations(text: str) -> str:
    """Розшифровує абревіатури в тексті посади (родовий відмінок)."""
    if not text:
        return text
    # Replace abbreviations with full forms (genitive case for subordinate unit descriptions)
    replacements = [
        ("ДПРЧ", "державної пожежно-рятувальної частини"),
        ("ДПРП", "державного пожежно-рятувального поста"),
        ("ДПРЗ", "державного пожежно-рятувального загону"),
        ("ГУ ДСНС", "Головного управління ДСНС"),
        (" ГУ ", " Головного управління "),
    ]
    result = text
    for abbr, full in replacements:
        result = result.replace(abbr, full)
    return result

def _plural_days(n: int) -> str:
    """Правильна форма слова 'день' для числа n: 1 день, 2 дні, 5 днів."""
    if 11 <= n % 100 <= 14:
        return "днів"
    last = n % 10
    if last == 1:
        return "день"
    if 2 <= last <= 4:
        return "дні"
    return "днів"

def _get_leave_type_text(leave_type: str) -> str:
    """Повертає повний текст типу відпустки."""
    types = {
        "shchorichna": "частину щорічної основної відпустки",
        "shchorichna_full": "щорічну основну відпустку",
        "dodatkova": "додаткову відпустку",
        "simeina": "відпустку за сімейними обставинами",
        "kalendarna": "календарну відпустку",
        "navchalna": "навчальну відпустку",
    }
    return types.get(leave_type, leave_type)


def format_list(items):
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} та {items[1]}"
    return ", ".join(items[:-1]) + f" та {items[-1]}"

def format_date(date_obj):
    return f"{date_obj.day:02} {uk_month_gen[date_obj.month - 1]} {date_obj.year} року"

def normalize_text(s: str) -> str:
    """Normalize text for robust equality: replace NBSP, collapse spaces, trim, lower preserved.
    Keep case as-is for display; only use this for comparisons and dedup.
    """
    if not isinstance(s, str):
        return ""
    # Replace non-breaking spaces and zero-width joiners, then collapse whitespace
    s = s.replace('\u00A0', ' ').replace('\u2060', '')
    s = re.sub(r"\s+", " ", s)
    return s.strip()

def sanitize_sho_text(s: str) -> str:
    """Sanitize 'sho' text before inserting into the document.
    - Remove surrounding quotes (", ', «», "“, „“, etc.), possibly nested.
    - Trim whitespace.
    - Ensure ending period if there is no sentence-ending punctuation.
    """
    if not isinstance(s, str):
        return ""
    text = s.strip()
    # Define quote pairs (opening -> closing). Include symmetric quotes too.
    pairs = {
        '"': '"',
        "'": "'",
        '«': '»',
        '“': '”',
        '„': '“',
        '‚': '‘',
        '‘': '’',
        '”': '”',  # handle stray same-quote cases gracefully
        '»': '»',
    }
    # Strip nested surrounding quotes
    stripped = True
    while stripped and len(text) >= 2:
        stripped = False
        first, last = text[0], text[-1]
        close = pairs.get(first)
        if close and last == close:
            text = text[1:-1].strip()
            stripped = True
        elif first == last and first in ('"', "'"):
            text = text[1:-1].strip()
            stripped = True
    # Ensure trailing sentence punctuation
    if text and not re.search(r"[\.!?…]$", text):
        text = text + "."
    return text

def protect_no_wrap(s: str) -> str:
    """Insert non-breaking characters to strictly prevent line breaks in key phrases.
    - Replace spaces in patterns like '1 ДПРЧ', '2 ДПРЗ', '3 ДПРП' with NBSP and add ZWJ.
    - Ensure 'м. <Місто>' and 'селище <Назва>' (optionally in parentheses) are non-breaking.
    - Convert hyphens inside those names to non-breaking hyphen (U+2011).
    """
    if not isinstance(s, str) or not s:
        return s

    def _nobreak_hyphens(text: str) -> str:
        # Replace standard hyphen between letters with non-breaking hyphen
        return re.sub(r"(?<=\w)-(?!\s)", "\u2011", text)

    out = s

    # 1) Number + DПР[ЧЗП]
    out = re.sub(r"\b(\d+)\s+(ДПР[ЧЗП])\b", lambda m: f"{m.group(1)}\u00A0\u2060{m.group(2)}", out)

    # Helper to protect inner location words (replace spaces with NBSP and hyphens with non-breaking hyphen)
    def _protect_inner(m):
        prefix = m.group(1)
        name = m.group(2)
        # Default behavior: make hyphens non-breaking and spaces NBSP
        name = _nobreak_hyphens(name)
        name = re.sub(r"\s+", "\u00A0", name)
        return f"{prefix}\u00A0{name}"

    # 2) Standalone 'м. <City>' and 'селище <Name>' and 'с-ще <Name>' (case-insensitive)
    out = re.sub(r"(?i)\b(м\.)\s+([A-Za-zА-Яа-яІіЇїЄєҐґ\-\s]+)", _protect_inner, out)
    out = re.sub(r"(?i)\b(селище)\s+([A-Za-zА-Яа-яІіЇїЄєҐґ\-\s]+)", _protect_inner, out)
    out = re.sub(r"(?i)\b(с-ще)\s+([A-Za-zА-Яа-яІіЇїЄєҐґ\-\s]+)", _protect_inner, out)

    # 3) Parenthesized '(м. <City>)' and '(селище <Name>)' — protect inside parentheses too
    def _protect_paren(m):
        inner = _protect_inner(m)
        return f"({inner})"

    out = re.sub(r"(?i)\((м\.)\s+([A-Za-zА-Яа-яІіЇїЄєҐґ\-\s]+)\)", _protect_paren, out)
    out = re.sub(r"(?i)\((селище)\s+([A-Za-zА-Яа-яІіЇїЄєҐґ\-\s]+)\)", _protect_paren, out)
    out = re.sub(r"(?i)\((с-ще)\s+([A-Za-zА-Яа-яІіЇїЄєҐґ\-\s]+)\)", _protect_paren, out)

    # 4) Specific region phrase: 'у Чернігівській області' should be non-breaking as a whole
    out = re.sub(r"(?i)(?<!\w)у\s+Чернігівській\s+області(?!\w)", "у\u00A0Чернігівській\u00A0області", out)

    # 5) Parentheses should not be separated from adjacent words
    # Prevent '(' from starting a new line alone and from breaking after it
    out = out.replace(" (", "\u00A0(")
    out = re.sub(r"\(\u2060?", "(\u2060", out)
    # Prevent ')' from wrapping alone at the end of a line
    out = re.sub(r"\u2060?\)", "\u2060)", out)

    # Finally, normalize hyphens globally where appropriate (e.g., Михайло-Коцюбинське)
    out = _nobreak_hyphens(out)
    return out

# --- Екранний ввід пароля --- #
def generate_doc(template_path, dodatok_type, executor, selected_date=None, user_state=None):
    tel_mode = bool(user_state.get("tel_mode")) if isinstance(user_state, dict) else False
    phone_mapping = get_phone_mapping() if tel_mode else {}
    doc = Document(template_path)
    if selected_date:
        date_str = format_date(selected_date)
    else:
        date_str = format_date(datetime.now())
    for paragraph in doc.paragraphs:
        if '{дата}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{дата}', date_str)
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '{дата}' in cell.text:
                    cell.text = cell.text.replace('{дата}', date_str)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table = doc.tables[0]
    if dodatok_type == "1":
        unit_key = "Підрозділ"
        person_fields = ["Спеціальне звання", "Посада", "ПІБ"]
        num_columns = 5 if tel_mode else 4
    else:
        unit_key = "Підрозділ, що комплектує"
        person_fields = ["Марка техніки", "Номерний знак"]
        num_columns = 5
    data = get_sheet_data(dodatok_type)
    quantities_by_unit = user_state.get("quantities_by_unit", {})
    i = 1
    for unit in user_state["unit_order"]:
        indices = user_state["selected_indices_by_unit"].get(unit, [])
        if not indices:
            continue
        persons = [row for row in data if row.get(unit_key, "") == unit]
        unit_people = [persons[j] for j in indices]
        if dodatok_type == "2":
            quantities = quantities_by_unit.get(unit, [])
        if unit_people:
            row = table.add_row()
            cell = row.cells[0]
            merged_cell = cell.merge(row.cells[num_columns - 1])
            try:
                display_unit = re.sub(r'^(\d+)\s+', '\\1\u00A0', unit)
            except Exception:
                display_unit = unit.replace(' ', '\u00A0', 1)
            merged_cell.text = protect_no_wrap(display_unit)
            merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in merged_cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D3D3D3')
            merged_cell._tc.get_or_add_tcPr().append(shd)
            set_cell_border(
                merged_cell,
                top={"sz": 4, "val": "single"},
                bottom={"sz": 4, "val": "single"},
                left={"sz": 4, "val": "single"},
                right={"sz": 4, "val": "single"}
            )
            for k, person in enumerate(unit_people):
                row = table.add_row().cells
                row[0].text = str(i)
                if dodatok_type == "1":
                    for j, field in enumerate(person_fields, 1):
                        value = person.get(field, "")
                        if field == "Посада" and isinstance(value, str):
                            try:
                                pattern = r'(\d{1,3})([\s\u00A0]+)(ДПР(?:Ч|З|П))'
                                def _nobreak(m):
                                    return f"{m.group(1)}\u00A0\u2060{m.group(3)}"
                                value = re.sub(pattern, _nobreak, value)
                            except Exception:
                                pass
                        if field == "Спеціальне звання" and isinstance(value, str) and value:
                            value = value[0].upper() + value[1:]
                        row[j].text = protect_no_wrap(value)
                else:
                    row[1].text = protect_no_wrap(unit)
                    row[2].text = person.get(person_fields[0], "")
                    row[3].text = person.get(person_fields[1], "")
                    row[4].text = str(quantities[k])
                    row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if tel_mode and dodatok_type == "1":
                    phone_value = ""
                    pib_value = person.get("ПІБ", "") if isinstance(person, dict) else ""
                    key = _normalize_lookup_key(pib_value)
                    if key:
                        phone_value = phone_mapping.get(key, "")
                    if phone_value:
                        phone_value = re.sub(r"[\s\-]+", "", phone_value)
                        if phone_value and not phone_value.startswith("0"):
                            phone_value = "0" + phone_value
                    phone_cell_index = len(person_fields) + 1
                    row[phone_cell_index].text = protect_no_wrap(phone_value)
                i += 1
                for c in row:
                    set_cell_border(
                        c,
                        top={"sz": 4, "val": "single"},
                        bottom={"sz": 4, "val": "single"},
                        left={"sz": 4, "val": "single"},
                        right={"sz": 4, "val": "single"}
                    )
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(14)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 4, "val": "single"},
                bottom={"sz": 4, "val": "single"},
                left={"sz": 4, "val": "single"},
                right={"sz": 4, "val": "single"}
            )
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for _ in range(3):
        doc.add_paragraph("")
    executor_file = EXECUTORS[executor]
    executor_doc = Document(executor_file)
    for para in executor_doc.paragraphs:
        new_para = doc.add_paragraph()
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
    filename = f"Додаток_{dodatok_type}.docx"
    pdf_filename = f"Додаток_{dodatok_type}.pdf"
    doc.save(filename)
    # Disable auto hyphenation globally to avoid mid-word breaks
    try:
        disable_auto_hyphenation_document(doc)
        doc.save(filename)
    except Exception:
        pass
    pdf_generated = False
    try:
        subprocess.run([LIBREOFFICE_PATH, '--headless', '--convert-to', 'pdf', filename], check=True, timeout=30)
        pdf_generated = True
    except subprocess.TimeoutExpired:
        print(f"PDF conversion timed out after 30 seconds")
    except Exception as e:
        print(f"PDF conversion failed: {e}")
    time.sleep(2)
    return filename, pdf_filename if pdf_generated else None

def generate_nakaz(user_state, dodatok_type="1"):
    if user_state.get("is_vchera"):
        nakaz_template = os.path.join(BASE_DIR, "nakaz_vchera.docx")
    else:
        nakaz_template = os.path.join(BASE_DIR, "nakaz.docx")
        
    if not os.path.exists(nakaz_template):
        raise FileNotFoundError(f"Template not found: {os.path.basename(nakaz_template)}")
    
    doc = Document(nakaz_template)

    # Use user-selected values directly; fields are independent
    selected_izchogo = user_state.get("izchogo", "")
    selected_sho = user_state.get("sho", "")
    selected_kontrol = user_state.get("kontrol", "")
    izchogo_value = selected_izchogo
    sho_value = sanitize_sho_text(selected_sho)
    kontrol_value = selected_kontrol

    units_display = [protect_no_wrap(re.sub(r'^(\d+)\s+', '\\1\u00A0', u)) for u in user_state.get("pidrozdily_units", [])]
    if units_display:
        s = list(units_display[0])
        for idx, ch in enumerate(s):
            if ch.isalpha():
                s[idx] = ch.upper()
                break
        units_display[0] = "".join(s)
    addresses_list = user_state.get("addresses_list") or []
    adresa_value = user_state.get("adresa", "")
    if addresses_list:
        adresa_value = _format_addresses_block(addresses_list)

    placeholders = {
        "{1date}": format_date(user_state.get("1date", datetime.now())),
        "{2date}": format_date(user_state.get("2date", datetime.now())),
        "{izchogo}": protect_no_wrap(izchogo_value),
        "{sho}": protect_no_wrap(sho_value),
        "{kontrol}": protect_no_wrap(kontrol_value),
        "{adresa}": protect_no_wrap(adresa_value),
        "{pidrozdily}": format_list(units_display),
    }
    # Build {nachalnyky} from user selections (columns O, P, Q)
    try:
        items = user_state.get("nachalnyky_items", [])
        sel_idx = user_state.get("nachalnyky_selected_indices", [])
        chosen = []
        for i in sel_idx:
            if 0 <= i < len(items):
                chosen.append(items[i].get("text", ""))
        # Capitalize first alphabetic character of the first item only
        if chosen:
            s = list(chosen[0])
            for idx, ch in enumerate(s):
                if ch.isalpha():
                    s[idx] = ch.upper()
                    break
            chosen[0] = "".join(s)
        # Join with commas, similar to format_list but without 'та'
        placeholders["{nachalnyky}"] = ", ".join([c for c in chosen if c])
    except Exception:
        placeholders["{nachalnyky}"] = ""

    # Build {kgp} from Google Sheet 1 for ONLY the single selected person:
    # columns I (8), K (10), J (9) in that order; keep J case as-is
    # Use warm-cached raw values of sheet1 if available for speed
    cached_vals = SHEET_CACHE.get("1_values")
    if cached_vals:
        values = cached_vals[0]
    else:
        try:
            creds = Credentials.from_service_account_file(CREDENTIALS_FILE, scopes=SCOPES)
            client = gspread.authorize(creds)
            sheet = client.open_by_url(GOOGLE_SHEET_URL1).sheet1
            values = sheet.get_all_values() or []
        except Exception:
            values = []

    kgp_text = ""
    kgp2_text = ""
    unit_order = user_state.get("unit_order", [])
    selected_by_unit = user_state.get("selected_indices_by_unit", {})
    if values and unit_order and selected_by_unit:
        data_rows = values[1:]
        # find first unit with a selected index
        picked_unit = None
        picked_index = None
        for u in unit_order:
            indices = selected_by_unit.get(u, [])
            if indices:
                picked_unit = u
                picked_index = indices[0]
                break
        if picked_unit is not None and picked_index is not None:
            # filter raw rows by unit column F == picked_unit, preserving order
            unit_rows = [r for r in data_rows if (len(r) > 5 and r[5].strip() == picked_unit)]
            if 0 <= picked_index < len(unit_rows):
                row = unit_rows[picked_index]
                col_i = (row[8].strip() if len(row) > 8 else "")
                col_h = (row[7].strip() if len(row) > 7 else "")
                col_k = (row[10].strip() if len(row) > 10 else "")
                col_j = (row[9].strip() if len(row) > 9 else "")
                # Compose: "I H K, J"; keep J case
                parts_left = []
                if col_i:
                    parts_left.append(col_i)
                if col_h:
                    parts_left.append(col_h)
                if col_k:
                    parts_left.append(col_k)
                left = " ".join(p for p in parts_left if p)
                if left and col_j:
                    kgp_text = f"{left}, {col_j}"
                else:
                    kgp_text = left or col_j

                # Build kgp2_text from R (17), S (18), T (19)
                col_r = (row[17].strip() if len(row) > 17 else "")
                col_s = (row[18].strip() if len(row) > 18 else "")
                col_t = (row[19].strip() if len(row) > 19 else "")
                parts_left2 = []
                if col_r:
                    parts_left2.append(col_r)
                if col_s:
                    parts_left2.append(col_s)
                left2 = " ".join(p for p in parts_left2 if p)
                if left2 and col_t:
                    kgp2_text = f"{left2}, {col_t}"
                else:
                    kgp2_text = left2 or col_t

    # Respect personal override; otherwise prefix with 'покласти на '
    if user_state.get("kgp_text_override"):
        final_kgp = user_state["kgp_text_override"]
    else:
        final_kgp = f"покласти на {kgp_text}" if kgp_text else ""
    
    if user_state.get("kgp2_text_override"):
        final_kgp2 = user_state["kgp2_text_override"]
    else:
        final_kgp2 = kgp2_text if kgp2_text else ""

    placeholders["{kgp}"] = protect_no_wrap(final_kgp)
    placeholders["{kgp2}"] = protect_no_wrap(final_kgp2)

    # Replace placeholders preserving images: full-paragraph replacement; preserve drawing runs
    def replace_in_paragraph(paragraph, placeholders):
        def _para_has_drawing(p) -> bool:
            try:
                for r in p.runs:
                    rr = r._r
                    if rr.xpath('.//w:drawing') or rr.xpath('.//w:pict'):
                        return True
            except Exception:
                return False
            return False
        try:
            has_drawing = _para_has_drawing(paragraph)
            # Build full paragraph text and replace across the whole text (handles placeholders split across runs)
            text = paragraph.text or ""
            changed = False
            for key, value in placeholders.items():
                if key in text:
                    text = text.replace(key, str(value) if value is not None else "")
                    changed = True
            if changed:
                # Write back into runs while preserving drawing runs
                first_text_written = False
                non_image_runs = []
                if paragraph.runs:
                    for run in paragraph.runs:
                        rr = getattr(run, "_r", None)
                        has_pic = False
                        try:
                            has_pic = bool(rr is not None and (rr.xpath('.//w:drawing') or rr.xpath('.//w:pict')))
                        except Exception:
                            has_pic = False
                        if has_pic:
                            # Keep image runs intact
                            continue
                        non_image_runs.append(run)
                    if non_image_runs:
                        for idx, run in enumerate(non_image_runs):
                            run.text = text if idx == 0 else ""
                        # Also clear text in other non-image runs not listed (already covered)
                    else:
                        # All runs are images or there are no text runs; append a new text run
                        try:
                            new_run = paragraph.add_run()
                            new_run.text = text
                        except Exception:
                            # Fallback to setting paragraph.text if add_run fails
                            paragraph.text = text
                else:
                    # Paragraph without runs: just set text
                    paragraph.text = text
        except Exception:
            # Fail-safe: do not modify structure if anything goes wrong
            pass

    # Body paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, placeholders)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, placeholders)
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            for paragraph in nested_cell.paragraphs:
                                replace_in_paragraph(paragraph, placeholders)

    # Headers and footers
    try:
        for section in doc.sections:
            # Header
            hdr = section.header
            for p in hdr.paragraphs:
                replace_in_paragraph(p, placeholders)
            for t in hdr.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_in_paragraph(p, placeholders)
            # Footer
            ftr = section.footer
            for p in ftr.paragraphs:
                replace_in_paragraph(p, placeholders)
            for t in ftr.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_in_paragraph(p, placeholders)
    except Exception:
        pass

    # Add pidpys document if selected
    pidpys_key = user_state.get("pidpys")
    if pidpys_key and pidpys_key in PIDPYS:
        pidpys_file = PIDPYS[pidpys_key]
        try:
            pidpys_doc = Document(pidpys_file)
            # Add spacing before pidpys
            for _ in range(3):
                doc.add_paragraph("")
            # Copy all paragraphs from pidpys_doc to main doc
            for para in pidpys_doc.paragraphs:
                new_para = doc.add_paragraph()
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
        except Exception as e:
            print(f"Помилка додавання pidpys файлу: {e}")

    # Save document
    filename = "Наказ.docx"
    pdf_filename = "Наказ.pdf"

    os.makedirs("generated_orders", exist_ok=True)
    filepath = os.path.join("generated_orders", filename)
    pdf_filepath = os.path.join("generated_orders", pdf_filename)

    doc.save(filepath)
    # Disable auto hyphenation globally to avoid mid-word breaks
    try:
        disable_auto_hyphenation_document(doc)
        doc.save(filepath)
    except Exception:
        pass

    # Convert to PDF if possible
    pdf_generated = False
    if os.path.exists(LIBREOFFICE_PATH):
        try:
            subprocess.run([
                LIBREOFFICE_PATH,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', 'generated_orders',
                filepath
            ], check=True, timeout=45)
            pdf_generated = True
        except subprocess.TimeoutExpired:
            # Якщо конвертація PDF займає занадто довго – пропускаємо її, щоб бот не "зависав"
            _log_debug("[PDF] Conversion timed out for dodatok document")
        except Exception as e:
            print(f"PDF conversion failed: {e}")

    return filepath, pdf_filepath if pdf_generated else None

# --- Календар --- #
uk_month_nom = ["січень", "лютий", "березень", "квітень", "травень", "червень", "липень", "серпень", "вересень", "жовтень", "листопад", "грудень"]
uk_month_gen = ["січня", "лютого", "березня", "квітня", "травня", "червня", "липня", "серпня", "вересня", "жовтня", "листопада", "грудня"]

def create_calendar(year=None, month=None):
    if year is None:
        year = datetime.now().year
    if month is None:
        month = datetime.now().month
    kb = InlineKeyboardBuilder()
    month_name = calendar.month_name[month]
    kb.button(text=f"{month_name} {year}", callback_data="ignore")
    kb.adjust(1)
    days = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Нд"]
    for day in days:
        kb.button(text=day, callback_data="ignore")
    month_calendar = calendar.monthcalendar(year, month)
    today = datetime.now().date()
    for week in month_calendar:
        for day in week:
            if day == 0:
                kb.button(text=" ", callback_data="ignore")
            else:
                btn_text = str(day)
                if today == datetime(year, month, day).date():
                    btn_text = f"📍 {btn_text}"
                kb.button(text=btn_text, callback_data=f"date_{year}_{month}_{day}")
    prev_month = month - 1 if month > 1 else 12
    prev_year = year if month > 1 else year - 1
    next_month = month + 1 if month < 12 else 1
    next_year = year if month < 12 else year + 1
    kb.button(text="← Попередній", callback_data=f"nav_{prev_year}_{prev_month}")
    kb.button(text="Сьогодні", callback_data="date_today")
    kb.button(text="Наступний →", callback_data=f"nav_{next_year}_{next_month}")
    kb.button(text="✅ Підтвердити дату", callback_data="confirm_date")
    kb.button(text="🔙 Назад", callback_data="back_to_executors")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(7, *[7 for _ in range(len(month_calendar))], 3, 2)
    return kb.as_markup()

def create_calendar_for_nakaz(year=None, month=None, date_key="1date"):
    if year is None:
        year = datetime.now().year
    if month is None:
        month = datetime.now().month
    kb = InlineKeyboardBuilder()
    month_name = uk_month_nom[month - 1]
    kb.button(text=f"{month_name} {year}", callback_data="ignore")
    kb.adjust(1)
    days = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Нд"]
    for day in days:
        kb.button(text=day, callback_data="ignore")
    month_calendar = calendar.monthcalendar(year, month)
    today = datetime.now().date()
    for week in month_calendar:
        for day in week:
            if day == 0:
                kb.button(text=" ", callback_data="ignore")
            else:
                btn_text = str(day)
                if today == datetime(year, month, day).date():
                    btn_text = f"📍 {btn_text}"
                kb.button(text=btn_text, callback_data=f"nakaz_date_{date_key}_{year}_{month}_{day}")
    prev_month = month - 1 if month > 1 else 12
    prev_year = year if month > 1 else year - 1
    next_month = month + 1 if month < 12 else 1
    next_year = year if month < 12 else year + 1
    kb.button(text="← Попередній", callback_data=f"nakaz_nav_{date_key}_{prev_year}_{prev_month}")
    kb.button(text="Сьогодні", callback_data=f"nakaz_date_{date_key}_today")
    kb.button(text="Наступний →", callback_data=f"nakaz_nav_{date_key}_{next_year}_{next_month}")
    kb.button(text="✅ Підтвердити дату", callback_data=f"nakaz_confirm_{date_key}")
    kb.button(text="⏭️ Пропустити", callback_data=f"skip_date_{date_key}")
    kb.button(text="🔙 Назад", callback_data="back_to_start")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(7, *[7 for _ in range(len(month_calendar))], 3, 2)
    return kb.as_markup()

def create_rozrakh_calendar(year, month, selected_dates):
    kb = InlineKeyboardBuilder()
    month_name = uk_month_nom[month - 1]
    kb.button(text=f"{month_name} {year}", callback_data="ignore")
    kb.adjust(1)
    days = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Нд"]
    for day in days:
        kb.button(text=day, callback_data="ignore")
    month_calendar = calendar.monthcalendar(year, month)
    today = datetime.now().date()
    for week in month_calendar:
        for day in week:
            if day == 0:
                kb.button(text=" ", callback_data="ignore")
            else:
                date_obj = datetime(year, month, day).date()
                if date_obj in selected_dates:
                    text = f"✅ {day}"
                elif date_obj == today:
                    text = f"📍 {day}"
                else:
                    text = str(day)
                kb.button(text=text, callback_data=f"roz_date_{year}_{month}_{day}")
    prev_month = month - 1 if month > 1 else 12
    prev_year = year if month > 1 else year - 1
    next_month = month + 1 if month < 12 else 1
    next_year = year if month < 12 else year + 1
    kb.button(text="← Попередній", callback_data=f"roz_nav_{prev_year}_{prev_month}")
    kb.button(text="Наступний →", callback_data=f"roz_nav_{next_year}_{next_month}")
    kb.button(text="✅ Підтвердити дати", callback_data="roz_confirm_date")
    kb.button(text="Очистити вибір", callback_data=f"roz_clear_{year}_{month}")
    kb.button(text="🔙 Назад", callback_data="roz_back_to_step1")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(7, *[7 for _ in range(len(month_calendar))], 2, 3)
    return kb.as_markup()

async def show_rozrakh_calendar(callback: types.CallbackQuery, uid: int, year=None, month=None):
    touch_activity(uid)
    selected = user_state[uid].get("rozrakh_selected_dates")
    if not isinstance(selected, set):
        selected = set()
        user_state[uid]["rozrakh_selected_dates"] = selected
    if year is None or month is None:
        if selected:
            min_date = min(selected)
            year, month = min_date.year, min_date.month
        else:
            now = datetime.now()
            year, month = now.year, now.month
    print(f"[show_rozrakh_calendar] User {uid} selected dates: {sorted(selected)}")
    markup = create_rozrakh_calendar(year, month, selected)
    selected_str = ", ".join(sorted(d.strftime("%d.%m.%Y") for d in selected)) if selected else "Нічого"
    text = f"📅 Оберіть дати (одну або декілька для діапазону)\nВибрано: {selected_str}"
    try:
        await callback.message.edit_text(text, reply_markup=markup)
    except TelegramBadRequest:
        await safe_answer_callback(callback)
    except Exception:
        await callback.message.answer(text, reply_markup=markup)

@dp.callback_query(F.data == "back_to_izchogo")
async def back_to_izchogo(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_izchogo_selection(callback, uid)
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "restart")
async def restart(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    clear_workflow_state(uid)
    await start(callback.message)
    await safe_answer_callback(callback)

# --- Розрахунки --- #
def load_docx(filename):
    doc = Document(filename)
    return doc

def merge_documents(filenames):
    if not filenames:
        return None
    master = load_docx(filenames[0])
    composer = Composer(master)
    for filename in filenames[1:]:
        doc_temp = load_docx(filename)
        composer.append(doc_temp)
    return master

async def show_rozrakh_menu_step1(message_or_callback, uid):
    files = ["Чернігів", "Ніжин", "Корюківка", "ГУ"]
    kb = InlineKeyboardBuilder()
    selected = user_state[uid].get("rozrakh_selected_step1", set())
    for f in files:
        text = f"✅ {f}" if f in selected else f
        kb.button(text=text, callback_data=f"rozrakh_step1_{f}")
    if selected:
        kb.button(text="➡️ Далі", callback_data="rozrakh_step1_next")
    kb.button(text="🔙 Назад", callback_data="back_to_start")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    text = "Оберіть перші файли (Чернігів, Ніжин, Корюківка, ГУ):\n\n<i>Файл автоматично обрано та перехід до наступного кроку</i>" if selected else "Оберіть перші файли (Чернігів, Ніжин, Корюківка, ГУ):"
    try:
        if hasattr(message_or_callback, 'message'):
            await message_or_callback.message.edit_text(text, reply_markup=kb.as_markup(), parse_mode="HTML")
        else:
            await message_or_callback.edit_text(text, reply_markup=kb.as_markup(), parse_mode="HTML")
    except Exception:
        await message_or_callback.answer(text, reply_markup=kb.as_markup(), parse_mode="HTML")

async def show_rozrakh_menu_step2(message_or_callback, uid):
    files = ["Коржов", "Усик", "Середа"]
    kb = InlineKeyboardBuilder()
    selected = user_state[uid].get("rozrakh_selected_step2", set())
    for f in files:
        text = f"✅ {f}" if f in selected else f
        kb.button(text=text, callback_data=f"rozrakh_step2_{f}")
    if selected:
        kb.button(text="➡️ Далі", callback_data="rozrakh_step2_next")
    kb.button(text="🔙 Назад", callback_data="rozrakh_step1_back")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    text = "Оберіть другу групу файлів (Коржов, Усик, Середа):\n\n<i>Файл автоматично обрано та перехід до наступного кроку</i>" if selected else "Оберіть другу групу файлів (Коржов, Усик, Середа):"
    try:
        if hasattr(message_or_callback, 'message'):
            await message_or_callback.message.edit_text(text, reply_markup=kb.as_markup(), parse_mode="HTML")
        else:
            await message_or_callback.edit_text(text, reply_markup=kb.as_markup(), parse_mode="HTML")
    except Exception:
        await message_or_callback.answer(text, reply_markup=kb.as_markup(), parse_mode="HTML")

async def show_rozrakh_menu_step3(message_or_callback, uid):
    files = ["Дилін", "Царенко", "Федун", "Павленко"]
    kb = InlineKeyboardBuilder()
    selected = user_state[uid].get("rozrakh_selected_step3", set())
    for f in files:
        text = f"✅ {f}" if f in selected else f
        kb.button(text=text, callback_data=f"rozrakh_step3_{f}")
    if selected:
        kb.button(text="➡️ Далі", callback_data="rozrakh_step3_next")
    kb.button(text="🔙 Назад", callback_data="rozrakh_step2_back")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    text = "Оберіть третю групу файлів (Дилін, Царенко, Федун, Павленко):\n\n<i>Файл автоматично обрано та перехід до наступного кроку</i>" if selected else "Оберіть третю групу файлів (Дилін, Царенко, Федун, Павленко):"
    try:
        if hasattr(message_or_callback, 'message'):
            await message_or_callback.message.edit_text(text, reply_markup=kb.as_markup(), parse_mode="HTML")
        else:
            await message_or_callback.edit_text(text, reply_markup=kb.as_markup(), parse_mode="HTML")
    except Exception:
        await message_or_callback.answer(text, reply_markup=kb.as_markup(), parse_mode="HTML")

def generate_rozrakh_docx_pdf(uid):
    base_dir = os.path.join(os.path.dirname(__file__), "rozrakh")
    step1_files = user_state[uid].get("rozrakh_selected_step1", set())
    step2_files = user_state[uid].get("rozrakh_selected_step2", set())
    step3_files = user_state[uid].get("rozrakh_selected_step3", set())
    number = user_state[uid].get("rozrakh_nomber", 1)
    doc_files = []
    for fname in step1_files:
        path = os.path.join(base_dir, fname + ".docx")
        if os.path.exists(path):
            doc_files.append(path)
    for fname in step2_files:
        path = os.path.join(base_dir, fname + ".docx")
        if os.path.exists(path):
            doc_files.append(path)
    for fname in step3_files:
        path = os.path.join(base_dir, fname + ".docx")
        if os.path.exists(path):
            doc_files.append(path)
    merged_doc = merge_documents(doc_files)
    if not merged_doc:
        return None, None
    date_text = user_state[uid].get("rozrakh_date_text", "")
    if date_text:
        for p in merged_doc.paragraphs:
            if "{date}" in p.text:
                p.text = p.text.replace("{date}", date_text)
        for table in merged_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "{date}" in p.text:
                            p.text = p.text.replace("{date}", date_text)
    number = user_state[uid].get("rozrakh_nomber", 1)
    # Замена плейсхолдера {nomber} (причем учитывая регистр, если нужно)
    replaced = False
    def replace_in_paragraph(paragraph, old_text, new_text):
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
    for p in merged_doc.paragraphs:
        replace_in_paragraph(p, "{nomber}", str(number))
        if "{nomber}" in p.text:
            p.text = p.text.replace("{nomber}", str(number))
            replaced = True
    for table in merged_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, "{nomber}", str(number))
                    if "{nomber}" in p.text:
                        p.text = p.text.replace("{nomber}", str(number))
                        replaced = True
    # Если плейсхолдер не найден, можно дозаписать количество в конец
    if not replaced:
        merged_doc.add_paragraph("")
        merged_doc.add_paragraph("")
        merged_doc.add_paragraph("")
        para = merged_doc.add_paragraph(f"Кількість додатків: {number}")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    docx_filename = "Розрахунки.docx"
    pdf_filename = "Розрахунки.pdf"
    merged_doc.save(docx_filename)
    pdf_generated = False
    try:
        subprocess.run([LIBREOFFICE_PATH, '--headless', '--convert-to', 'pdf', docx_filename], check=True, timeout=45)
        pdf_generated = True
    except subprocess.TimeoutExpired:
        print(f"PDF conversion timed out (rozrakh)")
    except Exception as e:
        print(f"PDF conversion failed: {e}")
    time.sleep(2)
    return docx_filename, pdf_filename if pdf_generated else None

# --- Очищення неактивності --- #
INACTIVITY_SECONDS = 1800  # 30 минут

async def inactivity_cleaner():
    while True:
        now_ts = time.time()
        stale = []
        for uid, st in list(user_state.items()):
            last = st.get('last_activity', now_ts)
            if now_ts - last > INACTIVITY_SECONDS:
                stale.append(uid)
        for uid in stale:
            try:
                sent = user_state.get(uid, {}).get('sent_message_ids', [])
                for chat_id, mid in sent:
                    try:
                        await bot.delete_message(chat_id, mid)
                    except Exception:
                        pass
            except Exception:
                pass
            user_state[uid] = {'authenticated': False, 'state': 'enter_password', 'sent_message_ids': []}
            try:
                await prompt_password(uid, uid)
            except Exception:
                pass
        await asyncio.sleep(60)

# --- Меню після генерації --- #
async def show_post_generate_menu(callback: types.CallbackQuery, uid: int):
    kb = InlineKeyboardBuilder()
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    try:
        await callback.message.answer("✅ Файли згенеровано.", reply_markup=kb.as_markup())
    except Exception:
        await callback.message.answer("✅ Файли згенеровано.", reply_markup=kb.as_markup())

# --- Безпечне відповідання на callback --- #
async def safe_answer_callback(callback: types.CallbackQuery, text: str = None, show_alert: bool = False):
    """Безпечне відповідання на callback з обробкою помилок"""
    try:
        await callback.answer(text, show_alert=show_alert)
    except TelegramBadRequest as e:
        if "query is too old" in str(e):
            print("⚠️ Ігноруємо застарілий callback")
        else:
            print(f"⚠️ Помилка при відповіді на callback: {e}")
    except Exception as e:
        print(f"⚠️ Не вдалося відповісти на callback: {e}")

# --- Глобальний обробник помилок --- #
@dp.errors()
async def global_error_handler(event: types.ErrorEvent):
    """Глобальний обробник помилок"""
    error = event.exception
    print(f"🔥 Глобальна помилка: {error}")

    if isinstance(error, TelegramBadRequest):
        if "query is too old" in str(error):
            print("⚠️ Ігноруємо застарілий callback")
            return True

    return False

# --- Хендлери --- #
@dp.message(Command("start"))
async def start(message: types.Message):
    uid = message.from_user.id
    touch_activity(uid)
    if uid not in user_state:
        user_state[uid] = {'authenticated': False, 'state': None}
    user_state[uid]['state'] = None
    # Force cleanup of any lingering password UI
    st = user_state.get(uid, {})
    try:
        chat_id, mid = st.get('pwd_msg', (None, None))
        if chat_id and mid:
            await bot.delete_message(chat_id, mid)
            st['pwd_msg'] = (None, None)
    except Exception:
        pass
    # Формируем статус бота для користувача
    try:
        if CACHE_READY:
            status_line = "✅ Бот онлайн, кеш завантажено."
        else:
            if CACHE_LAST_ERROR:
                # Кеш не прогрівся або сталася помилка (можливо, немає інтернету)
                status_line = f"⚠️ Обмежений режим: кеш не завантажено ({CACHE_PROGRESS}%)."
            else:
                status_line = f"⏳ Завантажую кеш… {CACHE_PROGRESS}%"
    except Exception:
        status_line = "ℹ️ Статус кеша наразі невідомий."

    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📜 Наказ на залучення", callback_data="nakaz")],
        [InlineKeyboardButton(text="📜 Створити наказ за вчора", callback_data="nakaz_vchera")],
        [InlineKeyboardButton(text="📜 Наказ відпустки/лікарняні", callback_data="nakaz_vidpustky")],
        [InlineKeyboardButton(text="🔄 Оновити о/с (AI відмінювання)", callback_data="update_os")],
        [InlineKeyboardButton(text="📄 Створити Додаток 1", callback_data="dodatok_1")],
        [InlineKeyboardButton(text="📄 Створити Додаток 1 + тел", callback_data="dodatok_1_tel")],
        [InlineKeyboardButton(text="📄 Створити Додаток 2", callback_data="dodatok_2")],
        [InlineKeyboardButton(text="📂 Розрахунки", callback_data="rozrakh")],
        [InlineKeyboardButton(text="🤖 Запитай у бота ШІ", callback_data="ai_chat")],
        [InlineKeyboardButton(text="⚙️ Вибір моделі ШІ", callback_data="ai_model_menu")],
        [InlineKeyboardButton(text="🔄 Оновити кеш", callback_data="refresh_cache")],
        [InlineKeyboardButton(text="⚙️ Налаштування", callback_data="settings"), InlineKeyboardButton(text="ℹ️ Допомога", callback_data="help")],
        [InlineKeyboardButton(text="🔄 Рестарт", callback_data="restart")]
    ])
    text = status_line + "\n\n" + "👋 Вітаю! Оберіть дію:"
    sent = await message.answer(text, reply_markup=kb)
    try:
        user_state.setdefault(uid, {}).setdefault('sent_message_ids', []).append((message.chat.id, sent.message_id))
    except Exception:
        pass


@dp.message(Command("status"))
async def status_command(message: types.Message):
    """Показати поточний статус бота (онлайн/кеш/мережа)"""
    uid = message.from_user.id
    touch_activity(uid)
    try:
        if CACHE_READY:
            cache_part = "✅ Кеш завантажено, робота з таблицями доступна."
        else:
            if CACHE_LAST_ERROR:
                cache_part = f"⚠️ Кеш не завантажено ({CACHE_PROGRESS}%). Можлива помилка мережі або доступу до Google Sheets."
            else:
                cache_part = f"⏳ Виконується прогрів кеша… {CACHE_PROGRESS}%"
    except Exception:
        cache_part = "ℹ️ Статус кеша наразі невідомий."

    text = "🤖 Статус бота:\n" + cache_part
    await message.answer(text)

@dp.callback_query(F.data == "ai_model_menu")
async def ai_model_menu(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    
    # Список безкоштовних/прев'ю моделей
    available_models = [
        ("gemini-2.0-flash", "⚡ Gemini 2.0 Flash"),
        ("gemini-2.0-flash-lite", "🍃 Gemini 2.0 Flash-Lite"),
        ("gemini-2.5-flash-lite", "🚀 Gemini 2.5 Flash-Lite"),
        ("gemini-1.5-flash", "🌩️ Gemini 1.5 Flash"),
        ("gemma-3-4b-it", "🧠 Gemma 3 4B"),
    ]
    
    current_model = _get_user_model(uid)
    
    kb = InlineKeyboardBuilder()
    for model_id, display_name in available_models:
        label = f"✅ {display_name}" if model_id == current_model else display_name
        kb.button(text=label, callback_data=f"set_ai_model_{model_id}")
    
    kb.button(text="🔙 Назад", callback_data="back_to_start")
    kb.adjust(1)
    
    text = f"🤖 **Налаштування моделі ШІ**\n\nПоточна модель: `{current_model}`\n\nОберіть модель для генерації преамбул та відповідей:"
    try:
        await callback.message.edit_text(text, reply_markup=kb.as_markup(), parse_mode="Markdown")
    except Exception:
        await callback.message.answer(text, reply_markup=kb.as_markup(), parse_mode="Markdown")
    await safe_answer_callback(callback)

@dp.callback_query(F.data.startswith("set_ai_model_"))
async def set_ai_model(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    model_id = callback.data.replace("set_ai_model_", "")
    
    user_state.setdefault(uid, {})["ai_model"] = model_id
    await safe_answer_callback(callback, f"Модель {model_id} обрана!", show_alert=True)
    await ai_model_menu(callback)

@dp.callback_query(F.data == "ai_chat")
async def ai_chat_entry(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state.setdefault(uid, {})["state"] = "ai_chat"
    kb = InlineKeyboardBuilder()
    kb.button(text="🔙 Назад", callback_data="back_to_start")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(2)
    try:
        await callback.message.edit_text("Напишіть запит для ШІ:", reply_markup=kb.as_markup())
    except Exception:
        await callback.message.answer("Напишіть запит для ШІ:", reply_markup=kb.as_markup())
    await safe_answer_callback(callback)

@dp.message()
async def handle_message(message: types.Message):
    uid = message.from_user.id
    touch_activity(uid)
    if uid not in user_state:
        user_state[uid] = {'authenticated': False, 'state': None}
    state = user_state[uid].get('state')
    
    # Обработка ввода адреса из временного файла
    if state == "enter_adresa":
        address_text = (message.text or "").strip()
        if not address_text:
            await _prompt_manual_address_input(message, uid)
        else:
            await _append_address_and_continue(message, uid, address_text)
        return
        
    global password
    if state == 'enter_password':
        # Ignore text entries; enforce on-screen keyboard usage
        await message.answer("ℹ️ Використовуйте екранну клавіатуру нижче для введення пароля.")
        return
    elif state == 'ai_chat':
        query = (message.text or '').strip()
        if not query:
            await message.answer("Напишіть запит для ШІ.")
            return
        loading = await message.answer("Думаю… 🤖")
        try:
            reply = await asyncio.to_thread(_call_gemini_chat_simple, query, uid)
        except Exception:
            reply = ""
        try:
            await loading.delete()
        except Exception:
            pass
        kb = InlineKeyboardBuilder()
        kb.button(text="🔙 Назад", callback_data="back_to_start")
        kb.button(text="🔁 Рестарт", callback_data="restart")
        kb.adjust(2)
        if reply:
            await message.answer(reply, reply_markup=kb.as_markup())
        else:
            await message.answer("Не вдалось отримати відповідь. Спробуйте ще раз.", reply_markup=kb.as_markup())
        return
    elif state == "enter_rozrakh_nomber":
        try:
            number = int(message.text.strip())
            if number < 1:
                raise ValueError
            user_state[uid]["rozrakh_nomber"] = number
            user_state[uid]["state"] = None
            await message.answer(f"✅ Кількість додатків встановлено: {number}")
            # Показуємо кнопку для генерації
            kb = InlineKeyboardBuilder()
            kb.button(text="✅ Сформувати розрахунки", callback_data="rozrakh_generate")
            kb.button(text="🔁 Рестарт", callback_data="restart")
            kb.adjust(1)
            await message.answer("Натисніть кнопку для формування розрахунків:", reply_markup=kb.as_markup())
        except Exception:
            await message.answer("❌ Введіть будь ласка ціле позитивне число для кількості додатків")
    elif state == 'change_password_old':
        if message.text.strip() == password:
            await message.answer("🔑 Введіть новий пароль:")
            user_state[uid]['state'] = 'change_password_new'
        else:
            await message.answer("❌ Невірний поточний пароль!")
    elif state == 'sho_ai_seed':
        # Save user's seed/details and immediately generate AI suggestions
        user_state[uid]["sho_ai_seed"] = message.text.strip()
        user_state[uid]['state'] = None
        seed = user_state.get(uid, {}).get("sho_ai_seed", "")
        iz = user_state.get(uid, {}).get("izchogo", "")
        try:
            suggestions = await asyncio.to_thread(_call_gemini_generate_sho, seed, iz, uid)
        except Exception:
            suggestions = []
        if not suggestions:
            kb = InlineKeyboardBuilder()
            kb.button(text="🔁 Спробувати ще раз", callback_data="sho_ai_generate")
            kb.button(text="🔙 Назад", callback_data="back_to_start")
            kb.adjust(2)
            await message.answer("Не вдалося отримати варіанти ШІ. Спробуйте уточнити деталі та спробуйте ще раз.", reply_markup=kb.as_markup())
            return
        user_state[uid]["sho_ai_suggestions"] = suggestions
        lines = [f"{i+1}. {s}" for i, s in enumerate(suggestions)]
        text = "Варіанти від ШІ:\n\n" + "\n\u2009\n".join(lines)
        kb = InlineKeyboardBuilder()
        for i, _ in enumerate(suggestions):
            kb.button(text=str(i+1), callback_data=f"sho_ai_pick_{i}")
        kb.button(text="🔁 Згенерувати ще раз", callback_data="sho_ai_generate")
        kb.button(text="🔙 Назад", callback_data="back_to_start")
        kb.button(text="🔁 Рестарт", callback_data="restart")
        kb.adjust(2)
        await message.answer(text, reply_markup=kb.as_markup())
    elif state == "vidp_enter_days":
        # Manual input for vidpustky days
        try:
            days = int(message.text.strip())
            if days < 1 or days > 365:
                raise ValueError
            user_state[uid]["vidp_days"] = days
            user_state[uid]["state"] = None
            await show_vidp_start_calendar(message, uid)
        except Exception:
            await message.answer("❌ Введіть ціле число від 1 до 365.")
    elif state == "vidp_enter_destination":
        # Manual input for destination
        dest = message.text.strip()
        user_state[uid]["vidp_destination"] = dest
        user_state[uid]["state"] = None
        await show_vidp_confirm_person(message, uid)
    elif state == "vidp_enter_ln":
        # LN number input
        text = message.text.strip()
        if text:
            user_state[uid]["vidp_ln_number"] = text
            user_state[uid]["vidp_ln_date"] = ""
        else:
            user_state[uid]["vidp_ln_number"] = ""
            user_state[uid]["vidp_ln_date"] = ""
        user_state[uid]["state"] = None
        await show_vidp_supervisor_selection(message, uid)
    else:
        if not is_authenticated(uid):
            await prompt_password(message, uid)
        else:
            # Якщо користувач автентифікований і надіслав довільне повідомлення,
            # просто покажемо головне меню замість повідомлення про невідому команду
            await start(message)

# Хендлери для розрахунків
@dp.callback_query(lambda c: c.data and c.data.startswith("roz_date_"))
async def roz_select_date(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    print(f"[roz_select_date] Callback data: {callback.data}")
    parts = callback.data.split("_")
    if len(parts) == 5:
        _, _, year_str, month_str, day_str = parts
        year, month, day = int(year_str), int(month_str), int(day_str)
        date_obj = datetime(year, month, day).date()
        if uid not in user_state:
            user_state[uid] = {}
        selected = user_state[uid].get("rozrakh_selected_dates")
        if not isinstance(selected, set):
            selected = set()
        if date_obj in selected:
            selected.remove(date_obj)
        else:
            selected.add(date_obj)
        user_state[uid]["rozrakh_selected_dates"] = selected.copy() # Обновляем состояние явно
        print(f"[roz_select_date] User {uid} selected dates: {sorted(selected)}")
        await show_rozrakh_calendar(callback, uid, year, month)
    await safe_answer_callback(callback, "Дата обрана")

@dp.callback_query(F.data.startswith("roz_nav_"))
async def roz_navigate_calendar(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    parts = callback.data.split("_")
    if len(parts) == 3:
        _, year_str, month_str = parts
        year, month = int(year_str), int(month_str)
        await show_rozrakh_calendar(callback, uid, year, month)
    await safe_answer_callback(callback)

@dp.callback_query(F.data.startswith("roz_clear_"))
async def roz_clear_dates(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    parts = callback.data.split("_")
    if len(parts) == 3:
        _, year_str, month_str = parts
        year, month = int(year_str), int(month_str)
        user_state[uid]["rozrakh_selected_dates"] = set()
        await show_rozrakh_calendar(callback, uid, year, month)
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "roz_confirm_date")
async def roz_confirm_dates(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    selected = user_state[uid].get("rozrakh_selected_dates", set())
    if not selected:
        await safe_answer_callback(callback, "Оберіть хоча б одну дату!", show_alert=True)
        return
    sorted_dates = sorted(selected)
    min_date = sorted_dates[0]
    max_date = sorted_dates[-1]
    if len(sorted_dates) == 1:
        date_text = f"{min_date.day:02} {uk_month_gen[min_date.month-1]} {min_date.year} року"
    else:
        min_month = uk_month_gen[min_date.month-1]
        max_month = uk_month_gen[max_date.month-1]
        date_text = f"з {min_date.day:02} {min_month} {min_date.year} року по {max_date.day:02} {max_month} {max_date.year} року"
    user_state[uid]["rozrakh_date_text"] = date_text
    await show_rozrakh_menu_step2(callback, uid)
    await safe_answer_callback(callback, f"✅ Дати встановлено: {date_text}")

@dp.callback_query(F.data == "roz_back_to_step1")
async def roz_back_to_step1(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state[uid].pop("rozrakh_selected_dates", None)
    await show_rozrakh_menu_step1(callback, uid)
    await safe_answer_callback(callback)

@dp.callback_query(F.data.startswith("rozrakh_step1_"))
async def rozrakh_select_step1(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    choice = callback.data.replace("rozrakh_step1_", "")
    selected = user_state[uid].setdefault("rozrakh_selected_step1", set())
    if choice in selected:
        selected.remove(choice)
    else:
        selected.add(choice)
    if selected:
        user_state[uid]["rozrakh_selected_dates"] = set()
        now = datetime.now()
        await show_rozrakh_calendar(callback, uid, now.year, now.month)
    else:
        await show_rozrakh_menu_step1(callback, uid)
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "rozrakh_step1_next")
async def rozrakh_step1_next(callback: types.CallbackQuery):
    uid = callback.from_user.id
    if not user_state[uid].get("rozrakh_selected_step1"):
        await safe_answer_callback(callback, "Оберіть хоча б один файл!", show_alert=True)
        return
    user_state[uid]["rozrakh_selected_dates"] = set()
    now = datetime.now()
    await show_rozrakh_calendar(callback, uid, now.year, now.month)
    await safe_answer_callback(callback)

@dp.callback_query(F.data.startswith("rozrakh_step2_"))
async def rozrakh_select_step2(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    choice = callback.data.replace("rozrakh_step2_", "")
    selected = user_state[uid].setdefault("rozrakh_selected_step2", set())
    if choice in selected:
        selected.remove(choice)
    else:
        selected.add(choice)
    if selected:
        await show_rozrakh_menu_step3(callback, uid)
    else:
        await show_rozrakh_menu_step2(callback, uid)
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "rozrakh_step2_next")
async def rozrakh_step2_next(callback: types.CallbackQuery):
    uid = callback.from_user.id
    if not user_state[uid].get("rozrakh_selected_step2"):
        await safe_answer_callback(callback, "Оберіть хоча б один файл!", show_alert=True)
        return
    await show_rozrakh_menu_step3(callback, uid)
    await safe_answer_callback(callback)

@dp.callback_query(F.data.startswith("rozrakh_step3_"))
async def rozrakh_select_step3(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    choice = callback.data.replace("rozrakh_step3_", "")
    selected = user_state[uid].setdefault("rozrakh_selected_step3", set())
    if choice in selected:
        selected.remove(choice)
    else:
        selected.add(choice)
    if selected:
        user_state[uid]["state"] = "enter_rozrakh_nomber"
        try:
            await callback.message.edit_text("Введіть кількість додатків (число):")
        except Exception:
            await callback.message.answer("Введіть кількість додатків (число):")
    else:
        await show_rozrakh_menu_step3(callback, uid)
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "rozrakh_step3_next")
async def rozrakh_step3_next(callback: types.CallbackQuery):
    uid = callback.from_user.id
    if not user_state[uid].get("rozrakh_selected_step3"):
        await safe_answer_callback(callback, "Оберіть хоча б один файл!", show_alert=True)
        return
    user_state[uid]["state"] = "enter_rozrakh_nomber"
    try:
        await callback.message.edit_text("Введіть кількість додатків (число):")
    except Exception:
        await callback.message.answer("Введіть кількість додатків (число):")
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "rozrakh")
async def start_rozrakh(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    clear_workflow_state(uid)
    user_state[uid]["rozrakh_selected_step1"] = set()
    user_state[uid]["rozrakh_selected_step2"] = set()
    user_state[uid]["rozrakh_selected_step3"] = set()
    user_state[uid]["state"] = None
    await show_rozrakh_menu_step1(callback, uid)
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "rozrakh_step1_back")
async def rozrakh_step1_back(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_rozrakh_menu_step1(callback, uid)
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "rozrakh_step2_back")
async def rozrakh_step2_back(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_rozrakh_menu_step2(callback, uid)
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "rozrakh_generate")
async def rozrakh_generate(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    loading_msg = await callback.message.answer("Генерація документа…")
    stop_event = asyncio.Event()
    spinner_task = asyncio.create_task(_spinner_message(loading_msg, "Генерація документа…", stop_event))
    try:
        docx_filename, pdf_filename = await asyncio.wait_for(
            asyncio.to_thread(generate_rozrakh_docx_pdf, uid),
            timeout=60
        )
    finally:
        stop_event.set()
        try:
            await spinner_task
        except Exception:
            pass
    if not docx_filename:
        await loading_msg.edit_text("Помилка при формуванні документу.")
        return
    await loading_msg.edit_text("Надсилання файлів… 📤")
    caption = f"📄 Розрахунки"
    sent_doc = await callback.message.answer_document(types.FSInputFile(docx_filename), caption=caption)
    try:
        user_state.setdefault(uid, {}).setdefault('sent_message_ids', []).append((callback.message.chat.id, sent_doc.message_id))
    except Exception:
        pass
    if pdf_filename:
        sent_pdf = await callback.message.answer_document(types.FSInputFile(pdf_filename), caption=caption)
        try:
            user_state.setdefault(uid, {}).setdefault('sent_message_ids', []).append((callback.message.chat.id, sent_pdf.message_id))
        except Exception:
            pass
        safe_remove(pdf_filename)
    safe_remove(docx_filename)
    clear_workflow_state(uid)
    try:
        await loading_msg.edit_text("Готово ✅")
        await asyncio.sleep(0.6)
        await loading_msg.delete()
    except Exception:
        pass
    await show_post_generate_menu(callback, uid)
    await safe_answer_callback(callback)

# Хендлери для додатків
@dp.callback_query(F.data.startswith("dodatok_"))
async def select_dodatok(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    parts = callback.data.split("_")
    dodatok_type = parts[1]
    tel_mode = len(parts) > 2 and parts[2] == "tel"
    user_state[uid] = {"dodatok": dodatok_type, 'authenticated': True, 'tel_mode': tel_mode}
    user_state[uid]["mode"] = "grouped"
    user_state[uid]["unit_order"] = []
    user_state[uid]["selected_indices_by_unit"] = {}
    user_state[uid]["quantities_by_unit"] = {}
    await show_unit_selection(callback, uid)
    await safe_answer_callback(callback)

async def show_unit_selection(callback_or_message, uid: int):
    touch_activity(uid)
    if not is_authenticated(uid):
        if isinstance(callback_or_message, types.CallbackQuery):
            await safe_answer_callback(callback_or_message, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        else:
            await callback_or_message.answer("Ви не автентифіковані! Використовуйте /start і введіть пароль.")
        return
    dodatok_type = user_state[uid]["dodatok"]
    if dodatok_type == "1":
        unit_key = "Підрозділ"
    else:
        unit_key = "Підрозділ, що комплектує"
    units = user_state[uid].get("units_list")
    if units is None:
        # Use precomputed units if available for instant rendering
        cache_key = "pre_units1" if dodatok_type == "1" else "pre_units2"
        cached_units = SHEET_CACHE.get(cache_key)
        if cached_units:
            units = cached_units[0]
        else:
            data = await get_sheet_data_async(dodatok_type)
            units = sort_units(list({row.get(unit_key, "") for row in data if row.get(unit_key)}))
        user_state[uid]["units_list"] = units
    # If in KGP mode, restrict units to user-selected kgp_units if provided
    if user_state[uid].get("kgp_mode"):
        allowed = user_state[uid].get("kgp_units", [])
        original_units = units[:]
        if allowed:
            # keep original ordering of 'units', filter only allowed
            units = [u for u in units if u in allowed]
        # If filtering produced empty list, fall back to original full list
        if not units:
            units = original_units
    selected_units = user_state[uid].get("unit_order", [])
    kb = InlineKeyboardBuilder()
    for i, unit in enumerate(units):
        button_text = f"✅ {unit}" if unit in selected_units else unit
        if user_state[uid].get("kgp_mode"):
            callback_data = f"kgp_select_unit_{i}"
        else:
            callback_data = f"select_unit_{i}"
        kb.button(text=button_text, callback_data=callback_data)
    # In KGP mode, allow choosing a personal option as the last item
    if user_state[uid].get("kgp_mode"):
        kb.button(text="буду здійснювати особисто", callback_data="kgp_personal")
    kb.button(text="Назад", callback_data="back_to_dodatok")
    kb.button(text="Рестарт", callback_data="restart")
    kb.adjust(1)
    text = "🏢 Оберіть підрозділ" if not user_state[uid].get("kgp_mode") else "🏢 Оберіть КГП"
    if isinstance(callback_or_message, types.CallbackQuery):
        try:
            await callback_or_message.message.edit_text(text, reply_markup=kb.as_markup())
        except TelegramBadRequest as e:
            if "message is not modified" in str(e):
                await safe_answer_callback(callback_or_message)
            else:
                await callback_or_message.message.answer(text, reply_markup=kb.as_markup())
    else:
        await callback_or_message.answer(text, reply_markup=kb.as_markup())

@dp.callback_query(F.data.startswith("kgp_select_unit_"))
async def kgp_select_unit(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    index = int(callback.data.replace("kgp_select_unit_", ""))
    units = user_state[uid].get("units_list", [])
    if not (0 <= index < len(units)):
        await safe_answer_callback(callback, "Невірний вибір підрозділу", show_alert=True)
        return
    unit = units[index]
    order = user_state[uid].setdefault("unit_order", [])
    if unit not in order:
        order.append(unit)
    user_state[uid]["current_unit"] = unit
    user_state[uid]["selected_person_indices"] = []
    await show_persons_selection_kgp(callback, uid)
    await safe_answer_callback(callback)

# Dodatky (non-KGP) unit selection handler
@dp.callback_query(F.data.startswith("select_unit_"))
async def select_unit_dodatok(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    index = int(callback.data.replace("select_unit_", ""))
    units = user_state[uid].get("units_list", [])
    if not (0 <= index < len(units)):
        await safe_answer_callback(callback, "Невірний вибір підрозділу", show_alert=True)
        return
    unit = units[index]
    order = user_state[uid].setdefault("unit_order", [])
    if unit not in order:
        order.append(unit)
    user_state[uid]["current_unit"] = unit
    user_state[uid]["selected_person_indices"] = []
    await show_persons_selection_dodatok(callback, uid)
    await safe_answer_callback(callback)

async def show_persons_selection_dodatok(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    unit = user_state[uid].get("current_unit")
    dodatok_type = user_state[uid].get("dodatok", "1")
    unit_key = "Підрозділ" if dodatok_type == "1" else "Підрозділ, що комплектує"
    data = await get_sheet_data_async(dodatok_type)
    persons = [row for row in data if row.get(unit_key, "") == unit]
    user_state[uid]["persons_list"] = persons
    selected = set(user_state[uid].get("selected_person_indices", []))
    kb = InlineKeyboardBuilder()
    for i, person in enumerate(persons):
        if dodatok_type == "1":
            label = (person.get("ПІБ", "") or "").strip() or f"Елемент {i+1}"
        else:
            label = f"{person.get('Марка техніки', '')} {person.get('Номерний знак', '')}".strip() or f"Елемент {i+1}"
        text = f"✅ {label}" if i in selected else label
        kb.button(text=text, callback_data=f"person_{i}")
    kb.button(text="➕ Додати підрозділ", callback_data="add_unit")
    kb.button(text="➡️ Далі", callback_data="next_to_executor")
    kb.button(text="🔙 Назад", callback_data="back_to_units_dodatok")
    kb.button(text="🔄 Рестарт", callback_data="restart")
    kb.adjust(1)
    title = f"👤 Оберіть елементи • {unit}"
    try:
        await callback.message.edit_text(title, reply_markup=kb.as_markup())
    except TelegramBadRequest:
        await callback.message.answer(title, reply_markup=kb.as_markup())

# Toggle person for Dodatky
@dp.callback_query(F.data.startswith("person_"))
async def toggle_person(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    index = int(callback.data.replace("person_", ""))
    sel = user_state[uid].get("selected_person_indices", [])
    if index in sel:
        sel.remove(index)
    else:
        sel.append(index)
    user_state[uid]["selected_person_indices"] = sel
    await show_persons_selection_dodatok(callback, uid)
    await safe_answer_callback(callback)

async def show_persons_selection_kgp(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    unit = user_state[uid].get("current_unit")
    # Read raw sheet values to ensure correct ordering and ПІБ from column B
    persons = []
    # Use warm-cached raw values when available
    cached_vals = SHEET_CACHE.get("1_values")
    if cached_vals:
        values = cached_vals[0]
        data_rows = values[1:] if values else []
        for r in data_rows:
            if len(r) > 5 and r[5].strip() == unit:
                pib_val = (r[1].strip() if len(r) > 1 else "")
                persons.append({"ПІБ": pib_val})
    else:
        try:
            creds = Credentials.from_service_account_file(CREDENTIALS_FILE, scopes=SCOPES)
            client = gspread.authorize(creds)
            sheet = client.open_by_url(GOOGLE_SHEET_URL1).sheet1
            values = sheet.get_all_values() or []
            data_rows = values[1:] if values else []
            for r in data_rows:
                if len(r) > 5 and r[5].strip() == unit:
                    pib_val = (r[1].strip() if len(r) > 1 else "")
                    persons.append({"ПІБ": pib_val})
        except Exception:
            persons = []
    user_state[uid]["persons_list"] = persons
    kb = InlineKeyboardBuilder()
    for i, person in enumerate(persons):
        pib = (person.get('ПІБ', '') or '').strip()
        text = pib if pib else f"Елемент {i+1}"
        kb.button(text=text, callback_data=f"kgp_person_{i}")
    kb.button(text="🔙 Назад", callback_data="back_to_units_dodatok")
    kb.button(text="🔄 Рестарт", callback_data="restart")
    kb.adjust(1)
    title = f"👤 Оберіть особовий склад що залучається• {unit}"
    try:
        await callback.message.edit_text(title, reply_markup=kb.as_markup())
    except TelegramBadRequest:
        await callback.message.answer(title, reply_markup=kb.as_markup())

@dp.callback_query(F.data == "kgp_personal")
async def kgp_personal(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    # Mark KGP as personal and proceed to final confirmation
    user_state.setdefault(uid, {})
    user_state[uid]["kgp_personal"] = True
    user_state[uid]["kgp_text_override"] = "буду здійснювати особисто"
    # Clear any previously selected unit/person choices for KGP if present
    user_state[uid].pop("unit_order", None)
    user_state[uid].pop("selected_indices_by_unit", None)
    await show_kontrol_for_selected_iz_sho(callback, uid)
    await safe_answer_callback(callback, "Обрано: буду здійснювати особисто")

@dp.callback_query(F.data.startswith("kgp_person_"))
async def choose_person_for_kgp(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    index = int(callback.data.replace("kgp_person_", ""))
    unit = user_state[uid].get("current_unit")
    persons = user_state[uid].get("persons_list", [])
    if unit is None or not (0 <= index < len(persons)):
        await safe_answer_callback(callback, "Невірний вибір", show_alert=True)
        return
    user_state[uid].setdefault("selected_indices_by_unit", {})[unit] = [index]
    await show_kontrol_for_selected_iz_sho(callback, uid)
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "add_unit")
async def add_unit_handler(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    current_unit = user_state[uid].get("current_unit")
    selected_indices = user_state[uid].get("selected_person_indices", [])
    if current_unit:
        user_state[uid]["selected_indices_by_unit"][current_unit] = selected_indices[:]
    if user_state[uid]["dodatok"] == "2" and len(selected_indices) > 0:
        user_state[uid]["quantities"] = []
        user_state[uid]["current_qty_index"] = 0
        user_state[uid]["next_action"] = "add_unit"
        await show_qty_input(callback, uid)
    else:
        user_state[uid].pop("selected_person_indices", None)
        user_state[uid].pop("current_unit", None)
        await show_unit_selection(callback, uid)
    await safe_answer_callback(callback)

async def show_qty_input(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    persons = user_state[uid]["persons_list"]
    selected = user_state[uid]["selected_person_indices"]
    current_index = user_state[uid]["current_qty_index"]
    idx = selected[current_index]
    person = persons[idx]
    name = f"{person.get('Марка техніки', '')} {person.get('Номерний знак', '')}"
    text = f"🔢 Оберіть кількість для {name} (1–6)"
    kb = InlineKeyboardBuilder()
    for n in range(1, 7):
        kb.button(text=str(n), callback_data=f"qty_{n}")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(5)
    try:
        await callback.message.edit_text(text, reply_markup=kb.as_markup())
    except TelegramBadRequest as e:
        if "message is not modified" in str(e):
            await safe_answer_callback(callback)
        else:
            raise

@dp.callback_query(F.data.startswith("qty_"))
async def handle_qty(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    num_str = callback.data.split("_")[1]
    num = int(num_str)
    if num < 1 or num > 6:
        await safe_answer_callback(callback, "Невірне значення!")
        return
    user_state[uid]["quantities"].append(num)
    user_state[uid]["current_qty_index"] += 1
    selected = user_state[uid]["selected_person_indices"]
    if user_state[uid]["current_qty_index"] < len(selected):
        await show_qty_input(callback, uid)
    else:
        current_unit = user_state[uid]["current_unit"]
        user_state[uid]["quantities_by_unit"][current_unit] = user_state[uid]["quantities"][:]
        del user_state[uid]["quantities"]
        del user_state[uid]["current_qty_index"]
        next_action = user_state[uid].pop("next_action")
        if next_action == "add_unit":
            user_state[uid].pop("selected_person_indices", None)
            user_state[uid].pop("current_unit", None)
            await show_unit_selection(callback, uid)
        elif next_action == "next_to_executor":
            await select_executor(callback)
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "next_to_executor")
async def next_to_executor(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    current_unit = user_state[uid].get("current_unit")
    selected_indices = user_state[uid].get("selected_person_indices", [])
    if current_unit:
        user_state[uid]["selected_indices_by_unit"][current_unit] = selected_indices[:]
    any_selected = any(len(indices) > 0 for indices in user_state[uid].get("selected_indices_by_unit", {}).values())
    if not any_selected:
        await safe_answer_callback(callback, "Оберіть хоча б один елемент!", show_alert=True)
        return
    if user_state[uid].get("kgp_mode"):
        # Proceed to select 'kontrol' after kgp selection
        await show_kontrol_for_selected_iz_sho(callback, uid)
        return
    if user_state[uid]["dodatok"] == "2" and current_unit and len(selected_indices) > 0:
        user_state[uid]["quantities"] = []
        user_state[uid]["current_qty_index"] = 0
        user_state[uid]["next_action"] = "next_to_executor"
        await show_qty_input(callback, uid)
    else:
        await select_executor(callback)
    await safe_answer_callback(callback)

async def select_executor(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    kb = InlineKeyboardBuilder()
    kb.button(text="Дилін", callback_data="executor_dilin")
    kb.button(text="Царенко", callback_data="executor_tsarenko")
    kb.button(text="Федун", callback_data="executor_fedun")
    kb.button(text="Павленко", callback_data="executor_pavlenko")
    kb.button(text="Назад", callback_data="back_to_persons")
    kb.button(text="Рестарт", callback_data="restart")
    kb.adjust(1)
    try:
        await callback.message.edit_text("🧾 Оберіть виконавця", reply_markup=kb.as_markup())
    except TelegramBadRequest as e:
        if "message is not modified" in str(e):
            await safe_answer_callback(callback)
        else:
            raise

@dp.callback_query(F.data.startswith("executor_"))
async def choose_executor(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    executor = callback.data.split("_")[1]
    user_state[uid]["executor"] = executor
    user_state[uid]["selected_date"] = datetime.now()
    await show_calendar(callback, uid)
    await safe_answer_callback(callback)

async def show_calendar(callback: types.CallbackQuery, uid: int, year=None, month=None):
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    if year is None or month is None:
        current_date = user_state[uid].get("selected_date", datetime.now())
        year = current_date.year
        month = current_date.month
    calendar_markup = create_calendar(year, month)
    executor_name = {
        "dilin": "Дилін",
        "tsarenko": "Царенко",
        "fedun": "Федун",
        "pavlenko": "Павленко",
    }[user_state[uid]["executor"]]
    selected_date = user_state[uid].get("selected_date", datetime.now())
    text = f"📅 Оберіть дату для документу\n" \
           f"Виконавець: {executor_name}\n" \
           f"Обрана дата: {selected_date.strftime('%d.%m.%Y')}"
    try:
        await callback.message.edit_text(text, reply_markup=calendar_markup)
    except TelegramBadRequest as e:
        if "message is not modified" in str(e):
            await safe_answer_callback(callback)
        else:
            raise

@dp.callback_query(F.data.startswith("date_"))
async def select_date(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    old_date = user_state[uid].get("selected_date", datetime.now())
    if callback.data == "date_today":
        new_date = datetime.now()
        year, month = new_date.year, new_date.month
    else:
        parts = callback.data.split("_")
        if len(parts) == 4:
            _, year_str, month_str, day_str = parts
            year, month, day = int(year_str), int(month_str), int(day_str)
            new_date = datetime(year, month, day)
        else:
            new_date = datetime.now()
            year, month = new_date.year, new_date.month
    if new_date.date() == old_date.date():
        await safe_answer_callback(callback)
        return
    user_state[uid]["selected_date"] = new_date
    await show_calendar(callback, uid, year, month)
    await safe_answer_callback(callback)

@dp.callback_query(F.data.startswith("nav_"))
async def navigate_calendar(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    parts = callback.data.split("_")
    if len(parts) == 3:
        _, year_str, month_str = parts
        year, month = int(year_str), int(month_str)
        await show_calendar(callback, uid, year, month)
    else:
        current_date = datetime.now()
        await show_calendar(callback, uid, current_date.year, current_date.month)
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "confirm_date")
async def confirm_date(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    selected_date = user_state[uid].get("selected_date", datetime.now())
    executor_name = {
        "dilin": "Дилін",
        "tsarenko": "Царенко",
        "fedun": "Федун",
        "pavlenko": "Павленко",
    }[user_state[uid]["executor"]]
    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Сформувати додаток", callback_data="generate")
    kb.button(text="🔙 Змінити дату", callback_data="back_to_calendar")
    kb.button(text="Рестарт", callback_data="restart")
    kb.adjust(1)
    text = f"✅ Дата обрана: {selected_date.strftime('%d.%m.%Y')}\n" \
           f"Виконавець: {executor_name}\n\n" \
           f"Натисніть 'Сформувати додаток' для завершення."
    try:
        await callback.message.edit_text(text, reply_markup=kb.as_markup())
    except TelegramBadRequest as e:
        if "message is not modified" in str(e):
            await safe_answer_callback(callback)
        else:
            raise

@dp.callback_query(F.data == "generate")
async def generate(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    if "executor" not in user_state[uid]:
        await safe_answer_callback(callback, "Оберіть виконавця!", show_alert=True)
        return
    dodatok_type = user_state[uid].get("dodatok", "1")
    executor = user_state[uid]["executor"]
    selected_date = user_state[uid].get("selected_date", datetime.now())
    tel_mode = user_state[uid].get("tel_mode")
    if dodatok_type == "1" and tel_mode:
        template = os.path.join(BASE_DIR, "Додаток 3.docx")
    else:
        template = os.path.join(BASE_DIR, "Додаток 1.docx" if dodatok_type == "1" else "Додаток 2.docx")
    
    # Check if template file exists
    if not os.path.exists(template):
        await safe_answer_callback(callback, f"Шаблон не знайдено: {os.path.basename(template)}", show_alert=True)
        return
    loading_msg = await callback.message.answer("Генерація файлів…")
    stop_event = asyncio.Event()
    spinner_task = asyncio.create_task(_spinner_message(loading_msg, "Генерація файлів…", stop_event))
    try:
        filename, pdf_filename = await asyncio.wait_for(
            asyncio.to_thread(
                generate_doc, template, dodatok_type, executor, selected_date, user_state[uid]
            ),
            timeout=120
        )
    finally:
        stop_event.set()
        try:
            await spinner_task
        except Exception:
            pass
    try:
        await loading_msg.edit_text("Надсилання файлів… 📤")
    except Exception:
        pass
    units_str = ", ".join(user_state[uid].get("unit_order", []))
    caption = f"📄 Додаток {dodatok_type} для {units_str}\nДата: {selected_date.strftime('%d.%m.%Y')}"
    sent_doc = await callback.message.answer_document(
        types.FSInputFile(filename),
        caption=caption
    )
    try:
        user_state.setdefault(uid, {}).setdefault('sent_message_ids', []).append((callback.message.chat.id, sent_doc.message_id))
    except Exception:
        pass
    if pdf_filename:
        sent_pdf = await callback.message.answer_document(
            types.FSInputFile(pdf_filename),
            caption=caption
        )
        try:
            user_state.setdefault(uid, {}).setdefault('sent_message_ids', []).append((callback.message.chat.id, sent_pdf.message_id))
        except Exception:
            pass
        safe_remove(pdf_filename)
    else:
        warn = await callback.message.answer("⚠️ PDF не згенеровано. Перевірте шлях до soffice.exe та встановлення LibreOffice.")
        try:
            user_state.setdefault(uid, {}).setdefault('sent_message_ids', []).append((callback.message.chat.id, warn.message_id))
        except Exception:
            pass
    safe_remove(filename)
    if uid not in user_state:
        user_state[uid] = {'authenticated': True, 'state': None}
    else:
        user_state[uid]['authenticated'] = True
        user_state[uid]['state'] = None
    clear_workflow_state(uid)
    try:
        await loading_msg.edit_text("Готово ✅")
        await asyncio.sleep(0.6)
        await loading_msg.delete()
    except Exception:
        pass
    await show_post_generate_menu(callback, uid)
    await safe_answer_callback(callback)

async def _start_nakaz_flow(callback: types.CallbackQuery, is_vchera: bool = False):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state.setdefault(uid, {})["nakaz_flow"] = "new"
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    clear_workflow_state(uid)
    user_state[uid]["mode"] = "nakaz"
    user_state[uid]["dodatok"] = "1"  # Для kgp використовуємо sheet1
    user_state[uid]["addresses_required"] = 1
    user_state[uid]["addresses_list"] = []
    user_state[uid]["is_vchera"] = is_vchera
    
    if is_vchera:
        yesterday = datetime.now() - timedelta(days=1)
        user_state[uid]["1date"] = yesterday
        user_state[uid]["2date"] = yesterday

    await show_calendar_for_nakaz(callback, uid, date_key="1date")
    await safe_answer_callback(callback)

# Нові хендлери для "Наказ на залучення"
@dp.callback_query(F.data == "nakaz")
async def start_nakaz(callback: types.CallbackQuery):
    await _start_nakaz_flow(callback)

@dp.callback_query(F.data == "nakaz_vchera")
async def start_nakaz_vchera(callback: types.CallbackQuery):
    await _start_nakaz_flow(callback, is_vchera=True)

async def show_calendar_for_nakaz(callback: types.CallbackQuery, uid: int, year=None, month=None, date_key="1date"):
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    if year is None or month is None:
        current_date = user_state[uid].get(date_key, datetime.now())
        year = current_date.year
        month = current_date.month
    calendar_markup = create_calendar_for_nakaz(year, month, date_key)
    selected_date = user_state[uid].get(date_key, datetime.now())
    text = f"📅 Оберіть дату коли сталась подія\nОбрана дата: {selected_date.strftime('%d.%m.%Y')}"
    try:
        await callback.message.edit_text(text, reply_markup=calendar_markup)
    except TelegramBadRequest as e:
        if "message is not modified" in str(e):
            await safe_answer_callback(callback)
        else:
            raise

@dp.callback_query(F.data.startswith("nakaz_date_"))
async def select_nakaz_date(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    parts = callback.data.split("_")
    date_key = parts[2]
    if parts[3] == "today":
        new_date = datetime.now()
        year, month = new_date.year, new_date.month
    else:
        year_str, month_str, day_str = parts[3], parts[4], parts[5]
        year, month, day = int(year_str), int(month_str), int(day_str)
        new_date = datetime(year, month, day)
    old_date = user_state[uid].get(date_key, datetime.now())
    if new_date.date() == old_date.date():
        await safe_answer_callback(callback)
        return
    user_state[uid][date_key] = new_date
    # After selecting the second date, go to the next step based on the current state
    if date_key == "2date":
        # Only show unit selection if we haven't already done so
        if not user_state[uid].get("pidrozdily_units"):
            await show_unit_selection_for_pidrozdily(callback, uid)
        else:
            # If units already selected, proceed to selecting 'контроль' before generation
            await show_kontrol_for_selected_iz_sho(callback, uid)
    else:
        await show_calendar_for_nakaz(callback, uid, year, month, date_key)
    await safe_answer_callback(callback)

@dp.callback_query(F.data.startswith("nakaz_nav_"))
async def navigate_nakaz_calendar(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    parts = callback.data.split("_")
    date_key = parts[2]
    year, month = int(parts[3]), int(parts[4])
    await show_calendar_for_nakaz(callback, uid, year, month, date_key)
    await safe_answer_callback(callback)

@dp.callback_query(F.data.startswith("nakaz_confirm_"))
async def confirm_nakaz_date(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not is_authenticated(uid):
        await safe_answer_callback(callback, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        return
    date_key = callback.data.replace("nakaz_confirm_", "")
    if date_key == "1date":
        # Skip 'izchogo' selection and immediately prompt for AI message input
        user_state[uid]["state"] = "sho_ai_seed"
        kb = InlineKeyboardBuilder()
        kb.button(text="⏭️ Пропустити", callback_data="skip_sho_ai_seed")
        kb.button(text="🔙 Назад", callback_data="back_to_start")
        kb.button(text="🔁 Рестарт", callback_data="restart")
        kb.adjust(1)
        text = (
            "Введіть повідомлення: де? що саме?\n"
            "Надішліть одним повідомленням деталі для ШІ."
        )
        try:
            await callback.message.edit_text(text, reply_markup=kb.as_markup())
        except Exception:
            await callback.message.answer(text, reply_markup=kb.as_markup())
    elif date_key == "2date":
        await show_unit_selection_for_pidrozdily(callback, uid)
    await safe_answer_callback(callback, f"✅ Дата {date_key} підтверджена.")

async def show_izchogo_selection(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    cached = SHEET_CACHE.get("pre_izchogo")
    if cached:
        izchogo_list = sorted(cached[0])
    else:
        data = await get_sheet_data_async("1")
        seen = {}
        for row in data:
            iz = row.get("izchogo", "")
            if not iz:
                continue
            key = normalize_text(iz)
            if key and key not in seen:
                seen[key] = iz
        izchogo_list = sorted(seen.values())
    user_state[uid]["izchogo_list"] = izchogo_list
    # Pagination and numeric selection for full-text readability
    per_page = 10
    page = user_state[uid].get("iz_page", 0)
    total = len(izchogo_list)
    start = page * per_page
    end = min(start + per_page, total)
    shown = izchogo_list[start:end]
    lines = []
    for idx, val in enumerate(shown, start=1):
        lines.append(f"{start+idx}. {val}")
    text = "Оберіть 'із чого'\n\n" + ("\n\u2009\n".join(lines) if lines else "Немає значень")
    kb = InlineKeyboardBuilder()
    for idx in range(start, end):
        kb.button(text=str(idx+1), callback_data=f"izchogo_{idx}")
    if page > 0:
        kb.button(text="←", callback_data=f"iz_page_{page-1}")
    if end < total:
        kb.button(text="→", callback_data=f"iz_page_{page+1}")
    kb.button(text="🔙 Назад", callback_data="back_to_start")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(2)
    try:
        await callback.message.edit_text(text, reply_markup=kb.as_markup())
    except Exception:
        await callback.message.answer(text, reply_markup=kb.as_markup())

@dp.callback_query(F.data.startswith("izchogo_"))
async def select_izchogo(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    index = int(callback.data.replace("izchogo_", ""))
    iz_list = user_state[uid]["izchogo_list"]
    iz = iz_list[index]
    user_state[uid]["izchogo"] = iz
    # New flow: immediately prompt for AI details with an option to choose from list
    user_state[uid]["state"] = "sho_ai_seed"
    text = (
        "Введіть повідомлення: де? що саме?\n"
        "Надішліть одним повідомленням деталі для ІІ."
    )
    try:
        await callback.message.edit_text(text)
    except Exception:
        await callback.message.answer(text)
    await safe_answer_callback(callback)

async def show_sho_selection(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    cached = SHEET_CACHE.get("pre_sho")
    if cached:
        sho_list = sorted(cached[0])
    else:
        data = await get_sheet_data_async("1")
        seen_sho = {}
        for row in data:
            sh = row.get("sho", "")
            if not sh:
                continue
            sk = normalize_text(sh)
            if sk and sk not in seen_sho:
                seen_sho[sk] = sh
        sho_list = sorted(seen_sho.values())
    user_state[uid]["sho_list"] = sho_list
    if not sho_list:
        await safe_answer_callback(callback, "Немає варіантів для 'що саме'!", show_alert=True)
        return
    per_page = 10
    page = user_state[uid].get("sho_page", 0)
    total = len(sho_list)
    start = page * per_page
    end = min(start + per_page, total)
    shown = sho_list[start:end]
    lines = []
    for idx, val in enumerate(shown, start=1):
        lines.append(f"{start+idx}. {val}")
    text = "Здійснено ураження чого?\n\n" + ("\n\u2009\n".join(lines) if lines else "Немає значень")
    kb = InlineKeyboardBuilder()
    for idx in range(start, end):
        kb.button(text=str(idx+1), callback_data=f"sho_{idx}")
    if page > 0:
        kb.button(text="←", callback_data=f"sho_page_{page-1}")
    if end < total:
        kb.button(text="→", callback_data=f"sho_page_{page+1}")
    kb.button(text="✏️ Ввести деталі для ІІ", callback_data="sho_ai_prompt")
    kb.button(text="🧠 Варіант ІІ", callback_data="sho_ai_generate")
    kb.button(text="🔙 Назад", callback_data="back_to_izchogo")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(2)
    try:
        await callback.message.edit_text(text, reply_markup=kb.as_markup())
    except Exception:
        await callback.message.answer(text, reply_markup=kb.as_markup())

@dp.callback_query(F.data.regexp(r"^sho_\d+$"))
async def select_sho(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    suffix = callback.data.replace("sho_", "", 1)
    if not suffix.isdigit():
        await safe_answer_callback(callback)
        return
    index = int(suffix)
    sho_list = user_state[uid]["sho_list"]
    sho = sho_list[index]
    user_state[uid]["sho"] = sho
    # After selecting 'sho', proceed to selecting the second date
    await show_calendar_for_nakaz(callback, uid, date_key="2date")
    await safe_answer_callback(callback)

def _build_kontrol_list_for_selected_iz_sho():
    # Helper to build kontrol list from current izchogo + sho
    # Returns sorted list
    return []

async def show_kontrol_for_selected_iz_sho(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    cached = SHEET_CACHE.get("pre_kontrol")
    if cached:
        kontrol_list = sorted(cached[0])
    else:
        data = await get_sheet_data_async("1")
        kontrol_seen = {}
        for row in data:
            ko = row.get("kontrol", "")
            if not ko:
                continue
            kk = normalize_text(ko)
            if kk and kk not in kontrol_seen:
                kontrol_seen[kk] = ko
        kontrol_list = sorted(kontrol_seen.values())
    user_state[uid]["kontrol_list"] = kontrol_list
    if not kontrol_list:
        user_state[uid]["kontрол"] = ""
        await show_pidpys_selection(callback, uid)
        return
    if len(kontrol_list) == 1:
        user_state[uid]["kontрол"] = kontrol_list[0]
        await show_pidpys_selection(callback, uid)
        return
    per_page = 10
    page = user_state[uid].get("kontrol_page", 0)
    total = len(kontrol_list)
    start = page * per_page
    end = min(start + per_page, total)
    shown = kontrol_list[start:end]
    lines = []
    for idx, val in enumerate(shown, start=1):
        lines.append(f"{start+idx}. {val}")
    text = "Оберіть 'контроль'\n\n" + ("\n\u2009\n".join(lines) if lines else "Немає значень")
    kb = InlineKeyboardBuilder()
    for idx in range(start, end):
        kb.button(text=str(idx+1), callback_data=f"kontrol_{idx}")
    if page > 0:
        kb.button(text="←", callback_data=f"kontrol_page_{page-1}")
    if end < total:
        kb.button(text="→", callback_data=f"kontrol_page_{page+1}")
    kb.button(text="⏭️ Пропустити", callback_data="skip_kontrol")
    kb.button(text="🔙 Назад", callback_data="back_to_start")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(2)
    try:
        await callback.message.edit_text(text, reply_markup=kb.as_markup())
    except Exception:
        await callback.message.answer(text, reply_markup=kb.as_markup())

async def show_pidpys_selection(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    kb = InlineKeyboardBuilder()
    kb.button(text="Коржов", callback_data="pidpys_korzhov")
    kb.button(text="Усик", callback_data="pidpys_usyk")
    kb.button(text="Середа", callback_data="pidpys_sereda")
    kb.button(text="⏭️ Пропустити", callback_data="skip_pidpys")
    kb.button(text="🔙 Назад", callback_data="back_to_kontrol")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    text = "Хто підписує наказ?"
    try:
        await callback.message.edit_text(text, reply_markup=kb.as_markup())
    except Exception:
        await callback.message.answer(text, reply_markup=kb.as_markup())

@dp.callback_query(F.data.startswith("pidpys_"))
async def select_pidpys(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    
    try:
        pidpys_key = callback.data.replace("pidpys_", "")
        if pidpys_key in PIDPYS:
            user_state[uid]["pidpys"] = pidpys_key
            await safe_answer_callback(callback, "Підпис обрано")
            await prompt_generate_nakaz(callback, uid)
        else:
            await safe_answer_callback(callback, "Невірний вибір", show_alert=True)
    except Exception as e:
        print(f"Помилка в select_pidpys: {e}")
        await safe_answer_callback(callback, "Сталася помилка", show_alert=True)

@dp.callback_query(F.data == "back_to_kontrol")
async def back_to_kontrol(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_kontrol_for_selected_iz_sho(callback, uid)
    await safe_answer_callback(callback)

async def prompt_generate_nakaz(callback: types.CallbackQuery, uid: int):
    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Сформувати наказ", callback_data="generate_nakaz")
    kb.button(text="🔙 Назад", callback_data="back_to_pidpys")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    try:
        await callback.message.edit_text("Перевірте вибір і сформуйте наказ:", reply_markup=kb.as_markup())
    except Exception:
        await callback.message.answer("Перевірте вибір і сформуйте наказ:", reply_markup=kb.as_markup())

@dp.callback_query(F.data == "back_to_pidpys")
async def back_to_pidpys(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_pidpys_selection(callback, uid)
    await safe_answer_callback(callback)

# Хендлери для пропуску кроків
@dp.callback_query(F.data.startswith("skip_date_"))
async def skip_date(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    date_key = callback.data.replace("skip_date_", "")
    # Встановлюємо поточну дату як значення за замовчуванням
    user_state[uid][date_key] = datetime.now()
    await safe_answer_callback(callback, "Дату пропущено")
    
    if date_key == "1date":
        # Переходимо до введення деталей для ШІ
        user_state[uid]["state"] = "sho_ai_seed"
        kb = InlineKeyboardBuilder()
        kb.button(text="⏭️ Пропустити", callback_data="skip_sho_ai_seed")
        kb.button(text="🔙 Назад", callback_data="back_to_start")
        kb.button(text="🔁 Рестарт", callback_data="restart")
        kb.adjust(1)
        text = (
            "Введіть повідомлення: де? що саме?\n"
            "Надішліть одним повідомленням деталі для ШІ."
        )
        try:
            await callback.message.edit_text(text, reply_markup=kb.as_markup())
        except Exception:
            await callback.message.answer(text, reply_markup=kb.as_markup())
    elif date_key == "2date":
        # Переходимо до вибору підрозділів
        await show_unit_selection_for_pidrozdily(callback, uid)

@dp.callback_query(F.data == "skip_sho_ai_seed")
async def skip_sho_ai_seed(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state[uid]["state"] = None
    user_state[uid]["sho"] = ""  # Порожнє значення
    await safe_answer_callback(callback, "Крок пропущено")
    # Переходимо до вибору другої дати
    await show_calendar_for_nakaz(callback, uid, date_key="2date")

@dp.callback_query(F.data == "skip_pidrozdily")
async def skip_pidrozdily(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    # Встановлюємо порожній список підрозділів
    user_state[uid]["pidrozdily_units"] = []
    await safe_answer_callback(callback, "Вибір підрозділів пропущено")
    # Переходимо до вибору начальників (або пропускаємо і їх)
    await show_nachalnyky_selection(callback, uid)

@dp.callback_query(F.data == "skip_nachalnyky")
async def skip_nachalnyky(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    # Встановлюємо порожній список начальників
    user_state[uid]["nachalnyky_selected_indices"] = []
    await safe_answer_callback(callback, "Вибір начальників пропущено")
    # Переходимо до введення адреси
    sho_text = user_state.get(uid, {}).get("sho", "").strip()
    if not sho_text:
        user_state[uid]["state"] = "enter_adresa"
        kb = InlineKeyboardBuilder()
        kb.button(text="⏭️ Пропустити", callback_data="skip_adresa")
        kb.button(text="🔙 Назад", callback_data="back_to_pid_units")
        kb.button(text="🔁 Рестарт", callback_data="restart")
        kb.adjust(1)
        await callback.message.answer("Введіть адресу:", reply_markup=kb.as_markup())
    else:
        await next_to_adresa(callback)

@dp.callback_query(F.data == "skip_adresa")
async def skip_adresa(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state[uid]["adresa"] = ""
    user_state[uid]["state"] = None
    cleanup_address_request(uid)
    await safe_answer_callback(callback, "Адресу пропущено")
    user_state[uid]["kgp_mode"] = True
    user_state[uid]["dodatok"] = "1"
    user_state[uid]["unit_order"] = []
    user_state[uid]["selected_indices_by_unit"] = {}
    user_state[uid]["quantities_by_unit"] = {}
    await show_unit_selection(callback, uid)

@dp.callback_query(F.data == "skip_kontrol")
async def skip_kontrol(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state[uid]["kontrol"] = ""  # Порожнє значення
    await safe_answer_callback(callback, "Контроль пропущено")
    # Переходимо до вибору підпису
    await show_pidpys_selection(callback, uid)

@dp.callback_query(F.data == "skip_pidpys")
async def skip_pidpys(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state[uid]["pidpys"] = None  # Не вибрано підпис
    await safe_answer_callback(callback, "Підпис пропущено")
    # Переходимо до фінального екрану
    await prompt_generate_nakaz(callback, uid)

async def show_unit_selection_for_pidrozdily(callback_or_message, uid: int):
    touch_activity(uid)
    if not is_authenticated(uid):
        if isinstance(callback_or_message, types.CallbackQuery):
            await safe_answer_callback(callback_or_message, "Ви не автентифіковані! Використовуйте /start і введіть пароль.", show_alert=True)
        else:
            await callback_or_message.answer("Ви не автентифіковані! Використовуйте /start і введіть пароль.")
        return
    
    dodatok_type = "1"  # Для pidrozdily та kgp використовуємо sheet1
    # Use precomputed units for instant response
    cached_units = SHEET_CACHE.get("pre_units1")
    if cached_units:
        units = cached_units[0]
        # Ensure we also have full data for later steps
        cached_data_entry = SHEET_CACHE.get("1")
        data = cached_data_entry[0] if cached_data_entry else get_sheet_data(dodatok_type)
    else:
        data = get_sheet_data(dodatok_type)
        unit_key = "Підрозділ"
        units = sorted(set(row.get(unit_key, "") for row in data if row.get(unit_key)))
    user_state[uid]["units_list"] = units
    
    # Store the full data for later use in KGP selection
    user_state[uid]["kgp_full_data"] = data
    
    selected_units = user_state[uid].get("pidrozdily_units", [])
    
    kb = InlineKeyboardBuilder()
    for i, unit in enumerate(units):
        button_text = f"✅ {unit}" if unit in selected_units else unit
        callback_data = f"select_pid_unit_{i}"
        kb.button(text=button_text, callback_data=callback_data)
    
    kb.button(text="➡️ Далі", callback_data="next_to_nachalnyky")
    kb.button(text="⏭️ Пропустити", callback_data="skip_pidrozdily")
    kb.button(text="🔙 Назад", callback_data="back_to_start")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    
    text = "Оберіть підрозділи які залучаються"
    
    if isinstance(callback_or_message, types.CallbackQuery):
        try:
            await callback_or_message.message.edit_text(text, reply_markup=kb.as_markup())
        except TelegramBadRequest:
            await callback_or_message.message.answer(text, reply_markup=kb.as_markup())
    elif isinstance(callback_or_message, types.Message):
        await callback_or_message.answer(text, reply_markup=kb.as_markup())

@dp.callback_query(F.data == "next_to_nachalnyky")
async def next_to_nachalnyky(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    # Ensure units are selected
    if not user_state.get(uid, {}).get("pidrozdily_units"):
        await safe_answer_callback(callback, "Оберіть хоча б один підрозділ!", show_alert=True)
        return
    
    # Skip supervisors selection for yesterday flow
    if user_state[uid].get("is_vchera"):
        await next_to_adresa(callback)
        return

    # Build and show nachalnyky list filtered by selected units and allowed categories (column D)
    await show_nachalnyky_selection(callback, uid)
    await safe_answer_callback(callback)

async def show_nachalnyky_selection(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    selected_units = user_state[uid].get("pidrozdily_units", [])
    # Use warm-cached raw values of sheet 1 for speed; fallback to live fetch only if missing
    cached_vals = SHEET_CACHE.get("1_values")
    if cached_vals:
        values = cached_vals[0]
    else:
        try:
            creds = Credentials.from_service_account_file(CREDENTIALS_FILE, scopes=SCOPES)
            client = gspread.authorize(creds)
            sheet = client.open_by_url(GOOGLE_SHEET_URL1).sheet1
            values = sheet.get_all_values() or []
        except Exception:
            values = []
    data_rows = values[1:] if values else []
    allowed_categories = {
        "начальник поста",
        "начальник частини",
        "заступник начальника частини",
    }
    items = []
    # Build items with unit filter (F, index 5) and category filter (D, index 3)
    for r in data_rows:
        unit_val = (r[5].strip() if len(r) > 5 else "")
        cat_val = (r[3].strip().lower() if len(r) > 3 else "")
        if unit_val in selected_units and cat_val in allowed_categories:
            col_o = (r[14].strip() if len(r) > 14 else "")
            col_p = (r[15].strip() if len(r) > 15 else "")
            col_q = (r[16].strip() if len(r) > 16 else "")
            display = ", ".join([p for p in [f"{col_o} {col_p}".strip(), col_q] if p])
            # Fallback: if display empty, skip
            if not display:
                continue
            # Store full text for document and P for menu display
            items.append({"unit": unit_val, "text": display, "p": col_p})
    user_state[uid]["nachalnyky_items"] = items
    await render_nachalnyky_menu(callback, uid)

async def render_nachalnyky_menu(callback: types.CallbackQuery, uid: int):
    selected_idx = set(user_state[uid].get("nachalnyky_selected_indices", []))
    items = user_state.get(uid, {}).get("nachalnyky_items", [])
    kb = InlineKeyboardBuilder()
    # Show grouped by unit for clarity; preserve original order
    for i, it in enumerate(items):
        p_val = it.get('p', '')
        label = p_val if p_val else it.get('text', '')
        txt = f"✅ {label}" if i in selected_idx else label
        kb.button(text=txt, callback_data=f"nach_{i}")
    kb.button(text="➡️ Далі", callback_data="next_to_adresa")
    kb.button(text="⏭️ Пропустити", callback_data="skip_nachalnyky")
    kb.button(text="🔙 Назад", callback_data="back_to_pid_units")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    title = "Оберіть начальників або заступників хто проводить інструктажі перед виконанням завдань"
    try:
        await callback.message.edit_text(title, reply_markup=kb.as_markup())
    except TelegramBadRequest:
        await callback.message.answer(title, reply_markup=kb.as_markup())

@dp.callback_query(F.data.startswith("nach_"))
async def toggle_nachalnyk(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    try:
        idx = int(callback.data.replace("nach_", ""))
    except Exception:
        await safe_answer_callback(callback)
        return
    items = user_state.get(uid, {}).get("nachalnyky_items", [])
    if not (0 <= idx < len(items)):
        await safe_answer_callback(callback)
        return
    sel = set(user_state[uid].get("nachalnyky_selected_indices", []))
    if idx in sel:
        sel.remove(idx)
    else:
        sel.add(idx)
    user_state[uid]["nachalnyky_selected_indices"] = sorted(sel)
    # Re-render menu from cached items without recomputing from sheets
    await render_nachalnyky_menu(callback, uid)
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "back_to_pid_units")
async def back_to_pid_units(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_unit_selection_for_pidrozdily(callback, uid)
    await safe_answer_callback(callback)

@dp.callback_query(F.data.startswith("select_pid_unit_"))
async def toggle_pid_unit(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    index = int(callback.data.replace("select_pid_unit_", ""))
    units = user_state[uid]["units_list"]
    unit = units[index]
    pidrozdily_units = user_state[uid].get("pidrozdily_units", [])
    if unit in pidrozdily_units:
        pidrozdily_units.remove(unit)
    else:
        pidrozdily_units.append(unit)
    user_state[uid]["pidrozdily_units"] = pidrozdily_units
    await show_unit_selection_for_pidrozdily(callback, uid)
    await safe_answer_callback(callback)

def get_addresses_from_sheet():
    """Fetch unique addresses from column K of the Google Sheet with caching"""
    # Check cache first
    cached = SHEET_CACHE.get("addresses")
    if cached:
        data, ts = cached
        if time.time() - ts < CACHE_TTL_SECONDS:
            return data
    
    # Fetch from sheet if not cached or expired
    try:
        creds = Credentials.from_service_account_file(CREDENTIALS_FILE, scopes=SCOPES)
        client = gspread.authorize(creds)
        sheet = client.open_by_url(GOOGLE_SHEET_URL1).sheet1
        # Column K is index 10 (0-based)
        addresses = sheet.col_values(11)[1:]  # Skip header row
        # Filter out empty strings and get unique values
        unique_addresses = sorted(list(set(addr.strip() for addr in addresses if addr.strip())))
        # Cache the result
        SHEET_CACHE["addresses"] = (unique_addresses, time.time())
        return unique_addresses
    except Exception as e:
        print(f"Error fetching addresses from sheet: {e}")
        return []

def get_address_hash(address: str) -> str:
    """Generate a short hash for the address to use in callback data"""
    import hashlib
    return hashlib.md5(address.encode('utf-8')).hexdigest()[:10]

@dp.callback_query(F.data == "next_to_adresa")
async def next_to_adresa(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if not user_state[uid].get("pidrozdily_units"):
        await safe_answer_callback(callback, "Оберіть хоча б один підрозділ!", show_alert=True)
        return
    
    # Use last AI SHO text to extract address via AI
    sho_text = user_state.get(uid, {}).get("sho", "").strip()
    
    # If sho is empty (was skipped), go directly to manual address input
    if not sho_text:
        await _prompt_manual_address_input(callback, uid)
        await safe_answer_callback(callback)
        return
    
    # Сохраняем запрос во временный файл
    save_address_request(uid, sho_text, {
        'pidrozdily_units': user_state[uid].get("pidrozdily_units", []),
        'nachalnyky_selected_indices': user_state[uid].get("nachalnyky_selected_indices", [])
    })
    
    loading = await callback.message.answer("Отримую адресу з ШІ…")
    try:
        address_text = await asyncio.to_thread(_call_gemini_extract_address_improved, sho_text, uid)
    except Exception:
        address_text = ""
    try:
        await loading.delete()
    except Exception:
        pass

    if address_text and "за адресою:" in address_text.lower():
        await _append_address_and_continue(callback, uid, address_text)
    else:
        await _reply_message(callback, "❌ Не вдалося автоматично визначити адресу. Будь ласка, введіть адресу вручну:")
        await _prompt_manual_address_input(callback, uid)

    await safe_answer_callback(callback)
    
@dp.callback_query(F.data == "refresh_cache")
async def refresh_cache(callback: types.CallbackQuery):
    """Ручне оновлення кешу Google Sheets з меню бота."""
    uid = callback.from_user.id
    touch_activity(uid)
    loading = await callback.message.answer("♻️ Оновлюю кеш з Google Sheets…")
    try:
        await warm_cache_loop()
        status_msg = "✅ Кеш успішно оновлено."
    except Exception:
        status_msg = "⚠️ Не вдалося оновити кеш. Перевірте підключення до інтернету."
    try:
        await loading.edit_text(status_msg)
    except Exception:
        await callback.message.answer(status_msg)
    await safe_answer_callback(callback)

@dp.callback_query(F.data.startswith("addr_page_"))
async def addr_page_nav(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    try:
        page = int(callback.data.replace("addr_page_", ""))
    except Exception:
        page = 0
    if page < 0:
        page = 0
    user_state[uid]["addr_page"] = page
    # Re-render the address selection view
    await next_to_adresa(callback)
    await safe_answer_callback(callback)

@dp.callback_query(F.data.startswith("addr:"))
async def select_address(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    
    # Get the address hash from callback data
    addr_hash = callback.data.split(":", 1)[1]
    
    # Look up the full address from our mapping
    address_mapping = user_state[uid].get("_address_mapping", {})
    if addr_hash not in address_mapping:
        await safe_answer_callback(callback, "Помилка: адресу не знайдено. Спробуйте ще раз.", show_alert=True)
        return
        
    address = address_mapping[addr_hash]
    if "_address_mapping" in user_state[uid]:
        del user_state[uid]["_address_mapping"]
    await safe_answer_callback(callback)
    await _append_address_and_continue(callback, uid, address)

@dp.callback_query(F.data == "enter_address_manually")
async def enter_address_manually(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state[uid]["state"] = "enter_adresa"
    await callback.message.answer("Введіть адресу для {adresa}:")
    await safe_answer_callback(callback)

@dp.callback_query(F.data == "back_to_units")
async def back_to_units(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_unit_selection_for_pidrozdily(callback, uid)
    await safe_answer_callback(callback)

# Для nakaz generate викликаємо після вибору kgp
@dp.callback_query(F.data == "generate_nakaz")
async def generate_nakaz_handler(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    
    # Негайно відповідаємо, що почали обробку
    await safe_answer_callback(callback, "Початок генерації наказу...")
    
    loading_msg = await callback.message.answer("Генерація файлів…")
    stop_event = asyncio.Event()
    spinner_task = asyncio.create_task(_spinner_message(loading_msg, "Генерація файлів…", stop_event))
    try:
        # Use only the first set of selected units (pidrozdily_units)
        if "pidrozdily_units2" in user_state[uid]:
            del user_state[uid]["pidrozdily_units2"]
        filename, pdf_filename = await asyncio.to_thread(generate_nakaz, user_state[uid])
        
        # Зупиняємо спінер перед відправкою файлів
        stop_event.set()
        try:
            await spinner_task
        except Exception:
            pass
            
        await loading_msg.edit_text("Надсилання файлів… 📤")
        
        caption = f"📜 Наказ на залучення"
        sent_doc = await callback.message.answer_document(
            types.FSInputFile(filename),
            caption=caption
        )
        safe_remove(filename)
        
        if pdf_filename:
            sent_pdf = await callback.message.answer_document(
                types.FSInputFile(pdf_filename),
                caption=caption
            )
            safe_remove(pdf_filename)
            
        clear_workflow_state(uid)
        await loading_msg.edit_text("Готово ✅")
        await asyncio.sleep(0.6)
        await loading_msg.delete()
        
    except Exception as e:
        stop_event.set()
        try:
            await spinner_task
        except Exception:
            pass
        await loading_msg.edit_text(f"❌ Помилка генерації: {e}")
        
    await show_post_generate_menu(callback, uid)

# ======================================================================== #
# ===  НАКАЗ ВІДПУСТКИ / ЛІКАРНЯНІ  ====================================== #
# ======================================================================== #

# --- Helper for Message/CallbackQuery compatibility --- #

async def _vidp_reply(target, text: str, reply_markup=None):
    """Безпечно відповісти/редагувати незалежно від типу target."""
    try:
        if isinstance(target, types.CallbackQuery):
            await target.message.edit_text(text, reply_markup=reply_markup)
        else:
            await target.answer(text, reply_markup=reply_markup)
    except TelegramBadRequest as e:
        err = str(e)
        if "message is not modified" in err:
            pass  # ignore
        elif "message can't be edited" in err:
            if isinstance(target, types.CallbackQuery):
                await target.message.answer(text, reply_markup=reply_markup)
            else:
                await target.answer(text, reply_markup=reply_markup)
        else:
            import traceback
            print(f"⚠️ _vidp_reply error: {e}")
            traceback.print_exc()

# --- Calendar for vidpustky dates --- #

def create_vidp_calendar(year=None, month=None, date_key="vidp_date"):
    """Календар для вибору дати наказу відпустки."""
    if year is None:
        year = datetime.now().year
    if month is None:
        month = datetime.now().month
    kb = InlineKeyboardBuilder()
    month_name = uk_month_nom[month - 1]
    kb.button(text=f"{month_name} {year}", callback_data="ignore")
    days = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Нд"]
    for day in days:
        kb.button(text=day, callback_data="ignore")
    month_calendar = calendar.monthcalendar(year, month)
    today = datetime.now().date()
    for week in month_calendar:
        for day in week:
            if day == 0:
                kb.button(text=" ", callback_data="ignore")
            else:
                btn_text = str(day)
                if today == datetime(year, month, day).date():
                    btn_text = f"📍 {btn_text}"
                kb.button(text=btn_text, callback_data=f"{date_key}_{year}_{month}_{day}")
    prev_month = month - 1 if month > 1 else 12
    prev_year = year if month > 1 else year - 1
    next_month = month + 1 if month < 12 else 1
    next_year = year if month < 12 else year + 1
    kb.button(text="← Попередній", callback_data=f"vidp_nav_{prev_year}_{prev_month}")
    kb.button(text="Сьогодні", callback_data=f"{date_key}_today")
    kb.button(text="Наступний →", callback_data=f"vidp_nav_{next_year}_{next_month}")
    kb.button(text="✅ Підтвердити дату", callback_data="vidp_confirm_date")
    kb.button(text="🔙 Назад", callback_data="back_to_start")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1, 7, *[7 for _ in range(len(month_calendar))], 3, 3)
    return kb.as_markup()


def create_vidp_start_calendar(year=None, month=None):
    """Календар для вибору дати початку відпустки (ручна збірка)."""
    if year is None:
        year = datetime.now().year
    if month is None:
        month = datetime.now().month

    month_name = uk_month_nom[month - 1]
    days = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Нд"]
    month_cal = calendar.monthcalendar(year, month)
    today = datetime.now().date()

    keyboard = []
    # Row 1: month name
    keyboard.append([InlineKeyboardButton(text=f"{month_name} {year}", callback_data="ignore")])
    # Row 2: day names
    keyboard.append([InlineKeyboardButton(text=d, callback_data="ignore") for d in days])
    # Calendar rows
    for week in month_cal:
        row = []
        for day in week:
            if day == 0:
                row.append(InlineKeyboardButton(text=" ", callback_data="ignore"))
            else:
                txt = f"[{day}]" if today == datetime(year, month, day).date() else str(day)
                row.append(InlineKeyboardButton(text=txt, callback_data=f"vsd_{year}_{month}_{day}"))
        keyboard.append(row)

    prev_month = month - 1 if month > 1 else 12
    prev_year = year if month > 1 else year - 1
    next_month = month + 1 if month < 12 else 1
    next_year = year if month < 12 else year + 1

    # Nav row
    keyboard.append([
        InlineKeyboardButton(text="< Попередній", callback_data=f"vsn_{prev_year}_{prev_month}"),
        InlineKeyboardButton(text="Сьогодні", callback_data="vsd_today"),
        InlineKeyboardButton(text="Наступний >", callback_data=f"vsn_{next_year}_{next_month}"),
    ])
    # Bottom row
    keyboard.append([
        InlineKeyboardButton(text=".. Назад", callback_data="vidp_back_to_days"),
        InlineKeyboardButton(text=".. Рестарт", callback_data="restart"),
    ])

    return InlineKeyboardMarkup(inline_keyboard=keyboard)


# --- Entry point --- #

@dp.callback_query(F.data == "nakaz_vidpustky")
async def vidp_start(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    clear_workflow_state(uid)
    user_state[uid]["mode"] = "vidpustky"
    user_state[uid]["vidpustky_items"] = []
    user_state[uid]["vidpustky_sick_items"] = []
    user_state[uid]["vidpustky_return_items"] = []
    user_state[uid]["vidp_data_nakazu"] = datetime.now()
    await show_vidp_calendar(callback, uid)
    await safe_answer_callback(callback)


async def show_vidp_calendar(callback: types.CallbackQuery, uid: int, year=None, month=None):
    touch_activity(uid)
    if year is None or month is None:
        d = user_state[uid].get("vidp_data_nakazu", datetime.now())
        year, month = d.year, d.month
    markup = create_vidp_calendar(year, month)
    sel_date = user_state[uid].get("vidp_data_nakazu", datetime.now())
    text = f"📅 Оберіть дату наказу\nОбрана дата: {sel_date.strftime('%d.%m.%Y')}"
    try:
        await callback.message.edit_text(text, reply_markup=markup)
    except TelegramBadRequest:
        await callback.message.answer(text, reply_markup=markup)


@dp.callback_query(F.data.startswith("vidp_date_"))
async def vidp_select_date(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    parts = callback.data.split("_")
    if parts[2] == "today":
        new_date = datetime.now()
        year, month = new_date.year, new_date.month
    else:
        year, month, day = int(parts[2]), int(parts[3]), int(parts[4])
        new_date = datetime(year, month, day)
    user_state[uid]["vidp_data_nakazu"] = new_date
    await show_vidp_calendar(callback, uid, year, month)
    await safe_answer_callback(callback)


@dp.callback_query(F.data.startswith("vidp_nav_"))
async def vidp_navigate_calendar(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    parts = callback.data.split("_")
    year, month = int(parts[2]), int(parts[3])
    await show_vidp_calendar(callback, uid, year, month)
    await safe_answer_callback(callback)


@dp.callback_query(F.data == "vidp_confirm_date")
async def vidp_confirm_date(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_vidp_main_menu(callback, uid)
    await safe_answer_callback(callback, "✅ Дату підтверджено")


# --- Main menu --- #

async def show_vidp_main_menu(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    n_vidp = len(user_state[uid].get("vidpustky_items", []))
    n_sick = len(user_state[uid].get("vidpustky_sick_items", []))
    n_ret = len(user_state[uid].get("vidpustky_return_items", []))

    stats = f"📋 Відпустки: {n_vidp} | Лікарняні: {n_sick} | Повернення: {n_ret}"

    kb = InlineKeyboardBuilder()
    kb.button(text="➕ Додати особу (відпустка)", callback_data="vidp_add_person")
    kb.button(text="➕ Додати особу (лікарняний)", callback_data="vidp_add_sick")
    kb.button(text="➕ Повернення після лікарняного", callback_data="vidp_add_return")
    if n_vidp > 0 or n_sick > 0 or n_ret > 0:
        kb.button(text="📋 Переглянути список", callback_data="vidp_show_list")
        kb.button(text="➡️ Далі (вибір підписанта)", callback_data="vidp_next_to_sign")
    kb.button(text="🔙 Назад", callback_data="back_to_start")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)

    sel_date = user_state[uid].get("vidp_data_nakazu", datetime.now())
    text = f"📜 Наказ відпустки\n📅 Дата наказу: {sel_date.strftime('%d.%m.%Y')}\n\n{stats}\n\nОберіть дію:"

    try:
        await callback.message.edit_text(text, reply_markup=kb.as_markup())
    except Exception:
        await callback.message.answer(text, reply_markup=kb.as_markup())


# --- Add person (vidpustka) flow --- #

@dp.callback_query(F.data == "vidp_add_person")
async def vidp_add_person(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state[uid]["vidp_mode"] = "vidpustka"
    await show_vidp_unit_selection(callback, uid)
    await safe_answer_callback(callback)


async def show_vidp_unit_selection(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    cached_units = SHEET_CACHE.get("pre_units1")
    if cached_units:
        units = cached_units[0]
    else:
        data = await get_sheet_data_async("1")
        units = sort_units(list({row.get("Підрозділ", "") for row in data if row.get("Підрозділ")}))
    user_state[uid]["vidp_units_list"] = units

    kb = InlineKeyboardBuilder()
    for i, unit in enumerate(units):
        kb.button(text=unit, callback_data=f"vidp_select_unit_{i}")
    kb.button(text="🔙 Назад", callback_data="vidp_back_to_menu")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    await callback.message.edit_text("🏢 Оберіть підрозділ:", reply_markup=kb.as_markup())


@dp.callback_query(F.data.startswith("vidp_select_unit_"))
async def vidp_select_unit(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    idx = int(callback.data.replace("vidp_select_unit_", ""))
    units = user_state[uid].get("vidp_units_list", [])
    if not (0 <= idx < len(units)):
        await safe_answer_callback(callback)
        return
    unit = units[idx]
    user_state[uid]["vidp_current_unit"] = unit
    await show_vidp_person_selection(callback, uid)
    await safe_answer_callback(callback)


async def show_vidp_person_selection(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    unit = user_state[uid].get("vidp_current_unit", "")
    # Use cached raw values
    cached_vals = SHEET_CACHE.get("1_values")
    persons = []
    if cached_vals:
        values = cached_vals[0]
        data_rows = values[1:] if values else []
        for r in data_rows:
            if len(r) > 5 and r[5].strip() == unit:
                pib = r[1].strip() if len(r) > 1 else ""  # Col B: ПІБ (називний)
                pib_dative = r[20].strip() if len(r) > 20 else ""  # Col U: ПІБ (давальний)
                pib_genitive = r[21].strip() if len(r) > 21 else ""  # Col V: ПІБ (родовий)
                rank_dative = r[14].strip() if len(r) > 14 else ""  # Col O: звання (давальний)
                rank_nom = r[2].strip() if len(r) > 2 else ""  # Col C: звання (називний)
                rank_genitive = r[22].strip() if len(r) > 22 else ""  # Col W: звання (родовий)
                position_full = r[16].strip() if len(r) > 16 else ""  # Col Q: повна посада (дав.)
                position_genitive = r[9].strip() if len(r) > 9 else ""  # Col J: посада (родовий)
                has_rank = bool(r[4].strip()) if len(r) > 4 else False  # Col E: ознака звання
                if pib:
                    persons.append({
                        "ПІБ": pib,
                        "pib_dative": pib_dative or pib,       # §1
                        "pib_genitive": pib_genitive or pib,    # §2, §3
                        "rank_dative": rank_dative,             # §1
                        "rank_genitive": rank_genitive or rank_nom,  # §2, §3
                        "rank": rank_nom or rank_dative,        # fallback
                        "position_full": position_full,         # §1 (дав.)
                        "position_genitive": position_genitive or position_full,  # §2, §3 (род.)
                        "position": position_full,              # для сумісності
                        "unit": unit,
                        "has_rank": has_rank,                   # колонка E не порожня
                    })
    user_state[uid]["vidp_persons_list"] = persons

    if not persons:
        kb = InlineKeyboardBuilder()
        kb.button(text="🔙 Назад", callback_data="vidp_back_to_units")
        await callback.message.edit_text(
            f"⚠️ Немає даних для підрозділу '{unit}'.\nСпробуйте інший підрозділ.",
            reply_markup=kb.as_markup())
        return

    kb = InlineKeyboardBuilder()
    for i, person in enumerate(persons):
        pib = person.get("ПІБ", "")
        kb.button(text=pib, callback_data=f"vidp_person_{i}")
    kb.button(text="🔙 Назад", callback_data="vidp_back_to_units")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    await callback.message.edit_text(f"👤 Оберіть особу • {unit}", reply_markup=kb.as_markup())


@dp.callback_query(F.data.startswith("vidp_person_"))
async def vidp_select_person(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    idx = int(callback.data.replace("vidp_person_", ""))
    persons = user_state[uid].get("vidp_persons_list", [])
    if not (0 <= idx < len(persons)):
        await safe_answer_callback(callback)
        return
    person = persons[idx]
    user_state[uid]["vidp_current_person"] = person

    mode = user_state[uid].get("vidp_mode", "vidpustka")
    if mode == "vidpustka":
        await show_vidp_leave_type(callback, uid)
    elif mode == "sick":
        await show_vidp_sick_date(callback, uid)
    elif mode == "return":
        await show_vidp_return_date(callback, uid)
    await safe_answer_callback(callback)


# --- Leave type selection --- #

async def show_vidp_leave_type(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    person = user_state[uid].get("vidp_current_person", {})
    pib = person.get("ПІБ", "")

    kb = InlineKeyboardBuilder()
    kb.button(text="Щорічна основна (частина)", callback_data="vidp_type_shchorichna")
    kb.button(text="Щорічна основна (повна)", callback_data="vidp_type_shchorichna_full")
    kb.button(text="Додаткова", callback_data="vidp_type_dodatkova")
    kb.button(text="За сімейними обставинами", callback_data="vidp_type_simeina")
    kb.button(text="Календарна (без збереження)", callback_data="vidp_type_kalendarna")
    kb.button(text="🔙 Назад", callback_data="vidp_back_to_persons")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)

    await callback.message.edit_text(
        f"👤 {pib}\n\nОберіть тип відпустки:",
        reply_markup=kb.as_markup())


@dp.callback_query(F.data.startswith("vidp_type_"))
async def vidp_select_leave_type(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    leave_type = callback.data.replace("vidp_type_", "")
    user_state[uid]["vidp_leave_type"] = leave_type
    await show_vidp_year_selection(callback, uid)
    await safe_answer_callback(callback)


# --- Year selection --- #

async def show_vidp_year_selection(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    current_year = datetime.now().year
    kb = InlineKeyboardBuilder()
    for y in range(current_year - 2, current_year + 2):
        kb.button(text=str(y), callback_data=f"vidp_year_{y}")
    kb.button(text="🔙 Назад", callback_data="vidp_back_to_type")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(3)
    await callback.message.edit_text("📅 Оберіть рік, за який надається відпустка:", reply_markup=kb.as_markup())


@dp.callback_query(F.data.startswith("vidp_year_"))
async def vidp_select_year(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    year = int(callback.data.replace("vidp_year_", ""))
    user_state[uid]["vidp_year"] = year
    await show_vidp_days_selection(callback, uid)
    await safe_answer_callback(callback)


# --- Days selection --- #

async def show_vidp_days_selection(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    kb = InlineKeyboardBuilder()
    preset_days = [5, 7, 10, 14, 15, 21, 24, 30]
    for d in preset_days:
        kb.button(text=str(d), callback_data=f"vidp_days_{d}")
    kb.button(text="✏️ Ввести вручну", callback_data="vidp_days_manual")
    kb.button(text="🔙 Назад", callback_data="vidp_back_to_year")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(4)
    await callback.message.edit_text(
        "🔢 Оберіть тривалість відпустки (календарних днів):",
        reply_markup=kb.as_markup())


@dp.callback_query(F.data.startswith("vidp_days_"))
async def vidp_select_days(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if callback.data == "vidp_days_manual":
        user_state[uid]["state"] = "vidp_enter_days"
        kb = InlineKeyboardBuilder()
        kb.button(text="🔙 Назад", callback_data="vidp_back_to_days")
        kb.adjust(1)
        await callback.message.edit_text(
            "✏️ Введіть кількість календарних днів (число):",
            reply_markup=kb.as_markup())
        await safe_answer_callback(callback)
        return

    days = int(callback.data.replace("vidp_days_", ""))
    user_state[uid]["vidp_days"] = days
    await show_vidp_start_calendar(callback, uid)
    await safe_answer_callback(callback)


async def show_vidp_start_calendar(callback: types.CallbackQuery, uid: int, year=None, month=None):
    touch_activity(uid)
    if year is None or month is None:
        now = datetime.now()
        year, month = now.year, now.month
    markup = create_vidp_start_calendar(year, month)
    mode = user_state[uid].get("vidp_mode", "vidpustka")
    if mode == "vidpustka":
        days = user_state[uid].get("vidp_days", 0)
        text = f"📅 Оберіть дату початку відпустки\nТривалість: {days} днів"
    elif mode == "sick":
        text = "📅 Оберіть дату початку лікарняного:"
    elif mode == "return":
        text = "📅 Оберіть дату, з якої особа приступила до виконання обов'язків:"
    else:
        text = "📅 Оберіть дату:"
    await _vidp_reply(callback, text, markup)


@dp.callback_query(F.data.startswith("vsd_"))
async def vidp_select_start_date(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if callback.data == "vsd_today":
        new_date = datetime.now()
    else:
        parts = callback.data.split("_")
        year, month, day = int(parts[1]), int(parts[2]), int(parts[3])
        new_date = datetime(year, month, day)

    mode = user_state[uid].get("vidp_mode", "vidpustka")

    if mode == "vidpustka":
        user_state[uid]["vidp_date_from"] = new_date
        # Auto-calculate end date
        days = user_state[uid].get("vidp_days", 1)
        end_date = new_date + timedelta(days=days - 1)
        user_state[uid]["vidp_date_to"] = end_date
        await show_vidp_destination(callback, uid)
    elif mode == "sick":
        user_state[uid]["vidp_date_from"] = new_date
        await show_vidp_supervisor_selection(callback, uid)
    elif mode == "return":
        user_state[uid]["vidp_date_return"] = new_date
        await show_vidp_confirm_person(callback, uid)

    await safe_answer_callback(callback)


@dp.callback_query(F.data.startswith("vsn_"))
async def vidp_start_navigate(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    parts = callback.data.split("_")
    year, month = int(parts[1]), int(parts[2])
    await show_vidp_start_calendar(callback, uid, year, month)
    await safe_answer_callback(callback)


# --- Destination --- #

async def show_vidp_destination(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    date_from = user_state[uid].get("vidp_date_from", datetime.now())
    date_to = user_state[uid].get("vidp_date_to", datetime.now())
    days = user_state[uid].get("vidp_days", 0)

    kb = InlineKeyboardBuilder()
    kb.button(text="м. Київ", callback_data="vidp_dest_kyiv")
    kb.button(text="м. Одеса", callback_data="vidp_dest_odesa")
    kb.button(text="м. Кременчук", callback_data="vidp_dest_kremenchuk")
    kb.button(text="м. Львів", callback_data="vidp_dest_lviv")
    kb.button(text="Пропустити", callback_data="vidp_dest_skip")
    kb.button(text="Ввести вручну", callback_data="vidp_dest_manual")
    kb.button(text="Назад", callback_data="vidp_back_to_start_date")
    kb.button(text="Рестарт", callback_data="restart")
    kb.adjust(2)

    text = (
        f"Перевірте дані:\n"
        f"З {format_date(date_from)}\n"
        f"По {format_date(date_to)}\n"
        f"Тривалість: {days} днів\n\n"
        f"Оберіть місце проведення відпустки (опціонально):"
    )
    await _vidp_reply(callback, text, kb.as_markup())


@dp.callback_query(F.data.startswith("vidp_dest_"))
async def vidp_select_destination(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if callback.data == "vidp_dest_skip":
        user_state[uid]["vidp_destination"] = ""
        user_state[uid]["vidp_travel_days"] = 0
        await show_vidp_confirm_person(callback, uid)
        await safe_answer_callback(callback)
        return
    elif callback.data == "vidp_dest_manual":
        user_state[uid]["state"] = "vidp_enter_destination"
        kb = InlineKeyboardBuilder()
        kb.button(text="⏭️ Пропустити", callback_data="vidp_dest_skip")
        kb.button(text="🔁 Рестарт", callback_data="restart")
        kb.adjust(1)
        await callback.message.edit_text(
            "✏️ Введіть місце проведення відпустки (наприклад: м. Київ):",
            reply_markup=kb.as_markup())
        await safe_answer_callback(callback)
        return
    else:
        dest_map = {
            "kyiv": "м. Київ",
            "odesa": "м. Одеса",
            "kremenchuk": "м. Кременчук, Полтавської області",
            "lviv": "м. Львів",
        }
        dest_code = callback.data.replace("vidp_dest_", "")
        user_state[uid]["vidp_destination"] = dest_map.get(dest_code, dest_code)
        print(f"[DEBUG] Destination selected: {user_state[uid]['vidp_destination']}")
    # After destination selection, ask for travel days ONLY if vacation > 14 days
    days = user_state[uid].get("vidp_days", 0)
    if days > 14:
        print(f"[DEBUG] Calling show_vidp_travel_days (days={days})...")
        await show_vidp_travel_days(callback, uid)
    else:
        user_state[uid]["vidp_travel_days"] = 0
        print(f"[DEBUG] Skipping travel days (days={days} <= 14)")
        await show_vidp_confirm_person(callback, uid)
    await safe_answer_callback(callback)


# --- Travel days menu --- #

async def show_vidp_travel_days(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    destination = user_state[uid].get("vidp_destination", "")
    if not destination:
        # No destination selected, skip travel days
        user_state[uid]["vidp_travel_days"] = 0
        await show_vidp_confirm_person(callback, uid)
        return
    print(f"[DEBUG] show_vidp_travel_days: dest={destination}")
    kb = InlineKeyboardBuilder()
    for d in range(5):
        label = f"{d} днів" if d != 1 else f"{d} день"
        if d == 0:
            label = "0 (без днів на дорогу)"
        kb.button(text=label, callback_data=f"vidp_travel_{d}")
    kb.button(text="🔙 Назад", callback_data="vidp_back_to_destination")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(3)
    text = f"🚗 Місце проведення: {destination}\n\nОберіть кількість днів на дорогу:"
    await _vidp_reply(callback, text, kb.as_markup())


@dp.callback_query(F.data.startswith("vidp_travel_"))
async def vidp_select_travel_days(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    print(f"[DEBUG] vidp_select_travel_days: data={callback.data}")
    travel = int(callback.data.replace("vidp_travel_", ""))
    user_state[uid]["vidp_travel_days"] = travel
    # Recalculate end date including travel days
    days = user_state[uid].get("vidp_days", 0)
    total = days + travel
    date_from = user_state[uid].get("vidp_date_from", datetime.now())
    date_to = date_from + timedelta(days=total - 1)
    user_state[uid]["vidp_date_to"] = date_to
    await show_vidp_confirm_person(callback, uid)
    await safe_answer_callback(callback)


@dp.callback_query(F.data == "vidp_back_to_destination")
async def vidp_back_to_destination(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_vidp_destination(callback, uid)
    await safe_answer_callback(callback)


# --- Supervisor selection for sick leave (like nachalnyky in engagement order) --- #

async def show_vidp_supervisor_selection(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    unit = user_state[uid].get("vidp_current_unit", "")
    # Same logic as show_nachalnyky_selection
    cached_vals = SHEET_CACHE.get("1_values")
    supervisors = []
    if cached_vals:
        values = cached_vals[0]
        data_rows = values[1:] if values else []
        allowed = {"начальник поста", "начальник частини", "заступник начальника частини",
                    "начальник", "заступник начальника", "т.в.о. начальника"}
        for r in data_rows:
            unit_val = (r[5].strip() if len(r) > 5 else "")
            cat_val = (r[3].strip().lower() if len(r) > 3 else "")
            if unit_val == unit and cat_val in allowed:
                col_o = (r[14].strip() if len(r) > 14 else "")
                col_p = (r[15].strip() if len(r) > 15 else "")
                col_q = (r[16].strip() if len(r) > 16 else "")
                pib_nom = (r[21].strip() if len(r) > 21 else "")  # column V for pidstava (родовий)
                display = ", ".join([p for p in [f"{col_o} {col_p}".strip(), col_q] if p])
                if display:
                    supervisors.append({"display": display, "pib": pib_nom, "p": col_p})
    user_state[uid]["vidp_supervisors"] = supervisors

    person = user_state[uid].get("vidp_current_person", {})
    pib = person.get("ПІБ", "")

    if not supervisors:
        user_state[uid]["vidp_rapport_pib"] = pib
        await show_vidp_confirm_person(callback, uid)
        return

    kb = InlineKeyboardBuilder()
    for i, s in enumerate(supervisors):
        label = s.get("p", s.get("display", ""))
        kb.button(text=label, callback_data=f"vidp_sup_{i}")
    kb.button(text="⏭️ Пропустити", callback_data="vidp_sup_skip")
    kb.button(text="🔙 Назад", callback_data="vidp_back_to_start_date_sick")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)

    text = f"👤 {pib}\nОберіть начальника (хто подає рапорт):"
    await _vidp_reply(callback, text, kb.as_markup())


@dp.callback_query(F.data == "vidp_back_to_start_date_sick")
async def vidp_back_to_start_date_sick(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_vidp_start_calendar(callback, uid)
    await safe_answer_callback(callback)


@dp.callback_query(F.data.startswith("vidp_sup_"))
async def vidp_select_supervisor(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    if callback.data == "vidp_sup_skip":
        person = user_state[uid].get("vidp_current_person", {})
        user_state[uid]["vidp_rapport_pib"] = person.get("ПІБ", "")
    else:
        idx = int(callback.data.replace("vidp_sup_", ""))
        supervisors = user_state[uid].get("vidp_supervisors", [])
        if 0 <= idx < len(supervisors):
            # Store the B column (nominative) name for pidstava
            user_state[uid]["vidp_rapport_pib"] = supervisors[idx].get("pib", "")
    await show_vidp_confirm_person(callback, uid)
    await safe_answer_callback(callback)


# --- Confirm person --- #

async def show_vidp_confirm_person(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    person = user_state[uid].get("vidp_current_person", {})
    mode = user_state[uid].get("vidp_mode", "vidpustka")

    if mode == "vidpustka":
        leave_type = user_state[uid].get("vidp_leave_type", "shchorichna")
        year = user_state[uid].get("vidp_year", datetime.now().year)
        days = user_state[uid].get("vidp_days", 0)
        date_from = user_state[uid].get("vidp_date_from", datetime.now())
        date_to = user_state[uid].get("vidp_date_to", datetime.now())
        destination = user_state[uid].get("vidp_destination", "")

        item_preview = {
            "pib": person.get("ПІБ", ""),
            "pib_dative": person.get("pib_dative", ""),
            "rank_dative": person.get("rank_dative", ""),
            "position_full": person.get("position_full", ""),
            "rank": person.get("rank_dative", ""),  # fallback for declension
            "position": person.get("position_full", ""),
            "unit": person.get("unit", ""),
            "leave_type": leave_type,
            "year": year,
            "days": days,
            "date_from": format_date(date_from),
            "date_to": format_date(date_to),
            "destination": destination,
            "travel_days": user_state[uid].get("vidp_travel_days", 0),
            "has_rank": person.get("has_rank", True),
        }
        text_item, text_pidstava = _format_vidpustky_item(item_preview, 1)
    elif mode == "sick":
        date_from = user_state[uid].get("vidp_date_from", datetime.now())
        ln_number = user_state[uid].get("vidp_ln_number", "")
        ln_date = user_state[uid].get("vidp_ln_date", "")

        item_preview = {
            "pib": person.get("ПІБ", ""),
            "pib_genitive": person.get("pib_genitive", ""),
            "rank_genitive": person.get("rank_genitive", ""),
            "position_genitive": person.get("position_genitive", ""),
            "rank": person.get("rank_genitive", ""),
            "position": person.get("position_genitive", ""),
            "unit": person.get("unit", ""),
            "date_from": format_date(date_from),
            "ln_number": ln_number,
            "ln_date": ln_date,
            "rapport_pib": user_state[uid].get("vidp_rapport_pib", ""),
        }
        text_item, text_pidstava = _format_sick_item(item_preview, 1)
    elif mode == "return":
        date_return = user_state[uid].get("vidp_date_return", datetime.now())

        item_preview = {
            "pib": person.get("ПІБ", ""),
            "pib_genitive": person.get("pib_genitive", ""),
            "rank_genitive": person.get("rank_genitive", ""),
            "position_genitive": person.get("position_genitive", ""),
            "rank": person.get("rank_genitive", ""),
            "position": person.get("position_genitive", ""),
            "unit": person.get("unit", ""),
            "date_return": format_date(date_return),
        }
        text_item, text_pidstava = _format_return_item(item_preview, 1)
    else:
        text_item = "Невідомий режим"
        text_pidstava = ""

    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Підтвердити та додати", callback_data="vidp_confirm_add")
    kb.button(text="❌ Скасувати", callback_data="vidp_back_to_menu")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)

    await _vidp_reply(callback,
        f"📋 ПЕРЕВІРКА ПУНКТУ:\n\n{text_item}\n\n{text_pidstava}",
        kb.as_markup())


@dp.callback_query(F.data == "vidp_confirm_add")
async def vidp_confirm_add(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    person = user_state[uid].get("vidp_current_person", {})
    mode = user_state[uid].get("vidp_mode", "vidpustka")

    item = {
        "pib": person.get("ПІБ", ""),
        "pib_dative": person.get("pib_dative", ""),
        "pib_genitive": person.get("pib_genitive", ""),
        "rank_dative": person.get("rank_dative", ""),
        "rank_genitive": person.get("rank_genitive", ""),
        "position_full": person.get("position_full", ""),
        "position_genitive": person.get("position_genitive", ""),
        "rank": person.get("rank_genitive", ""),
        "position": person.get("position_genitive", ""),
        "unit": person.get("unit", ""),
        "has_rank": person.get("has_rank", False),
    }

    if mode == "vidpustka":
        item["leave_type"] = user_state[uid].get("vidp_leave_type", "shchorichna")
        item["year"] = user_state[uid].get("vidp_year", datetime.now().year)
        item["days"] = user_state[uid].get("vidp_days", 0)
        item["date_from"] = format_date(user_state[uid].get("vidp_date_from", datetime.now()))
        item["date_to"] = format_date(user_state[uid].get("vidp_date_to", datetime.now()))
        item["destination"] = user_state[uid].get("vidp_destination", "")
        item["travel_days"] = user_state[uid].get("vidp_travel_days", 0)
        user_state[uid].setdefault("vidpustky_items", []).append(item)
    elif mode == "sick":
        item["date_from"] = format_date(user_state[uid].get("vidp_date_from", datetime.now()))
        item["ln_number"] = user_state[uid].get("vidp_ln_number", "")
        item["ln_date"] = user_state[uid].get("vidp_ln_date", "")
        item["rapport_pib"] = user_state[uid].get("vidp_rapport_pib", "")
        item["extra_notes"] = user_state[uid].get("vidp_extra_notes", "")
        user_state[uid].setdefault("vidpustky_sick_items", []).append(item)
    elif mode == "return":
        item["date_return"] = format_date(user_state[uid].get("vidp_date_return", datetime.now()))
        item["rapport_pib"] = user_state[uid].get("vidp_rapport_pib", "")
        user_state[uid].setdefault("vidpustky_return_items", []).append(item)

    # Clear temporary fields
    for k in ["vidp_current_person", "vidp_leave_type", "vidp_year", "vidp_days",
              "vidp_date_from", "vidp_date_to", "vidp_destination",
              "vidp_ln_number", "vidp_ln_date", "vidp_rapport_pib",
              "vidp_date_return", "vidp_extra_notes", "vidp_mode",
              "vidp_persons_list", "vidp_current_unit"]:
        user_state[uid].pop(k, None)

    user_state[uid]["state"] = None
    await safe_answer_callback(callback, "✅ Особу додано!")
    await show_vidp_main_menu(callback, uid)


# --- Sick leave flow --- #

@dp.callback_query(F.data == "vidp_add_sick")
async def vidp_add_sick(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state[uid]["vidp_mode"] = "sick"
    await show_vidp_unit_selection(callback, uid)
    await safe_answer_callback(callback)


async def show_vidp_sick_date(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    person = user_state[uid].get("vidp_current_person", {})
    pib = person.get("ПІБ", "")

    # Show simple date selection
    now = datetime.now()
    markup = create_vidp_start_calendar(now.year, now.month)
    await callback.message.edit_text(
        f"👤 {pib}\n📅 Оберіть дату початку лікарняного:",
        reply_markup=markup)


@dp.callback_query(F.data == "vidp_back_to_days")
async def vidp_back_to_days(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    mode = user_state[uid].get("vidp_mode", "vidpustka")
    if mode == "vidpustka":
        await show_vidp_days_selection(callback, uid)
    else:
        await show_vidp_main_menu(callback, uid)
    await safe_answer_callback(callback)


# --- LN number input (for sick leave) --- #

async def show_vidp_ln_input(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    user_state[uid]["state"] = "vidp_enter_ln"
    kb = InlineKeyboardBuilder()
    kb.button(text="⏭️ Пропустити", callback_data="vidp_ln_skip")
    kb.button(text="🔙 Назад", callback_data="vidp_back_to_start_date")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    await callback.message.edit_text(
        "🔢 Введіть номер листка непрацездатності (ЛН) або натисніть Пропустити:",
        reply_markup=kb.as_markup())


@dp.callback_query(F.data == "vidp_ln_skip")
async def vidp_ln_skip(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state[uid]["vidp_ln_number"] = ""
    user_state[uid]["vidp_ln_date"] = ""
    user_state[uid]["state"] = None
    await show_vidp_supervisor_selection(callback, uid)
    await safe_answer_callback(callback)

# Add handler for start date selection in sick/return mode
# Override: when in sick mode, after date selection go to LN input
# We handle this in vidp_select_start_date by checking mode


# --- Return from sick leave flow --- #

@dp.callback_query(F.data == "vidp_add_return")
async def vidp_add_return(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state[uid]["vidp_mode"] = "return"
    await show_vidp_unit_selection(callback, uid)
    await safe_answer_callback(callback)


async def show_vidp_return_date(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    person = user_state[uid].get("vidp_current_person", {})
    pib = person.get("ПІБ", "")
    now = datetime.now()
    markup = create_vidp_start_calendar(now.year, now.month)
    await callback.message.edit_text(
        f"👤 {pib}\n📅 Оберіть дату, з якої особа приступила до виконання обов'язків:",
        reply_markup=markup)

# Override start date selection for sick/return modes
# The existing vidp_select_start_date handler auto-calculates end date for vidpustka.
# For sick mode, date_from is the sick leave start date.
# For return mode, date_from should be mapped to date_return.
# We handle this by checking mode inside the handler.


# --- Show list / Delete --- #

@dp.callback_query(F.data == "vidp_show_list")
async def vidp_show_list(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)

    items1 = user_state[uid].get("vidpustky_items", [])
    items2 = user_state[uid].get("vidpustky_sick_items", [])
    items3 = user_state[uid].get("vidpustky_return_items", [])

    if not items1 and not items2 and not items3:
        await safe_answer_callback(callback, "Список порожній!", show_alert=True)
        await show_vidp_main_menu(callback, uid)
        return

    text_lines = []
    kb = InlineKeyboardBuilder()

    if items1:
        text_lines.append("=== §1 НАДАТИ ===")
        for i, item in enumerate(items1):
            item_preview = {
                "pib": item.get("pib", ""),
                "pib_dative": item.get("pib_dative", ""),
                "rank_dative": item.get("rank_dative", ""),
                "position_full": item.get("position_full", ""),
                "rank": item.get("rank_dative", ""),
                "position": item.get("position_full", ""),
                "unit": item.get("unit", ""),
                "leave_type": item.get("leave_type", "shchorichna"),
                "year": item.get("year", ""),
                "days": item.get("days", ""),
                "date_from": item.get("date_from", ""),
                "date_to": item.get("date_to", ""),
                "destination": item.get("destination", ""),
                "travel_days": item.get("travel_days", 0),
                "has_rank": item.get("has_rank", True),
            }
            txt, _ = _format_vidpustky_item(item_preview, i + 1)
            text_lines.append(txt)
            kb.button(text=f"❌ Видалити відп. #{i+1}", callback_data=f"vidp_del_vidp_{i}")

    if items2:
        text_lines.append("\n=== §2 ЗВІЛЬНИТИ (лікарняні) ===")
        for i, item in enumerate(items2):
            txt, _ = _format_sick_item(item, i + 1)
            text_lines.append(txt)
            kb.button(text=f"❌ Видалити лік. #{i+1}", callback_data=f"vidp_del_sick_{i}")

    if items3:
        text_lines.append("\n=== §3 ПРИСТУПИЛИ ===")
        for i, item in enumerate(items3):
            txt, _ = _format_return_item(item, i + 1)
            text_lines.append(txt)
            kb.button(text=f"❌ Видалити пов. #{i+1}", callback_data=f"vidp_del_ret_{i}")

    kb.button(text="🔙 Назад", callback_data="vidp_back_to_menu")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)

    full_text = "\n\n".join(text_lines)
    if len(full_text) > 4000:
        full_text = full_text[:4000] + "\n\n... (текст скорочено)"

    await callback.message.edit_text(full_text, reply_markup=kb.as_markup())
    await safe_answer_callback(callback)


@dp.callback_query(F.data.startswith("vidp_del_"))
async def vidp_delete_item(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    parts = callback.data.split("_")
    item_type = parts[2]  # vidp, sick, ret
    idx = int(parts[3])

    if item_type == "vidp":
        items = user_state[uid].get("vidpustky_items", [])
        if 0 <= idx < len(items):
            del items[idx]
            user_state[uid]["vidpustky_items"] = items
    elif item_type == "sick":
        items = user_state[uid].get("vidpustky_sick_items", [])
        if 0 <= idx < len(items):
            del items[idx]
            user_state[uid]["vidpustky_sick_items"] = items
    elif item_type == "ret":
        items = user_state[uid].get("vidpustky_return_items", [])
        if 0 <= idx < len(items):
            del items[idx]
            user_state[uid]["vidpustky_return_items"] = items

    await safe_answer_callback(callback, "✅ Видалено")
    await vidp_show_list(callback)


# --- Signatory selection --- #

@dp.callback_query(F.data == "vidp_next_to_sign")
async def vidp_next_to_sign(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    total = (len(user_state[uid].get("vidpustky_items", [])) +
             len(user_state[uid].get("vidpustky_sick_items", [])) +
             len(user_state[uid].get("vidpustky_return_items", [])))
    if total == 0:
        await safe_answer_callback(callback, "Додайте хоча б одну особу!", show_alert=True)
        return
    await show_vidp_signatory(callback, uid)
    await safe_answer_callback(callback)


async def show_vidp_signatory(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    kb = InlineKeyboardBuilder()
    kb.button(text="Коржов", callback_data="vidp_pidpys_korzhov")
    kb.button(text="Усик", callback_data="vidp_pidpys_usyk")
    kb.button(text="Середа", callback_data="vidp_pidpys_sereda")
    kb.button(text="🔙 Назад", callback_data="vidp_back_to_menu")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)
    await callback.message.edit_text("Хто підписує наказ?", reply_markup=kb.as_markup())


@dp.callback_query(F.data.startswith("vidp_pidpys_"))
async def vidp_select_signatory(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    pidpys_key = callback.data.replace("vidp_pidpys_", "")
    user_state[uid]["vidp_pidpys"] = pidpys_key
    await safe_answer_callback(callback, "Підписанта обрано")
    await show_vidp_generate_prompt(callback, uid)


async def show_vidp_generate_prompt(callback: types.CallbackQuery, uid: int):
    touch_activity(uid)
    n_vidp = len(user_state[uid].get("vidpustky_items", []))
    n_sick = len(user_state[uid].get("vidpustky_sick_items", []))
    n_ret = len(user_state[uid].get("vidpustky_return_items", []))
    pidpys = user_state[uid].get("vidp_pidpys", "")

    pidpys_names = {"korzhov": "Коржов", "usyk": "Усик", "sereda": "Середа"}
    pidpys_display = pidpys_names.get(pidpys, pidpys)

    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Сформувати наказ", callback_data="vidp_generate")
    kb.button(text="🔙 Назад", callback_data="vidp_next_to_sign")
    kb.button(text="🔁 Рестарт", callback_data="restart")
    kb.adjust(1)

    text = (f"📜 ГОТОВО ДО ГЕНЕРАЦІЇ:\n"
            f"📅 Дата наказу: {user_state[uid].get('vidp_data_nakazu', datetime.now()).strftime('%d.%m.%Y')}\n"
            f"👥 Відпустки: {n_vidp} | Лікарняні: {n_sick} | Повернення: {n_ret}\n"
            f"✍️ Підписант: {pidpys_display}\n\n"
            f"Натисніть кнопку для генерації.")
    await callback.message.edit_text(text, reply_markup=kb.as_markup())


# --- Navigation handlers --- #

@dp.callback_query(F.data == "vidp_back_to_menu")
async def vidp_back_to_menu(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    user_state[uid]["state"] = None
    await show_vidp_main_menu(callback, uid)
    await safe_answer_callback(callback)


@dp.callback_query(F.data == "vidp_back_to_units")
async def vidp_back_to_units(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_vidp_unit_selection(callback, uid)
    await safe_answer_callback(callback)


@dp.callback_query(F.data == "vidp_back_to_persons")
async def vidp_back_to_persons(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_vidp_person_selection(callback, uid)
    await safe_answer_callback(callback)


@dp.callback_query(F.data == "vidp_back_to_type")
async def vidp_back_to_type(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_vidp_leave_type(callback, uid)
    await safe_answer_callback(callback)


@dp.callback_query(F.data == "vidp_back_to_year")
async def vidp_back_to_year(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_vidp_year_selection(callback, uid)
    await safe_answer_callback(callback)


@dp.callback_query(F.data == "vidp_back_to_start_date")
async def vidp_back_to_start_date(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await show_vidp_start_calendar(callback, uid)
    await safe_answer_callback(callback)


# ======================================================================== #
# ===  GENERATE VIDPUSTKY DOCUMENT  ====================================== #
# ======================================================================== #

def generate_nakaz_vidpustky(user_state: dict) -> tuple:
    """Генерує наказ відпустки. Повертає (docx_path, pdf_path)."""
    template_path = os.path.join(BASE_DIR, "nakaz_vidpustky.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: nakaz_vidpustky.docx")

    doc = Document(template_path)

    # 1. Replace date placeholders
    data_nakazu = user_state.get("vidp_data_nakazu", datetime.now())
    data_str = format_date(data_nakazu)
    data_short = data_nakazu.strftime("%d") + uk_month_gen[data_nakazu.month - 1] + f" {data_nakazu.year} року"

    def replace_in_paragraph(paragraph, placeholders):
        try:
            text = paragraph.text or ""
            changed = False
            for key, value in placeholders.items():
                if key in text:
                    text = text.replace(key, str(value) if value is not None else "")
                    changed = True
            if changed:
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = text
                else:
                    p_run = paragraph.add_run(text)
                    p_run.font.name = 'Times New Roman'
                    p_run.font.size = Pt(14)
        except Exception:
            pass

    placeholders = {
        "{data_nakazu_short}": data_short,
        "{data_nakazu}": data_str,
    }

    # Apply to all paragraphs, tables, headers, footers
    for p in doc.paragraphs:
        replace_in_paragraph(p, placeholders)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, placeholders)
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph(p, placeholders)
        for p in section.footer.paragraphs:
            replace_in_paragraph(p, placeholders)

    # 2. Find {items_start} marker and insert dynamic content
    items_start_idx = None
    nadaty_idx = None
    para1_idx = None  # §1 paragraph index
    for i, p in enumerate(doc.paragraphs):
        if "{items_start}" in p.text:
            items_start_idx = i
            # Clear the marker paragraph
            for run in p.runs:
                run.text = ""
        if p.text.strip() == "§1":
            para1_idx = i
        if p.text.strip() == "НАДАТИ:":
            nadaty_idx = i

    if items_start_idx is None:
        raise ValueError("Marker {items_start} not found in template!")

    # 3. Build and insert §1 items (vidpustky)
    vidpustky_items = user_state.get("vidpustky_items", [])
    section2_items = user_state.get("vidpustky_sick_items", [])
    section3_items = user_state.get("vidpustky_return_items", [])

    # Collect unit names for distribution
    all_units = set()

    def insert_text_paragraph_after(ref_para, text, font_name='Times New Roman', font_size=14, bold=False, italic=False, indent=True):
        """Insert a text paragraph with formatting after ref_para (which is an lxml element)."""
        new_p_elem = OxmlElement('w:p')
        ref_para.addnext(new_p_elem)
        # Add paragraph properties (first line indent for items)
        pPr = OxmlElement('w:pPr')
        if indent:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:firstLine'), '720')  # ~1.25cm indent
            pPr.append(ind)
        # Set font to Times New Roman 14pt
        rPrDefault = OxmlElement('w:rPrDefault')
        rPr = OxmlElement('w:rPr')
        rFonts_def = OxmlElement('w:rFonts')
        rFonts_def.set(qn('w:ascii'), font_name)
        rFonts_def.set(qn('w:hAnsi'), font_name)
        rFonts_def.set(qn('w:cs'), font_name)
        rPr.append(rFonts_def)
        sz_def = OxmlElement('w:sz')
        sz_def.set(qn('w:val'), str(font_size * 2))
        rPr.append(sz_def)
        rPrDefault.append(rPr)
        pPr.append(rPrDefault)
        new_p_elem.append(pPr)
        # Create run
        r_elem = OxmlElement('w:r')
        rpr_elem = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rpr_elem.append(rFonts)
        if bold:
            b = OxmlElement('w:b')
            rpr_elem.append(b)
        if italic:
            i_elem = OxmlElement('w:i')
            rpr_elem.append(i_elem)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size * 2))
        rpr_elem.append(sz)
        r_elem.append(rpr_elem)
        t_elem = OxmlElement('w:t')
        t_elem.text = text
        t_elem.set(qn('xml:space'), 'preserve')
        r_elem.append(t_elem)
        new_p_elem.append(r_elem)
        return new_p_elem

    # Start inserting after items_start
    last_element = doc.paragraphs[items_start_idx]._element

    # --- §1: НАДАТИ ---
    if vidpustky_items:
        for idx, item in enumerate(vidpustky_items):
            all_units.add(item.get("unit", ""))
            item_preview = {
                "pib": item.get("pib", ""),
                "pib_dative": item.get("pib_dative", ""),
                "rank_dative": item.get("rank_dative", ""),
                "position_full": item.get("position_full", ""),
                "rank": item.get("rank_dative", ""),
                "position": item.get("position_full", ""),
                "unit": item.get("unit", ""),
                "leave_type": item.get("leave_type", "shchorichna"),
                "year": item.get("year", ""),
                "days": item.get("days", ""),
                "date_from": item.get("date_from", ""),
                "date_to": item.get("date_to", ""),
                "destination": item.get("destination", ""),
                "travel_days": item.get("travel_days", 0),
                "has_rank": item.get("has_rank", True),
            }
            text_item, text_pidstava = _format_vidpustky_item(item_preview, idx + 1)
            # Insert item text (with indent)
            insert_text_paragraph_after(last_element, text_item, indent=True)
            # Move to the new paragraph
            last_element = last_element.getnext()
            # Insert pidstava (no indent)
            insert_text_paragraph_after(last_element, text_pidstava, indent=True)
            last_element = last_element.getnext()
            # Blank line between items (except last in section)
            if idx < len(vidpustky_items) - 1:
                insert_text_paragraph_after(last_element, "", indent=False)
                last_element = last_element.getnext()

    # --- §2: ЗВІЛЬНИТИ (лікарняні) ---
    has_section2 = bool(section2_items)

    # Find {section2_start} marker
    # We need to handle section markers: if section is empty, remove all paragraphs
    # between section_start and section_end (inclusive). Otherwise, fill content.

    def fill_section(section_start_marker: str, section_end_marker: str, items_list: list, format_func):
        """Fill a section with items or remove it if empty."""
        # Find markers
        start_idx = None
        end_idx = None
        for i, p in enumerate(doc.paragraphs):
            if section_start_marker in p.text:
                start_idx = i
            if section_end_marker in p.text and start_idx is not None:
                end_idx = i
                break

        if start_idx is None:
            return  # Section not in template

        if not items_list:
            # Remove everything from section_start to section_end (inclusive)
            # Remove backwards to avoid index shifting
            for _ in range(end_idx - start_idx + 1):
                p_to_remove = doc.paragraphs[start_idx]
                p_to_remove._element.getparent().remove(p_to_remove._element)
            return

        # Clear markers
        for i, p in enumerate(doc.paragraphs):
            if section_start_marker in p.text:
                for run in p.runs:
                    run.text = ""
            if section_end_marker in p.text:
                for run in p.runs:
                    run.text = ""

        # Find the {sectionX_items} marker (the paragraph where items should go)
        items_marker_idx = None
        for i, p in enumerate(doc.paragraphs):
            # Extract section number from marker like '{section2_start}'
            section_num = section_start_marker.split("section")[1].split("_")[0]
            if f"{{section{section_num}_items}}" in p.text:
                items_marker_idx = i
                # Clear the marker
                for run in p.runs:
                    run.text = ""
                break

        if items_marker_idx is None:
            return

        ref_para = doc.paragraphs[items_marker_idx]
        last_el = ref_para._element

        for idx, item in enumerate(items_list):
            all_units.add(item.get("unit", ""))
            text_item, text_pidstava = format_func(item, idx + 1)
            insert_text_paragraph_after(last_el, text_item, indent=True)
            last_el = last_el.getnext()
            insert_text_paragraph_after(last_el, text_pidstava, indent=True)
            last_el = last_el.getnext()
            if idx < len(items_list) - 1:
                insert_text_paragraph_after(last_el, "", indent=False)
                last_el = last_el.getnext()

    # Fill sections
    if section2_items:
        all_units.update(item.get("unit", "") for item in section2_items)
    if section3_items:
        all_units.update(item.get("unit", "") for item in section3_items)

    # Process §2
    fill_section("{section2_start}", "{section2_end}", section2_items, _format_sick_item)

    # Process §3
    fill_section("{section3_start}", "{section3_end}", section3_items, _format_return_item)

    # --- Renumber sections dynamically --- #
    section_number = 1
    for marker in ["§1", "§2", "§3"]:
        for p in doc.paragraphs:
            if p.text.strip() == marker:
                new_text = f"§{section_number}"
                for run in p.runs:
                    if marker in run.text:
                        run.text = run.text.replace(marker, new_text)
                        break
                else:
                    # No runs, set paragraph text
                    if p.runs:
                        p.runs[0].text = new_text
                section_number += 1
                break

    # --- Post-processing: remove empty sections, center-align headers --- #
    section_headers = {"§1", "НАДАТИ:", "§2", "ЗВІЛЬНИТИ від виконання службових обов'язків у зв'язку із тимчасовою непрацездатністю:",
                       "§3", "ВВАЖАТИ такими, що приступили до виконання службових обов'язків після тимчасової непрацездатності:"}

    # Remove §1 header if no vidpustky items
    if not vidpustky_items:
        # Find and remove "§1" and "НАДАТИ:" paragraphs
        for target_text in ["§1", "НАДАТИ:"]:
            for p in doc.paragraphs:
                if p.text.strip() == target_text:
                    p._element.getparent().remove(p._element)
                    break

    # Align section headers
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "НАДАТИ:" or t.startswith("ЗВІЛЬНИТИ") or t.startswith("ВВАЖАТИ"):
            try:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
            except Exception:
                pass
        elif t in section_headers:
            try:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
            except Exception:
                pass

    # 4. Fill signatory
    pidpys_key = user_state.get("vidp_pidpys", "")
    pidpys_names = {
        "korzhov": "Ігор КОРЖОВ",
        "usyk": "Олександр УСИК",
        "sereda": "Євген СЕРЕДА",
    }
    pidpys_name = pidpys_names.get(pidpys_key, "")

    sign_placeholders = {
        "{pidpys_name}": pidpys_name if pidpys_key else "",
        "{pidpys_sereda_name}": "Євген СЕРЕДА" if pidpys_key != "sereda" else "",
        "{pidpys_usyk_name}": "Олександр УСИК" if pidpys_key != "usyk" else "",
    }

    for p in doc.paragraphs:
        replace_in_paragraph(p, sign_placeholders)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, sign_placeholders)

    # 5. Fill distribution
    if all_units:
        # Extract unit numbers for ASKOD format (e.g., "1 ДПРЧ", "2 ДПРЧ")
        unit_parts = []
        for u in sorted(all_units):
            # Extract just the unit code like "1 ДПРЧ"
            match = re.match(r"^(\d+\s*(ДПРЧ|ДПРП|ДПРЗ))\b", u, re.IGNORECASE)
            if match:
                unit_parts.append(match.group(1))
            else:
                unit_parts.append(u)
        askod_text = ", ".join(unit_parts)
    else:
        askod_text = ""

    for p in doc.paragraphs:
        replace_in_paragraph(p, {"{pidrozdily_askod}": askod_text})

    # 6. Apply protect_no_wrap to all paragraphs
    apply_protect_no_wrap_document(doc)

    # 7. Disable auto hyphenation
    try:
        disable_auto_hyphenation_document(doc)
    except Exception:
        pass

    # 8. Save document
    os.makedirs("generated_orders", exist_ok=True)
    filepath = os.path.join("generated_orders", "Наказ_відпустки.docx")
    pdf_filepath = os.path.join("generated_orders", "Наказ_відпустки.pdf")

    doc.save(filepath)

    # 9. Convert to PDF
    pdf_generated = False
    if os.path.exists(LIBREOFFICE_PATH):
        try:
            subprocess.run([
                LIBREOFFICE_PATH,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', 'generated_orders',
                filepath
            ], check=True, timeout=45)
            pdf_generated = True
        except subprocess.TimeoutExpired:
            _log_debug("[PDF] Conversion timed out for vidpustky document")
        except Exception as e:
            print(f"PDF conversion failed: {e}")

    return filepath, pdf_filepath if pdf_generated else None


# --- Generate handler --- #

@dp.callback_query(F.data == "vidp_generate")
async def vidp_generate(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)

    total = (len(user_state[uid].get("vidpustky_items", [])) +
             len(user_state[uid].get("vidpustky_sick_items", [])) +
             len(user_state[uid].get("vidpustky_return_items", [])))
    if total == 0:
        await safe_answer_callback(callback, "Немає осіб для генерації!", show_alert=True)
        return

    await safe_answer_callback(callback, "Початок генерації наказу...")

    loading_msg = await callback.message.answer("Генерація файлів…")
    stop_event = asyncio.Event()
    spinner_task = asyncio.create_task(_spinner_message(loading_msg, "Генерація файлів…", stop_event))

    try:
        filename, pdf_filename = await asyncio.to_thread(generate_nakaz_vidpustky, user_state[uid])

        stop_event.set()
        try:
            await spinner_task
        except Exception:
            pass

        await loading_msg.edit_text("Надсилання файлів… 📤")

        caption = "📜 Наказ відпустки/лікарняні"
        sent_doc = await callback.message.answer_document(
            types.FSInputFile(filename),
            caption=caption
        )
        safe_remove(filename)

        if pdf_filename:
            sent_pdf = await callback.message.answer_document(
                types.FSInputFile(pdf_filename),
                caption=caption
            )
            safe_remove(pdf_filename)
        else:
            await callback.message.answer("⚠️ PDF не згенеровано. Перевірте LibreOffice.")

        clear_workflow_state(uid)
        await loading_msg.edit_text("Готово ✅")
        await asyncio.sleep(0.6)
        await loading_msg.delete()

    except Exception as e:
        stop_event.set()
        try:
            await spinner_task
        except Exception:
            pass
        await loading_msg.edit_text(f"❌ Помилка генерації: {e}")

    await show_post_generate_menu(callback, uid)


# ======================================================================== #
# ===  END VIDPUSTKY  =================================================== #
# ======================================================================== #


# ======================================================================== #
# ===  ОНОВИТИ О/С (AI відмінювання ПІБ)  =============================== #
# ======================================================================== #


async def process_declension_update(progress_callback=None):
    """Основний процес: читає колонку B, відмінює ПО ОДНОМУ через AI, ОДРАЗУ записує в H/P/U/V."""
    import traceback

    # Connect to sheet with WRITE access
    creds = Credentials.from_service_account_file(CREDENTIALS_FILE, scopes=[
        "https://www.googleapis.com/auth/spreadsheets"
    ])
    client = gspread.authorize(creds)
    sheet = client.open_by_url(GOOGLE_SHEET_URL1).sheet1

    # Read all data
    values = sheet.get_all_values()
    if not values or len(values) < 2:
        return {"error": "Таблиця порожня"}

    data_rows = values[1:]  # Skip header

    # Find names in column B
    tasks = []  # [(row_number_1based, name), ...]
    for i, row in enumerate(data_rows):
        if len(row) > 1:
            name = row[1].strip()
            if name and len(name.split()) >= 2:
                tasks.append((i + 2, name))  # +2 = 1-indexed + header

    total = len(tasks)
    if total == 0:
        return {"error": "Немає ПІБ у колонці B"}

    api_keys = _get_gemini_api_keys()
    if not api_keys:
        return {"error": "Немає API ключів Gemini"}

    # Track which keys are rate-limited and when they reset
    key_cooldown = {}  # key_index -> time when usable again

    if progress_callback:
        await progress_callback(f"Знайдено {total} ПІБ. Ключів: {len(api_keys)}. Починаю...")

    processed = 0
    updated = 0
    errors = 0

    # First try AI, fall back to built-in declension
    # Try AI for first 3 names to check if keys work
    use_ai = True
    api_keys = _get_gemini_api_keys()
    ai_failures = 0

    if progress_callback:
        await progress_callback(f"Знайдено {total} ПІБ. Починаю...")

    processed = 0
    updated = 0
    errors = 0

    for idx, (row_num, name) in enumerate(tasks):
        result = {}

        # Try AI first (if keys available)
        if use_ai and api_keys and ai_failures < 1:
            key_num = idx % len(api_keys)
            api_key = api_keys[key_num]
            result = await asyncio.to_thread(_call_gemini_decline_one, name, api_key)
            if not result:
                ai_failures += 1
                if ai_failures >= 1:
                    use_ai = False
                    if progress_callback:
                        await progress_callback("AI недоступний, використовую вбудоване відмінювання...")

        # Fall back to built-in declension if AI failed
        if not result:
            parts = name.split()
            if len(parts) >= 2:
                try:
                    # Get full declined forms
                    decl_dat = decline_name_dative(name)   # "ПРІЗВИЩЕ Ім'я По-батькові" in dative
                    decl_gen = decline_name_genitive(name)  # "ПРІЗВИЩЕ Ім'я По-батькові" in genitive
                    parts_dat = decl_dat.split()
                    parts_gen = decl_gen.split()

                    # U: АНДРОСЕНКУ Ігорю Сергійовичу (full dative, surname UPPER)
                    u_val = decl_dat
                    if parts_dat:
                        parts_dat[0] = parts_dat[0].upper()
                        u_val = " ".join(parts_dat)

                    # V: АНДРОСЕНКА Ігоря Сергійовича (full genitive, surname UPPER)
                    v_val = decl_gen
                    if parts_gen:
                        parts_gen[0] = parts_gen[0].upper()
                        v_val = " ".join(parts_gen)

                    # H: Ігоря АНДРОСЕНКА (short genitive = name_gen + surname_gen UPPER)
                    name_gen = parts_gen[1] if len(parts_gen) > 1 else parts_gen[0]
                    surname_gen = parts_gen[0].upper() if parts_gen else ""
                    h_val = f"{name_gen} {surname_gen}"

                    # P: Ігорю АНДРОСЕНКУ (short dative = name_dat + surname_dat UPPER)
                    name_dat = parts_dat[1] if len(parts_dat) > 1 else parts_dat[0]
                    surname_dat = parts_dat[0].upper() if parts_dat else ""
                    p_val = f"{name_dat} {surname_dat}"

                    result = {"H": h_val, "P": p_val, "U": u_val, "V": v_val}
                except Exception as e:
                    print(f"[FALLBACK ERR] {name}: {e}")
                    result = {}

        if result:
            # Write to sheet with retry on 429
            for retry in range(3):
                try:
                    if result.get("H"):
                        await asyncio.to_thread(sheet.update_cell, row_num, 8, result["H"])
                        await asyncio.sleep(0.5)
                    if result.get("P"):
                        await asyncio.to_thread(sheet.update_cell, row_num, 16, result["P"])
                        await asyncio.sleep(0.5)
                    if result.get("U"):
                        await asyncio.to_thread(sheet.update_cell, row_num, 21, result["U"])
                        await asyncio.sleep(0.5)
                    if result.get("V"):
                        await asyncio.to_thread(sheet.update_cell, row_num, 22, result["V"])
                    updated += 1
                    print(f"[OK] Row {row_num}: {name} | H={result.get('H','')} | P={result.get('P','')} | U={result.get('U','')} | V={result.get('V','')}")
                    break
                except Exception as e:
                    err = str(e)
                    if "429" in err and retry < 2:
                        await asyncio.sleep(15)  # wait 15s on rate limit
                    else:
                        errors += 1
                        print(f"[WRITE ERR] Row {row_num}: {err[:100]}")
                        break
        else:
            errors += 1
            print(f"[FAIL] Row {row_num}: {name}")

        processed += 1

        if progress_callback and (processed % 5 == 0 or processed == total):
            await progress_callback(
                f"Обробка: {processed}/{total} | Оновлено: {updated} | Помилок: {errors}"
            )

        await asyncio.sleep(1)  # 2s per row to stay under 60/min write limit

    # Clear cache
    global SHEET_CACHE
    SHEET_CACHE.pop("1", None)
    SHEET_CACHE.pop("1_values", None)

    return {"total": total, "processed": processed, "updated": updated, "errors": errors}


# --- Bot handler --- #

@dp.callback_query(F.data == "update_os")
async def update_os_start(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await safe_answer_callback(callback)

    # Confirmation screen
    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Так, оновити", callback_data="update_os_confirm")
    kb.button(text="❌ Скасувати", callback_data="back_to_start")
    kb.adjust(1)

    text = (
        "🔄 ОНОВЛЕННЯ ОСОБОВОГО СКЛАДУ\n\n"
        "Ця функція зчитає ПІБ з колонки B Google Sheets,\n"
        "за допомогою ШІ (Gemini) згенерує відмінки:\n"
        "• H — коротка форма (родовий)\n"
        "• P — коротка форма (давальний)\n"
        "• U — повна форма (давальний)\n"
        "• V — повна форма (родовий)\n\n"
        "⚠️ Процес може тривати кілька хвилин.\n"
        "Кількість запитів залежить від кількості рядків."
    )
    await callback.message.edit_text(text, reply_markup=kb.as_markup())


@dp.callback_query(F.data == "update_os_confirm")
async def update_os_confirm(callback: types.CallbackQuery):
    uid = callback.from_user.id
    touch_activity(uid)
    await safe_answer_callback(callback, "Запускаю обробку...")

    loading_msg = await callback.message.answer("🔄 Підготовка...")

    async def progress(msg):
        try:
            await loading_msg.edit_text(f"🔄 {msg}")
        except Exception:
            pass

    # Run the update process
    try:
        result = await process_declension_update(progress_callback=progress)
    except Exception as e:
        result = {"error": str(e)}

    if "error" in result:
        await loading_msg.edit_text(f"❌ Помилка: {result['error']}")
    else:
        await loading_msg.edit_text(
            f"✅ ОНОВЛЕННЯ ЗАВЕРШЕНО!\n\n"
            f"Всього ПІБ: {result['total']}\n"
            f"Оброблено: {result['processed']}\n"
            f"Оновлено рядків: {result['updated']}\n"
            f"Колонки: H, P, U, V"
        )

    # Back button
    kb = InlineKeyboardBuilder()
    kb.button(text="🔙 В меню", callback_data="back_to_start")
    kb.adjust(1)
    await callback.message.answer("Оберіть дію:", reply_markup=kb.as_markup())


# ======================================================================== #
# ===  END ОНОВИТИ О/С  ================================================ #
# ======================================================================== #

# --- Запуск --- #
async def main():
    print("🚀 Скрипт запущено в режимі підвищеної стабільності.")
    
    # Спочатку повністю прогріваємо кеш (Google Sheets → пам'ять)
    try:
        await _warm_cache_once()
    except Exception as e:
        print(f"⚠️ Помилка прогріву кешу: {e}")

    # Цикл "невмирущого" бота
    while True:
        current_bot = None
        try:
            current_bot = create_bot()
            print("📡 Підключення до Telegram встановлено. Очікування повідомлень...")
            await dp.start_polling(current_bot, handle_signals=False)
        except Exception as e:
            print(f"🆘 КРИТИЧНИЙ ЗБІЙ: {e}")
            print("🔄 Спроба повного перезавантаження через 10 секунд...")
        finally:
            if current_bot and current_bot.session:
                try:
                    await current_bot.session.close()
                    print("🧹 Стару сесію закрито.")
                except Exception:
                    pass
            await asyncio.sleep(10)

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n🛑 Бот зупинений користувачем.")
    except Exception as e:
        print(f"\n❌ Непередбачувана помилка: {e}")